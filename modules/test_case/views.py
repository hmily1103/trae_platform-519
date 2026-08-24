#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试用例管理 - Flask Views
"""

import os
from datetime import datetime
import uuid
import json
import requests
from flask import Blueprint, render_template, request, jsonify, send_file
from utils.response import success_response, error_response
from utils.logger import setup_logger
from .storage import TestCaseStorage
from .models import TestCase, TestSuite, TestCaseExecution, TestStep, PromptConfig

test_case_bp = Blueprint('test_case', __name__, template_folder='templates', static_folder='static', url_prefix='/test_case')
logger = setup_logger('test_case_module')

# 初始化存储（使用模块目录作为数据存储路径）
STORAGE_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(STORAGE_DIR, exist_ok=True)
# 平台级 LLM 配置单一真源（与 utils.llm_client.DEFAULT_LLM_CONFIG 保持一致，消除模块内双配置漂移）
LLM_CONFIG_FILE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'config', 'llm_config.json'
))
METERSPHERE_CONFIG_FILE = os.path.join(STORAGE_DIR, 'metersphere_config.json')

try:
    test_case_storage = TestCaseStorage(STORAGE_DIR)
except Exception as e:
    logger.error(f"初始化测试用例存储失败: {e}", exc_info=True)
    test_case_storage = None

DEFAULT_PRD_AUDIT_PROMPT_FILE = os.path.join(STORAGE_DIR, 'prd_audit_prompt_default.txt')
# 三段式 Stage3 极简审计报告生成（用 LLM 按八段格式输出，替代 Python 渲染）
STAGE3_MINIMAL_PROMPT_FILE = os.path.join(STORAGE_DIR, 'prd_audit_prompt_stage3_minimal.txt')

FALLBACK_PRD_PROMPT = (
    "你是 PRD 审计专家。请基于输入 PRD 输出结构化风险报告，至少包含："
    "总体结论、漏洞与风险清单、测试重点、研发重点、计划建议。"
    "仅输出报告正文，不要开场白。\n\nPRD 内容：\n{content}"
)

def _ensure_list(value):
    if isinstance(value, list):
        return [str(x).strip() for x in value if str(x).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []

def _numbered_lines(lines):
    arr = []
    if isinstance(lines, list):
        for x in lines:
            s = str(x or "").strip()
            if s:
                arr.append(s)
    elif isinstance(lines, str):
        for x in lines.splitlines():
            s = str(x or "").strip()
            if s:
                arr.append(s)
    if not arr:
        return ""
    return "\n".join([f"{i + 1}. {x}" for i, x in enumerate(arr)])

def _build_metersphere_row(case_data, idx):
    case_name = str(case_data.get("name") or f"未命名用例_{idx}").strip()
    module_name = str(case_data.get("category") or "未分类模块").strip()
    preconditions = str(case_data.get("preconditions") or "").strip()
    steps = case_data.get("steps") if isinstance(case_data.get("steps"), list) else []
    step_actions = [str(s.get("action") or "").strip() for s in steps if isinstance(s, dict)]
    step_expected = [str(s.get("expected") or "").strip() for s in steps if isinstance(s, dict)]
    step_actions = [x for x in step_actions if x]
    step_expected = [x for x in step_expected if x]
    row = {
        "用例编号": f"MS_TC_{datetime.now().strftime('%Y%m%d')}_{idx:03d}",
        "用例名称": case_name,
        "所属模块": module_name,
        "标签": "",
        "前置条件": preconditions,
        "备注": "",
        "步骤描述": _numbered_lines(step_actions),
        "预期结果": _numbered_lines(step_expected),
        "编辑模式": "STEP",
        "用例等级": "P1",
        "责任人(ID)": "fengnan",
        "用例状态": "Prepare",
    }
    return row

def _mask_secret(value):
    s = str(value or "").strip()
    if len(s) <= 7:
        return "****" if s else ""
    return s[:3] + "****" + s[-4:]

def _load_metersphere_config():
    if os.path.exists(METERSPHERE_CONFIG_FILE):
        with open(METERSPHERE_CONFIG_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def _save_metersphere_config(config):
    with open(METERSPHERE_CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump(config, f, ensure_ascii=False, indent=2)

@test_case_bp.route('/api/metersphere_config', methods=['GET', 'POST'])
def api_metersphere_config():
    if request.method == 'POST':
        try:
            data = request.get_json() or {}
            old = _load_metersphere_config()
            new_cfg = old.copy()
            # Added support for csrf_token and workspace_id
            for key in ["base_url", "import_api_path", "project_id", "auth_header_name", "payload_data_key", "workspace_id", "csrf_token"]:
                if key in data:
                    new_cfg[key] = data.get(key)
            auth_token = str(data.get("auth_token") or "").strip()
            cookie = str(data.get("cookie") or "").strip()
            if auth_token and "****" not in auth_token:
                new_cfg["auth_token"] = auth_token
            elif "auth_token" in old:
                new_cfg["auth_token"] = old.get("auth_token", "")
            if cookie and "****" not in cookie:
                new_cfg["cookie"] = cookie
            elif "cookie" in old:
                new_cfg["cookie"] = old.get("cookie", "")
            _save_metersphere_config(new_cfg)
            return success_response(message="MeterSphere 配置已保存")
        except Exception as e:
            logger.exception("保存 MeterSphere 配置失败")
            return error_response(str(e), status_code=500)
    else:
        try:
            cfg = _load_metersphere_config()
            if cfg.get("auth_token"):
                cfg["auth_token"] = _mask_secret(cfg.get("auth_token"))
            if cfg.get("cookie"):
                cfg["cookie"] = _mask_secret(cfg.get("cookie"))
            if cfg.get("csrf_token"):
                cfg["csrf_token"] = _mask_secret(cfg.get("csrf_token"))
            return success_response(data=cfg)
        except Exception as e:
            logger.exception("获取 MeterSphere 配置失败")
            return error_response(str(e), status_code=500)

@test_case_bp.route('/api/metersphere/import', methods=['POST'])
def api_import_to_metersphere():
    try:
        data = request.get_json() or {}
        # 优先使用前端解析好的 raw_rows (预览后的数据)
        rows = data.get("raw_rows")
        
        if not rows:
            # 如果没有 raw_rows，尝试从 cases 构建 (旧逻辑)
            cases = data.get("cases") or []
            if isinstance(cases, list) and cases:
                rows = [_build_metersphere_row(c if isinstance(c, dict) else {}, i + 1) for i, c in enumerate(cases)]
        
        if not rows:
            return error_response("没有可导入的用例数据", status_code=400)

        cfg = _load_metersphere_config()
        base_url = str(cfg.get("base_url") or "").rstrip("/")
        api_path = str(cfg.get("import_api_path") or "").strip()
        if not base_url or not api_path:
            return error_response("请先配置 MeterSphere 地址与导入API路径", status_code=400)
        if api_path.startswith("http://") or api_path.startswith("https://"):
            import_url = api_path
        else:
            if not api_path.startswith("/"):
                api_path = "/" + api_path
            import_url = base_url + api_path
        headers = {"Content-Type": "application/json"}
        auth_header_name = str(cfg.get("auth_header_name") or "Authorization").strip()
        token = str(cfg.get("auth_token") or "").strip()
        if token:
            if auth_header_name.lower() == "authorization":
                headers["Authorization"] = "Bearer " + token
            else:
                headers[auth_header_name] = token
        cookie = str(cfg.get("cookie") or "").strip()
        if cookie:
            headers["Cookie"] = cookie
        
        # Add new headers
        if cfg.get("project_id"):
            headers["Project"] = cfg.get("project_id")
        if cfg.get("workspace_id"):
            headers["Workspace"] = cfg.get("workspace_id")
        csrf_token = str(cfg.get("csrf_token") or "").strip()
        if csrf_token:
             headers["Csrf-Token"] = csrf_token

        data_key = str(cfg.get("payload_data_key") or "cases").strip()
        payload = {data_key: rows}
        if cfg.get("project_id"):
             # Some endpoints might expect projectId in body even if in header
             payload["projectId"] = cfg.get("project_id")
        
        # Log the request for debugging (mask sensitive data)
        safe_headers = headers.copy()
        if "Authorization" in safe_headers: safe_headers["Authorization"] = "***"
        if "Cookie" in safe_headers: safe_headers["Cookie"] = "***"
        if "Csrf-Token" in safe_headers: safe_headers["Csrf-Token"] = "***"
        logger.info(f"Sending to MeterSphere: {import_url}, Headers: {safe_headers}")

        resp = requests.post(import_url, json=payload, headers=headers, timeout=30)
        if 200 <= resp.status_code < 300:
            return success_response(data={
                "imported_count": len(rows),
                "target_url": import_url,
                "rows": rows,
                "metersphere_response": resp.text[:2000],
            }, message=f"已发送 {len(rows)} 条用例到 MeterSphere")
        return error_response(
            f"导入失败（HTTP {resp.status_code}）: {resp.text[:300]}",
            status_code=500
        )
    except Exception as e:
        logger.exception("导入 MeterSphere 失败")
        return error_response(str(e), status_code=500)

def _to_md_items(items):
    arr = _ensure_list(items)
    if not arr:
        return "- 【PRD未说明】"
    return "\n".join([f"- {x}" for x in arr])

def _calc_quality_score(defects):
    p0 = sum(1 for d in defects if str(d.get("risk_level", "")).upper() == "P0")
    p1 = sum(1 for d in defects if str(d.get("risk_level", "")).upper() == "P1")
    p2 = sum(1 for d in defects if str(d.get("risk_level", "")).upper() == "P2")
    score = 10.0 - p0 * 2.0 - p1 * 1.0 - p2 * 0.5
    if score < 0.0:
        score = 0.0
    if score >= 9.0:
        level = "高质量"
    elif score >= 7.0:
        level = "基本可开发"
    elif score >= 5.0:
        level = "存在明显风险"
    else:
        level = "不具备开发条件"
    return round(score, 1), level

def _risk_weight(level):
    v = str(level or "").upper()
    if v == "P0":
        return 3
    if v == "P1":
        return 2
    return 1

def _max_risk(levels):
    best = "P2"
    for lv in levels:
        if _risk_weight(lv) > _risk_weight(best):
            best = str(lv or "P2").upper()
    return best

def _core_group_name(defect):
    t = str(defect.get("type") or "")
    if any(k in t for k in ["状态", "跳转"]):
        return "状态机不完整"
    if any(k in t for k in ["流程", "并发", "重试"]):
        return "流程闭环缺失"
    if any(k in t for k in ["权限", "安全", "越权", "防护"]):
        return "权限与安全控制不足"
    if any(k in t for k in ["字段", "数据", "一致性"]):
        return "数据契约与一致性不足"
    if "逻辑矛盾" in t:
        return "规则冲突与口径不一致"
    return f"{str(defect.get('module') or '全局')}问题聚合"

def _merge_core_issues(defects):
    groups = {}
    for d in defects:
        name = _core_group_name(d)
        g = groups.setdefault(name, {
            "name": name,
            "risk_levels": [],
            "anchors": [],
            "modules": [],
            "types": [],
            "descriptions": [],
            "reasons": [],
            "suggestions": [],
            "count": 0,
        })
        g["count"] += 1
        g["risk_levels"].append(str(d.get("risk_level") or "P2").upper())
        anchor = str(d.get("anchor") or d.get("module") or "").strip()
        module = str(d.get("module") or "").strip()
        d_type = str(d.get("type") or "").strip()
        desc = str(d.get("description") or "").strip()
        reason = str(d.get("reason") or "").strip()
        sug = str(d.get("suggestion") or "").strip()
        if anchor and anchor not in g["anchors"]:
            g["anchors"].append(anchor)
        if module and module not in g["modules"]:
            g["modules"].append(module)
        if d_type and d_type not in g["types"]:
            g["types"].append(d_type)
        if desc and desc not in g["descriptions"]:
            g["descriptions"].append(desc)
        if reason and reason not in g["reasons"]:
            g["reasons"].append(reason)
        if sug and sug not in g["suggestions"]:
            g["suggestions"].append(sug)
    merged = []
    for g in groups.values():
        merged.append({
            "name": g["name"],
            "risk_level": _max_risk(g["risk_levels"]),
            "anchors": g["anchors"][:8],
            "modules": g["modules"][:6],
            "types": g["types"][:8],
            "description": "；".join(g["descriptions"][:3]) or "【PRD未说明】",
            "reason": "；".join(g["reasons"][:3]) or "【PRD未说明】",
            "suggestion": g["suggestions"][0] if g["suggestions"] else "补充该类问题的可执行规则与验收标准。",
            "count": g["count"],
        })
    merged.sort(key=lambda x: (_risk_weight(x["risk_level"]), x["count"]), reverse=True)
    return merged

def _build_core_risk_summary(defects, merged_issues):
    has_conflict = any("逻辑矛盾" in str(d.get("type") or "") for d in defects)
    has_state = any("状态" in str(d.get("type") or "") for d in defects)
    has_concurrency = any("并发" in str(d.get("type") or "") for d in defects)
    if has_conflict and has_state and has_concurrency:
        one_liner = "这是一个定义了“谁优先级高”但没有定义“怎么切换”的系统，在真实并发场景下必然混乱。"
    elif has_state and has_concurrency:
        one_liner = "系统具备功能描述，但关键状态切换与并发处理规则不足，上线后易出现行为不一致。"
    else:
        one_liner = "当前 PRD 存在多处关键规则缺口，建议先完成核心风险闭环再进入开发。"
    top3 = merged_issues[:3]
    bullets = []
    for item in top3:
        bullets.append(f"{item.get('name')}（{item.get('risk_level')}，{item.get('count')}项）")
    return {"one_liner": one_liner, "top3": bullets}

def _run_stage3_llm_report(prd_content, stage1_output, stage2_output, llm_config_path, timeout=90):
    """
    若存在极简审计 Stage3 prompt 文件，则用 LLM 按八段格式生成报告（结合你的六层思维/七维评分/合并原则）。
    返回 Markdown 字符串；失败或文件不存在时返回 None，调用方回退到 Python Stage3+Stage4。
    """
    prompt_file = STAGE3_MINIMAL_PROMPT_FILE
    if not os.path.exists(prompt_file):
        return None
    try:
        with open(prompt_file, 'r', encoding='utf-8') as f:
            template = f.read().strip()
        if not template or '{structure_json}' not in template or '{defects_json}' not in template:
            return None
        import json
        structure_json = json.dumps(stage1_output or {}, ensure_ascii=False, indent=2)
        defects = (stage2_output or {}).get('defects') if isinstance(stage2_output, dict) else []
        defects_json = json.dumps(defects if isinstance(defects, list) else [], ensure_ascii=False, indent=2)
        prd_snippet = (prd_content or '')[:12000].strip() or '【无PRD原文】'
        prompt_text = template.replace('{prd_content}', prd_snippet)
        prompt_text = prompt_text.replace('{structure_json}', structure_json)
        prompt_text = prompt_text.replace('{defects_json}', defects_json)
        from utils.llm_client import call_llm
        report = call_llm(
            [{"role": "user", "content": prompt_text}],
            config_path=llm_config_path,
            stream=False,
            timeout=timeout,
            max_tokens=16384,
        )
        return (report or '').strip()
    except Exception as e:
        logger.warning("Stage3 LLM report failed, fallback to Python: %s", e)
        return None

def _build_stage3_report(stage1_output, stage2_output):
    defects = stage2_output.get("defects") if isinstance(stage2_output, dict) else []
    defects = defects if isinstance(defects, list) else []
    score, risk_level = _calc_quality_score(defects)
    main_problem = defects[0].get("description") if defects else "未发现明显漏洞"
    risks = []
    for d in defects:
        risk = f"{d.get('module', '【PRD未说明】')}：{d.get('description', '【PRD未说明】')}"
        if risk not in risks:
            risks.append(risk)
    test_focus = []
    dev_focus = []
    for d in defects:
        t = str(d.get("type") or "")
        if t and t not in test_focus:
            test_focus.append(t)
        m = str(d.get("module") or "")
        if m and m not in dev_focus:
            dev_focus.append(m)
    if not test_focus:
        test_focus = ["核心流程回归", "边界与异常场景"]
    if not dev_focus:
        dev_focus = _ensure_list(stage1_output.get("modules")) or ["关键功能模块"]
    plan = []
    if any((d.get("risk_level") or "").upper() == "P0" for d in defects):
        plan.append("先完成 P0 漏洞澄清并冻结关键流程")
    plan.append("按 P0→P1→P2 优先级推进修复与复审")
    plan.append("同步更新测试点并执行回归验证")
    source_stats = {"rule": 0, "llm": 0, "hybrid": 0}
    for d in defects:
        src = str(d.get("source") or "llm").strip().lower()
        if src in source_stats:
            source_stats[src] += 1
    merged_issues = _merge_core_issues(defects)
    core_summary = _build_core_risk_summary(defects, merged_issues)
    p0_count = sum(1 for d in defects if str(d.get("risk_level", "")).upper() == "P0")
    report_title = f"【审计报告】PRD：工具扫描+人工复核版（含{len(defects)}项缺陷，P0级{p0_count}项）"
    return {
        "summary": {
            "quality_score": score,
            "risk_level": risk_level,
            "main_problem": main_problem,
        },
        "report_title": report_title,
        "core_risk_summary": core_summary,
        "merged_issues": merged_issues,
        "defects": defects,
        "scan_stats": source_stats,
        "risks": risks[:20],
        "test_focus": test_focus[:20],
        "dev_focus": dev_focus[:20],
        "plan": plan,
    }

def _render_stage4_markdown(stage3_json):
    summary = (stage3_json or {}).get("summary") or {}
    report_title = (stage3_json or {}).get("report_title") or "【审计报告】PRD：工具扫描+人工复核版"
    core_summary = (stage3_json or {}).get("core_risk_summary") or {}
    merged_issues = (stage3_json or {}).get("merged_issues") or []
    defects_data = (stage3_json or {}).get("defects") or []
    if not isinstance(defects_data, list):
        defects_data = []
    score = summary.get('quality_score', 0)
    try:
        score = float(score) if score is not None else 0
    except (TypeError, ValueError):
        score = 0
    lines = [f"# {report_title}", "", "## 一、总体结论", ""]
    lines.append(f"- 审计结论：{summary.get('main_problem', '【PRD未说明】')}")
    lines.append(f"- 综合质量评分：{score}/10（基于七维评分）")
    lines.append("### 核心风险摘要")
    lines.append(f"- 一句话总结：{core_summary.get('one_liner', '【PRD未说明】')}")
    top3 = _ensure_list(core_summary.get("top3"))
    if top3:
        lines.append("- 三个致命伤：")
        for x in top3:
            lines.append(f"  - {x}")
    lines.append("")
    lines.append("### 七维质量评分明细（必填）")
    lines.append("")
    lines.append("| 维度 | 评分 | 说明 |")
    lines.append("| :--- | :--- | :--- |")
    for dim in ["需求完整度", "规则明确度", "流程一致性", "状态机完备度", "异常覆盖度", "可测试性", "技术可实现性"]:
        lines.append(f"| {dim} | {score}/10 | 基于缺陷综合评估 |")
    lines.append("")
    lines.extend(["", "## 二、核心问题矩阵（合并版）", ""])
    if merged_issues:
        lines.append("| 风险等级 | 核心问题 | 涉及锚点 | 问题描述 | 风险分析 | 审计建议 |")
        lines.append("| :--- | :--- | :--- | :--- | :--- | :--- |")
        for m in merged_issues:
            anchors = "; ".join(_ensure_list(m.get("anchors"))) or "【PRD未说明】"
            lines.append(f"| {m.get('risk_level', 'P2')} | {m.get('name', '【PRD未说明】')} | {anchors} | {m.get('description', '【PRD未说明】')} | {m.get('reason', '【PRD未说明】')} | {m.get('suggestion', '【PRD未说明】')} |")
        lines.append("")
    lines.extend(["", "### 详细漏洞清单（原始版）", ""])
    stats = (stage3_json or {}).get("scan_stats") or {}
    lines.append(f"- 扫描来源：规则库 {stats.get('rule', 0)} 条，LLM {stats.get('llm', 0)} 条，混合 {stats.get('hybrid', 0)} 条")
    lines.append("")
    if defects_data:
        for i, d in enumerate(defects_data, start=1):
            lines.append(f"### 漏洞{i}")
            lines.append(f"- 模块：{d.get('module', '【PRD未说明】')}")
            lines.append(f"- 类型：{d.get('type', '【PRD未说明】')}")
            lines.append(f"- 风险等级：{d.get('risk_level', '【PRD未说明】')}")
            lines.append(f"- 来源：{d.get('source', 'llm')}")
            lines.append(f"- 锚点：{d.get('anchor', d.get('module', '【PRD未说明】'))}")
            lines.append(f"- 描述：{d.get('description', '【PRD未说明】')}")
            lines.append(f"- 原因：{d.get('reason', '【PRD未说明】')}")
            lines.append(f"- 建议：{d.get('suggestion', '【PRD未说明】')}")
            lines.append("")
    else:
        lines.append("- 未发现漏洞")
        lines.append("")
    lines.extend(["", "## 四、待确认清单", ""])
    lines.append("| 优先级 | 待确认项 | 涉及模块 | 具体问题 | 影响 |")
    lines.append("| :--- | :--- | :--- | :--- | :--- |")
    pending = [d for d in defects_data if "待确认" in str(d.get("suggestion") or "") or "【PRD未说明】" in str(d.get("description") or "")]
    for d in pending[:15]:
        lines.append(f"| P1 | {d.get('type', '')} | {d.get('module', '')} | {d.get('description', '')[:80]} | 需澄清后开发 |")
    if not pending:
        lines.append("| - | 无 | - | - | - |")
    lines.append("")
    lines.extend(["## 五、测试重点", _to_md_items((stage3_json or {}).get("test_focus")), ""])
    lines.extend(["## 六、研发重点", _to_md_items((stage3_json or {}).get("dev_focus")), ""])
    lines.extend(["## 七、项目风险", _to_md_items((stage3_json or {}).get("risks")), ""])
    lines.extend(["## 八、计划建议", _to_md_items((stage3_json or {}).get("plan")), ""])
    return "\n".join(lines).strip()

@test_case_bp.route('/')
def index():
    """主页面"""
    default_prd_prompt = FALLBACK_PRD_PROMPT
    try:
        if os.path.exists(DEFAULT_PRD_AUDIT_PROMPT_FILE):
            with open(DEFAULT_PRD_AUDIT_PROMPT_FILE, 'r', encoding='utf-8') as f:
                s = f.read().strip()
                if s and len(s) > 100:
                    default_prd_prompt = s
                    logger.debug("Loaded PRD audit prompt from file, len=%s", len(s))
                else:
                    logger.warning("PRD audit prompt file empty or too short, using fallback")
        else:
            logger.warning("PRD audit prompt file not found: %s", DEFAULT_PRD_AUDIT_PROMPT_FILE)
    except Exception as e:
        logger.warning("Failed to read PRD audit prompt file: %s", e)
    if not isinstance(default_prd_prompt, str):
        default_prd_prompt = FALLBACK_PRD_PROMPT
    llm_profiles = []
    try:
        if os.path.exists(LLM_CONFIG_FILE):
            with open(LLM_CONFIG_FILE, 'r', encoding='utf-8') as f:
                cfg = json.load(f)
            profiles = cfg.get('profiles') if isinstance(cfg.get('profiles'), dict) else {}
            default_profile = (cfg.get('default_profile') or cfg.get('llm_provider') or 'deepseek').strip()
            for k, v in profiles.items():
                if not isinstance(v, dict):
                    continue
                llm_profiles.append({
                    'key': k,
                    'llm_provider': (v.get('llm_provider') or k),
                    'model': (v.get('model') or ''),
                })
            if not llm_profiles:
                llm_profiles.append({
                    'key': default_profile,
                    'llm_provider': cfg.get('llm_provider') or default_profile,
                    'model': cfg.get('model') or '',
                })
            llm_profiles.sort(key=lambda x: (0 if x.get('key') == default_profile else 1, str(x.get('key') or '')))
    except Exception:
        pass
    return render_template(
        'test_case_index.html',
        default_prd_audit_prompt=default_prd_prompt,
        fallback_prd_audit_prompt=FALLBACK_PRD_PROMPT,
        llm_profiles=llm_profiles,
    )

@test_case_bp.route('/api/default_prd_prompt', methods=['GET'])
def api_get_default_prd_prompt():
    """返回 PRD 默认审计提示词（从 prd_audit_prompt_default.txt 读取），供前端展示"""
    try:
        if os.path.exists(DEFAULT_PRD_AUDIT_PROMPT_FILE):
            with open(DEFAULT_PRD_AUDIT_PROMPT_FILE, 'r', encoding='utf-8') as f:
                s = f.read().strip()
                if s and len(s) > 100:
                    return success_response(data={'prompt': s})
        return success_response(data={'prompt': FALLBACK_PRD_PROMPT})
    except Exception as e:
        logger.warning("Failed to read default PRD prompt: %s", e)
        return success_response(data={'prompt': FALLBACK_PRD_PROMPT})

@test_case_bp.route('/list')
def list_page():
    """测试用例列表页面"""
    return render_template('test_case_list.html')

@test_case_bp.route('/prompts')
def prompt_config_page():
    """提示词配置页面"""
    return render_template('prompt_config.html')

@test_case_bp.route('/knowledge')
def knowledge_center_page():
    """测试知识库中心页面"""
    return render_template('knowledge_center.html')

KNOWLEDGE_RULES_FILE = os.path.join(STORAGE_DIR, 'knowledge_rules.json')

@test_case_bp.route('/api/knowledge', methods=['GET'])
def api_get_knowledge():
    """获取测试知识库（功能图谱 + Bug 模式）"""
    try:
        if os.path.exists(KNOWLEDGE_RULES_FILE):
            with open(KNOWLEDGE_RULES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = {'features': [], 'bug_patterns': []}
        return success_response(data=data)
    except Exception as e:
        logger.exception("获取知识库失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/knowledge', methods=['POST'])
def api_save_knowledge():
    """保存测试知识库"""
    try:
        data = request.get_json() or {}
        features = data.get('features', [])
        bug_patterns = data.get('bug_patterns', [])
        to_write = {'features': features, 'bug_patterns': bug_patterns}
        with open(KNOWLEDGE_RULES_FILE, 'w', encoding='utf-8') as f:
            json.dump(to_write, f, ensure_ascii=False, indent=2)
        return success_response(data=to_write)
    except Exception as e:
        logger.exception("保存知识库失败")
        return error_response(str(e), status_code=500)

# ========== PRD 规则库 ==========

PRD_SCAN_RULES_FILE = os.path.join(STORAGE_DIR, 'prd_scan_rules.json')

@test_case_bp.route('/prd_rules')
def prd_rules_page():
    """PRD 漏洞扫描规则库页面"""
    return render_template('prd_rules.html')

@test_case_bp.route('/api/prd_scan_rules', methods=['GET'])
def api_get_prd_scan_rules():
    """获取 PRD 漏洞扫描规则库"""
    try:
        if os.path.exists(PRD_SCAN_RULES_FILE):
            with open(PRD_SCAN_RULES_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
        else:
            data = {'version': '1.0', 'description': '', 'categories': [], 'rules': []}
        return success_response(data=data)
    except Exception as e:
        logger.exception("获取 PRD 规则库失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/prd_scan_rules', methods=['POST'])
def api_save_prd_scan_rules():
    """保存 PRD 漏洞扫描规则库"""
    try:
        data = request.get_json() or {}
        with open(PRD_SCAN_RULES_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return success_response(data=data)
    except Exception as e:
        logger.exception("保存 PRD 规则库失败")
        return error_response(str(e), status_code=500)

# ========== 提示词 API ==========

@test_case_bp.route('/api/prompts', methods=['GET'])
def api_get_prompts():
    """获取提示词列表"""
    try:
        if not test_case_storage:
            return error_response('存储未初始化', status_code=500)
        
        prompts = list(test_case_storage.prompts.values())
        # 按创建时间倒序
        prompts.sort(key=lambda x: x.created_at or datetime.min, reverse=True)
        
        return success_response(data=[p.to_dict() for p in prompts])
    except Exception as e:
        logger.exception("获取提示词列表失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/prompts', methods=['POST'])
def api_create_prompt():
    """创建提示词"""
    try:
        if not test_case_storage:
            return error_response('存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        if not data.get('name') or not data.get('content'):
            return error_response('名称和内容不能为空', status_code=400)
            
        prompt = PromptConfig(
            id=uuid.uuid4().hex,
            name=data['name'],
            type=data.get('type', 'case_writing'),
            content=data['content'],
            tags=data.get('tags', []),
            is_active=data.get('is_active', True),
            created_by=data.get('created_by', 'admin')
        )
        
        test_case_storage.prompts[prompt.id] = prompt
        test_case_storage.save_prompts()
        
        return success_response(data=prompt.to_dict())
    except Exception as e:
        logger.exception("创建提示词失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/prompts/<prompt_id>', methods=['PUT'])
def api_update_prompt(prompt_id):
    """更新提示词"""
    try:
        if not test_case_storage:
            return error_response('存储未初始化', status_code=500)
        
        if prompt_id not in test_case_storage.prompts:
            return error_response('提示词不存在', status_code=404)
            
        data = request.get_json() or {}
        prompt = test_case_storage.prompts[prompt_id]
        
        if 'name' in data:
            prompt.name = data['name']
        if 'type' in data:
            prompt.type = data['type']
        if 'content' in data:
            prompt.content = data['content']
        if 'tags' in data:
            prompt.tags = data['tags']
        if 'is_active' in data:
            prompt.is_active = data['is_active']
            
        prompt.updated_at = datetime.now()
        test_case_storage.save_prompts()
        
        return success_response(data=prompt.to_dict())
    except Exception as e:
        logger.exception("更新提示词失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/prompts/<prompt_id>', methods=['DELETE'])
def api_delete_prompt(prompt_id):
    """删除提示词"""
    try:
        if not test_case_storage:
            return error_response('存储未初始化', status_code=500)
        
        if prompt_id in test_case_storage.prompts:
            del test_case_storage.prompts[prompt_id]
            test_case_storage.save_prompts()
            return success_response(message='删除成功')
        else:
            return error_response('提示词不存在', status_code=404)
    except Exception as e:
        logger.exception("删除提示词失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/prompts/load_defaults', methods=['POST'])
def api_load_default_prompts():
    """加载默认提示词"""
    try:
        if not test_case_storage:
            return error_response('存储未初始化', status_code=500)
        
        defaults = [
            {
                "name": "测试大纲生成提示词",
                "type": "outline_writing",
                "content": "作为一个资深测试工程师，请根据以下需求文档内容，编写一份完整的测试大纲。要求结合 BDD 与 TDD 思路：\n\n- **BDD**：关键测试场景用 Given-When-Then 描述（Given 前置/上下文，When 操作/触发，Then 预期/验收结果）。\n- **TDD**：每个测试点或模块先写出「验收标准」或「通过条件」，再列出测试范围与策略，体现「先定义怎样算通过，再设计怎么测」。\n\n大纲需包含：测试范围、测试策略、功能模块划分、关键测试点（尽量用 Given-When-Then 或验收标准先行的方式呈现）。\n\n需求内容：\n{content}",
                "tags": ["大纲生成", "默认"]
            },
            {
                "name": "用例编写提示词",
                "type": "case_writing",
                "content": "作为一个资深测试工程师 (BDD + TDD)，请根据以下需求描述编写测试用例。\n\n- **BDD**：每条用例严格按 Given-When-Then 结构：前置条件 (Given) → 操作步骤 (When) → 预期结果 (Then)；步骤与预期结果一一对应。\n- **TDD**：设计时先写「预期结果/验收标准」，再写操作步骤，体现验收先行、可验证。\n\n需求描述：\n{content}",
                "tags": ["BDD", "默认"]
            },
            {
                "name": "PRD 漏洞分析（八方面报告）",
                "type": "prd_review",
                "content": "你是资深需求评审与质量分析专家，只做 PRD 漏洞与风险分析，不编写测试用例或测试大纲。\n\n请根据用户提供的 PRD 文档内容，输出一份结构化的评审报告。必须严格按以下 8 个方面、使用如下二级标题（Markdown），不得遗漏或合并：\n\n## 一、总体结论\n- 可评审性/可测试性结论（一句话）\n- 综合质量评分（X/10）及简要理由\n\n## 二、漏洞与风险清单\n按维度分条列出，每条包含：问题描述、涉及位置/模块、建议（或标为【待确认】）。\n- 逻辑矛盾\n- 功能/场景缺失\n- 描述模糊与歧义\n- 可测试性不足\n- 体验与流程风险\n- 技术与实现风险\n\n## 三、待确认清单\n汇总所有需产品/业务确认的问题，便于评审会逐条过。\n\n## 四、测试重点\n建议的测试重点：核心场景、高风险模块、必测路径、边界与异常、性能/安全等专项；可按模块或优先级列。\n\n## 五、研发评估\n从研发视角：实现难度/工作量粗估、技术风险、依赖与阻塞、建议前期澄清或预研的点。\n\n## 六、计划建议\n建议的节奏：需求澄清 → 研发评估 → 测试设计 → 用例编写等时间顺序或里程碑建议。\n\n## 七、优先级与责任建议\n对上述问题做 P0/P1/P2 分级；责任归属建议（产品/开发/测试）。\n\n## 八、后续动作建议\n接下来几步该做什么（澄清哪些、谁先做、再做什么）。\n\n约束：不臆测需求细节；拿不准的标为【待确认】。直接输出报告正文，不要开场白。\n\nPRD 文档内容如下：\n\n{content}",
                "tags": ["PRD评审", "默认"]
            }
        ]
        
        added_count = 0
        for d in defaults:
            # 简单的查重：按名称检查
            exists = any(p.name == d['name'] for p in test_case_storage.prompts.values())
            if not exists:
                prompt = PromptConfig(
                    id=uuid.uuid4().hex,
                    name=d['name'],
                    type=d['type'],
                    content=d['content'],
                    tags=d['tags']
                )
                test_case_storage.prompts[prompt.id] = prompt
                added_count += 1
        
        if added_count > 0:
            test_case_storage.save_prompts()
            
        return success_response(message=f'成功加载 {added_count} 个默认提示词')
    except Exception as e:
        logger.exception("加载默认提示词失败")
        return error_response(str(e), status_code=500)

# ========== LLM 配置 API ==========

@test_case_bp.route('/api/generate', methods=['POST'])
def api_generate():
    """执行 AI 生成"""
    try:
        data = request.get_json() or {}
        gen_type = data.get('type')  # outline_writing, case_writing
        content = data.get('content')
        custom_prompt = data.get('prompt')  # 支持前端传入自定义提示词
        llm_profile = (data.get('llm_profile') or '').strip()
        
        missing = []
        if not gen_type:
            missing.append('type（生成类型）')
        if not content or (isinstance(content, str) and not content.strip()):
            missing.append('content（需求内容）')
        if missing:
            return error_response('缺少必要参数: ' + '、'.join(missing), status_code=400)

        content = content.strip() if isinstance(content, str) else content

        # PRD 漏洞分析：若输入为飞书文档链接，先拉取正文
        if gen_type == 'prd_review' and isinstance(content, str):
            try:
                from .feishu_client import is_feishu_doc_url, fetch_feishu_doc_content
                if is_feishu_doc_url(content):
                    ok, result = fetch_feishu_doc_content(content)
                    if not ok:
                        return error_response(
                            result if result else '无法拉取飞书文档，请复制内容粘贴或导出后上传',
                            status_code=400
                        )
                    content = result
                    logger.info("Feishu doc content fetched, length=%s", len(content))
            except Exception as e:
                logger.exception("Feishu fetch failed")
                return error_response('拉取飞书文档失败: ' + str(e), status_code=500)
            
        # 1. 获取 LLM 配置（支持 Gemini / DeepSeek / OpenAI）
        try:
            from utils.llm_client import load_llm_config
            llm_config = load_llm_config(LLM_CONFIG_FILE)
        except FileNotFoundError as e:
            return error_response(f'请先配置 LLM (用例管理 → LLM 全局配置)', status_code=400)
            
        profiles = llm_config.get('profiles') if isinstance(llm_config.get('profiles'), dict) else {}
        default_profile = (llm_config.get('default_profile') or llm_config.get('llm_provider') or 'deepseek').strip()
        chosen_profile = llm_profile or default_profile
        active_config = profiles.get(chosen_profile) if isinstance(profiles.get(chosen_profile), dict) else None
        if not active_config:
            active_config = {
                'llm_provider': llm_config.get('llm_provider') or default_profile,
                'base_url': llm_config.get('base_url'),
                'api_key': llm_config.get('api_key'),
                'model': llm_config.get('model'),
            }
        for k in ['fallback_enabled', 'fallback_provider', 'fallback_api_key', 'fallback_base_url', 'fallback_model']:
            if k in llm_config and active_config.get(k) is None:
                active_config[k] = llm_config.get(k)

        api_key = (active_config.get('api_key') or '').strip()
        has_env = bool((os.environ.get('LLM_API_KEY') or '').strip()) or \
                  bool((os.environ.get('DEEPSEEK_API_KEY') or '').strip()) or \
                  bool((os.environ.get('OPENAI_API_KEY') or '').strip()) or \
                  bool((os.environ.get('GEMINI_API_KEY') or '').strip()) or \
                  bool((os.environ.get('VOLCENGINE_API_KEY') or '').strip()) or \
                  bool((os.environ.get('ARK_API_KEY') or '').strip())
        if not api_key and not has_env:
            return error_response('API Key 未配置', status_code=400)

        # PRD 漏洞分析：委托 prd_audit 流水线（单源实现）
        if gen_type == 'prd_review':
            from flask import stream_with_context, Response
            from modules.prd_audit import pipeline as prd_pipeline

            def generate_prd():
                try:
                    for chunk in prd_pipeline.run_prd_audit_stream(
                        content,
                        llm_config_path=LLM_CONFIG_FILE,
                        llm_config_override=active_config,
                        timeout=180,
                        custom_prompt=custom_prompt.strip() if custom_prompt and custom_prompt.strip() else None,
                    ):
                        yield chunk
                except Exception as e:
                    logger.exception("PRD generate failed")
                    yield json.dumps({"type": "error", "text": str(e)}, ensure_ascii=False) + '\n'

            return Response(stream_with_context(generate_prd()), content_type='application/x-ndjson; charset=utf-8')

        # 2. 获取 Prompt (优先使用前端传入的，否则查库)
        prompt_template = ""
        if custom_prompt:
            prompt_template = custom_prompt
        else:
            # 获取对应的 Prompt
            if not test_case_storage:
                 return error_response('存储未初始化', status_code=500)
            
            # 查找启用的、类型匹配的 Prompt (优先取最近更新的)
            prompts = list(test_case_storage.prompts.values())
            prompts = [p for p in prompts if p.type == gen_type and p.is_active]
            prompts.sort(key=lambda x: x.updated_at or datetime.min, reverse=True)
            
            if not prompts:
                 return error_response(f'未找到类型为 {gen_type} 的启用提示词', status_code=404)
                 
            prompt_template = prompts[0].content
        
        # 3. 组装 Messages
        messages = []
        if '{content}' in prompt_template:
            # 如果提示词包含占位符，说明它是完整的指令模板
            final_content = prompt_template.replace('{content}', content)
            messages.append({"role": "user", "content": final_content})
        else:
            # 如果提示词不包含占位符，将其作为 System Prompt (角色设定/背景信息)
            # 而将用户的需求作为 User Message
            logger.info("Prompt template does not contain '{content}'. Using it as System Prompt.")
            messages.append({"role": "system", "content": prompt_template})
            messages.append({"role": "user", "content": content})
            
            # 为了调试，记录一下
            final_prompt = f"[System]: {prompt_template[:50]}...\n[User]: {content[:50]}..."
            logger.debug(f"Constructed messages: {final_prompt}")

        # 4. 调用 LLM（支持 Gemini / DeepSeek / OpenAI，优先 Gemini）
        from flask import stream_with_context, Response
        from utils.llm_client import stream_llm

        provider = (active_config.get('llm_provider') or 'openai').lower()
        model = active_config.get('model', '')
        provider_names = {'gemini': 'Google Gemini', 'deepseek': 'DeepSeek', 'openai': 'OpenAI', 'custom': 'Custom API', 'volcengine': '火山引擎（方舟）'}
        provider_display = provider_names.get(provider, provider)
        status_msg = f'正在连接 {provider_display}'
        if model:
            status_msg += f' ({model})'
        status_msg += '...'

        # 构建 fallback 配置（Gemini 失败时自动切换 DeepSeek）
        fallback_config = None
        fallback_api_key = (llm_config.get('fallback_api_key') or '').strip()
        if fallback_api_key and llm_config.get('fallback_enabled'):
            fb_provider = (llm_config.get('fallback_provider') or 'deepseek').lower()
            fallback_config = {
                'llm_provider': fb_provider,
                'api_key': fallback_api_key,
                'base_url': (llm_config.get('fallback_base_url') or 'https://api.deepseek.com/v1').rstrip('/'),
                'model': llm_config.get('fallback_model') or 'deepseek-chat'
            }

        logger.info(f"Calling LLM: provider={provider}, model={model}, fallback={bool(fallback_config)}")

        def generate():
            try:
                yield json.dumps({"type": "status", "text": status_msg}, ensure_ascii=False) + '\n'
                yield json.dumps({"type": "status", "text": "已连接，等待模型响应...\n"}, ensure_ascii=False) + '\n'
                try:
                    for chunk in stream_llm(messages, config_override=active_config, timeout=60):
                        if chunk and isinstance(chunk, dict):
                            yield json.dumps(chunk, ensure_ascii=False) + '\n'
                        elif chunk and isinstance(chunk, str):
                            yield json.dumps({"type": "content", "text": chunk}, ensure_ascii=False) + '\n'
                except Exception as e:
                    if fallback_config:
                        logger.warning("Primary LLM failed, switching to fallback: %s", e)
                        yield json.dumps({"type": "status", "text": f"连接失败 ({str(e)[:50]}...)，正在切换到 DeepSeek...\n"}, ensure_ascii=False) + '\n'
                        for chunk in stream_llm(messages, config_override=fallback_config, timeout=60):
                            if chunk and isinstance(chunk, dict):
                                yield json.dumps(chunk, ensure_ascii=False) + '\n'
                            elif chunk and isinstance(chunk, str):
                                yield json.dumps({"type": "content", "text": chunk}, ensure_ascii=False) + '\n'
                    else:
                        raise
            except Exception as e:
                yield json.dumps({"type": "error", "text": str(e)}, ensure_ascii=False) + '\n'

        return Response(stream_with_context(generate()), content_type='application/x-ndjson; charset=utf-8')

    except Exception as e:
        logger.exception("AI 生成失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/parse_pdf', methods=['POST'])
def api_parse_pdf():
    """解析 PDF 文件"""
    try:
        if 'file' not in request.files:
            return error_response('未上传文件', status_code=400)
        
        file = request.files['file']
        if not file.filename:
            return error_response('文件名为空', status_code=400)
            
        if not file.filename.lower().endswith('.pdf'):
            return error_response('请上传 PDF 文件', status_code=400)

        from io import BytesIO
        file_data = file.read()
        file_stream = BytesIO(file_data)

        text_chunks = []
        page_errors = 0
        try:
            import pypdf
            reader = pypdf.PdfReader(file_stream, strict=False)
            for page in reader.pages:
                try:
                    text_chunks.append(page.extract_text() or "")
                except Exception:
                    page_errors += 1
        except ImportError:
            try:
                import PyPDF2
                file_stream.seek(0)
                reader = PyPDF2.PdfReader(file_stream)
                for page in reader.pages:
                    try:
                        text_chunks.append(page.extract_text() or "")
                    except Exception:
                        page_errors += 1
            except ImportError:
                return error_response('未安装 PDF 解析库 (pypdf 或 PyPDF2)。请联系管理员安装: pip install pypdf', status_code=500)
        except Exception as e:
            logger.error(f"PDF解析错误: {e}")
            return error_response(f'PDF 解析失败: {str(e)}', status_code=500)

        text = "\n".join([c for c in text_chunks if isinstance(c, str) and c.strip()]).strip()
        if text:
            data = {'text': text}
            if page_errors > 0:
                data['warning'] = f'部分页面解析失败（{page_errors}页），已返回可提取内容。'
            return success_response(data=data)

        fallback_hint = (
            "【PDF解析提示】该PDF可能是扫描件/图片版，未提取到可复制文字。\n"
            "建议：1）导出可复制文本的PDF；2）上传Word(.docx)；3）直接粘贴PRD正文。"
        )
        return success_response(data={'text': fallback_hint, 'warning': '未提取到可复制文本，已返回导入提示。'})

    except Exception as e:
        logger.exception("PDF 上传解析失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/parse_docx', methods=['POST'])
def api_parse_docx():
    """解析 Word (.docx) 文件，提取正文供 PRD 分析"""
    try:
        if 'file' not in request.files:
            return error_response('未上传文件', status_code=400)
        file = request.files['file']
        if not file.filename:
            return error_response('文件名为空', status_code=400)
        if not file.filename.lower().endswith('.docx'):
            return error_response('请上传 .docx 文件', status_code=400)
        try:
            from docx import Document
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        except ImportError:
            return error_response('未安装 python-docx。请联系管理员: pip install python-docx', status_code=500)
        except Exception as e:
            logger.error("DOCX 解析错误: %s", e)
            return error_response('Word 解析失败: ' + str(e), status_code=500)
        if not text.strip():
            return error_response('未从 Word 中提取到文本', status_code=400)
        return success_response(data={'text': text})
    except Exception as e:
        logger.exception("Word 上传解析失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/export_report_docx', methods=['POST'])
def api_export_report_docx():
    """将评审报告 Markdown 转为 Word (.docx) 并返回文件"""
    try:
        data = request.get_json() or {}
        content = (data.get('content') or '').strip()
        stage3_json = data.get('stage3_json') if isinstance(data.get('stage3_json'), dict) else None
        if not content and not stage3_json:
            return error_response('报告内容为空', status_code=400)
        try:
            from docx import Document
            doc = Document()
            if stage3_json and isinstance(stage3_json, dict):
                title = stage3_json.get("report_title") or '【审计报告】PRD：工具扫描+人工复核版'
            else:
                title = 'PRD 评审报告'
            doc.add_heading(title, 0)
            if stage3_json:
                summary = stage3_json.get("summary") or {}
                defects = stage3_json.get("defects") or []
                core_summary = stage3_json.get("core_risk_summary") or {}
                merged_issues = stage3_json.get("merged_issues") or []
                doc.add_heading('一、总体结论', level=1)
                doc.add_paragraph(f"质量评分：{summary.get('quality_score', '【PRD未说明】')}/10")
                doc.add_paragraph(f"风险等级：{summary.get('risk_level', '【PRD未说明】')}")
                doc.add_paragraph(f"主要问题：{summary.get('main_problem', '【PRD未说明】')}")
                doc.add_heading('核心风险摘要', level=2)
                doc.add_paragraph(f"一句话总结：{core_summary.get('one_liner', '【PRD未说明】')}")
                for item in _ensure_list(core_summary.get("top3")):
                    doc.add_paragraph(f"- {item}")
                doc.add_heading('二、漏洞与风险清单', level=1)
                stats = stage3_json.get("scan_stats") or {}
                doc.add_paragraph(f"扫描来源统计：规则库 {stats.get('rule', 0)} 条，LLM {stats.get('llm', 0)} 条，混合 {stats.get('hybrid', 0)} 条")
                if isinstance(merged_issues, list) and merged_issues:
                    doc.add_heading('合并后的核心问题', level=2)
                    for i, m in enumerate(merged_issues, start=1):
                        doc.add_heading(f"核心问题{i}：{m.get('name', '【PRD未说明】')}（{m.get('risk_level', 'P2')}）", level=3)
                        doc.add_paragraph(f"涉及锚点：{'; '.join(_ensure_list(m.get('anchors'))) or '【PRD未说明】'}")
                        doc.add_paragraph(f"问题描述：{m.get('description', '【PRD未说明】')}")
                        doc.add_paragraph(f"风险分析：{m.get('reason', '【PRD未说明】')}")
                        doc.add_paragraph(f"建议：{m.get('suggestion', '【PRD未说明】')}")
                if isinstance(defects, list) and defects:
                    doc.add_heading('详细漏洞清单', level=2)
                    for i, d in enumerate(defects, start=1):
                        doc.add_heading(f"漏洞{i}", level=2)
                        doc.add_paragraph(f"模块：{d.get('module', '【PRD未说明】')}")
                        doc.add_paragraph(f"类型：{d.get('type', '【PRD未说明】')}")
                        doc.add_paragraph(f"风险等级：{d.get('risk_level', '【PRD未说明】')}")
                        doc.add_paragraph(f"来源：{d.get('source', 'llm')}")
                        doc.add_paragraph(f"锚点：{d.get('anchor', d.get('module', '【PRD未说明】'))}")
                        doc.add_paragraph(f"描述：{d.get('description', '【PRD未说明】')}")
                        doc.add_paragraph(f"原因：{d.get('reason', '【PRD未说明】')}")
                        doc.add_paragraph(f"建议：{d.get('suggestion', '【PRD未说明】')}")
                else:
                    doc.add_paragraph("未发现漏洞")
                doc.add_heading('三、测试重点', level=1)
                for item in _ensure_list(stage3_json.get("test_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(f"- {item}")
                doc.add_heading('四、研发重点', level=1)
                for item in _ensure_list(stage3_json.get("dev_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(f"- {item}")
                doc.add_heading('五、项目风险', level=1)
                for item in _ensure_list(stage3_json.get("risks")) or ["【PRD未说明】"]:
                    doc.add_paragraph(f"- {item}")
                doc.add_heading('六、计划建议', level=1)
                for item in _ensure_list(stage3_json.get("plan")) or ["【PRD未说明】"]:
                    doc.add_paragraph(f"- {item}")
            else:
                lines = content.replace('\r\n', '\n').split('\n')
                i = 0
                while i < len(lines):
                    line = lines[i]
                    if line.strip().startswith('## '):
                        title = line.strip()[3:].strip()
                        doc.add_heading(title, level=1)
                    elif line.strip().startswith('### '):
                        title = line.strip()[4:].strip()
                        doc.add_heading(title, level=2)
                    elif line.strip():
                        doc.add_paragraph(line.rstrip())
                    else:
                        doc.add_paragraph()
                    i += 1
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            from flask import send_file
            return send_file(
                buf,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='PRD评审报告.docx'
            )
        except ImportError:
            return error_response('未安装 python-docx。请联系管理员: pip install python-docx', status_code=500)
        except Exception as e:
            logger.exception("导出 Word 失败: %s", e)
            return error_response('导出 Word 失败: ' + str(e), status_code=500)
    except Exception as e:
        logger.exception("export_report_docx 失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/analyze_prd', methods=['POST'])
def api_analyze_prd():
    try:
        data = request.get_json() or {}
        prd_text = (data.get('prd_text') or data.get('content') or '').strip()
        if not prd_text:
            return error_response('prd_text 不能为空', status_code=400)
        from modules.prd_audit import pipeline as prd_pipeline
        merged_report, stage1_output, stage2_output, stage3_output = prd_pipeline.run_prd_audit_sync(
            prd_text, llm_config_path=LLM_CONFIG_FILE, timeout=180
        )
        score = stage3_output.get("summary", {}).get("quality_score")
        quality_summary = {
            "overall": score,
            "dimensions": {},
            "rule_engine_score": score,
        }
        states = _ensure_list(stage1_output.get("states"))
        state_diagram_mermaid = ""
        if states:
            cleaned = [s.replace(" ", "_") for s in states]
            lines = ["stateDiagram-v2", f"  [*] --> {cleaned[0]}"]
            for i in range(len(cleaned) - 1):
                lines.append(f"  {cleaned[i]} --> {cleaned[i + 1]}")
            state_diagram_mermaid = "\n".join(lines)
        defects = stage3_output.get("defects") or []
        prd_gaps = "\n".join([f"- {d.get('module', '【PRD未说明】')}：{d.get('description', '【PRD未说明】')}" for d in defects[:20]])
        risks = "\n".join([f"- {x}" for x in (stage3_output.get("risks") or [])])
        rd_focus = "\n".join([f"- {x}" for x in (stage3_output.get("dev_focus") or [])])
        test_focus = "\n".join([f"- {x}" for x in (stage3_output.get("test_focus") or [])])
        plan_list = stage3_output.get("plan") or []
        plan_level = "P1"
        if any((d.get("risk_level") or "").upper() == "P0" for d in defects):
            plan_level = "P0"
        elif defects and all((d.get("risk_level") or "").upper() == "P2" for d in defects):
            plan_level = "P2"
        plan = {
            "level": plan_level,
            "order_advice": plan_list[:2],
            "prep_advice": plan_list[2:4],
        }
        section_1 = "## 一、总体结论\n\n- 质量评分：{}/10\n- 风险等级：{}\n- 主要问题：{}".format(
            stage3_output.get("summary", {}).get("quality_score", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("risk_level", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("main_problem", "【PRD未说明】"),
        )
        section_2 = "## 二、漏洞与风险清单\n\n" + (prd_gaps or "- 未发现漏洞")
        section_4 = "## 四、测试重点\n\n" + (test_focus or "- 【PRD未说明】")
        section_5 = "## 五、研发重点\n\n" + (rd_focus or "- 【PRD未说明】")
        section_6 = "## 六、计划建议\n\n" + ("\n".join([f"- {x}" for x in plan_list]) if plan_list else "- 【PRD未说明】")
        result = {
            "raw_report_markdown": merged_report,
            "status": ["Stage1完成", "Stage2完成", "Stage3完成", "Stage4完成"],
            "summary": section_1,
            "modules": stage1_output.get("modules") or [],
            "prd_gaps": prd_gaps,
            "risks": risks,
            "rd_focus": rd_focus,
            "test_focus": test_focus,
            "test_points": "",
            "plan": plan,
            "quality": quality_summary,
            "stage1_output": stage1_output,
            "stage2_output": stage2_output,
            "stage3_output": stage3_output,
            "system_model": {
                "states": stage1_output.get("states") or [],
                "events": [],
                "rules": stage1_output.get("business_rules") or [],
                "data_rules": stage1_output.get("exceptions") or [],
                "data_structures": stage1_output.get("data_structures") or [],
                "permissions": stage1_output.get("permissions") or [],
                "edge_cases": stage1_output.get("edge_cases") or [],
                "dependencies": stage1_output.get("dependencies") or [],
                "transitions": [],
                "high_priority_events": [],
            },
            "rule_analysis": stage2_output,
            "test_point_matrix": {"state_matrix": [], "concurrent_matrix": [], "boundary_matrix": []},
            "sections": {
                "1": section_1,
                "2": section_2,
                "3": "## 三、待确认清单\n\n- 【PRD未说明】",
                "4": section_4,
                "5": section_5,
                "6": section_6,
                "7": "## 七、优先级与责任建议\n\n- 按 P0/P1/P2 闭环处理",
                "8": "## 八、后续动作建议\n\n- 组织复审并更新测试点",
            },
            "state_diagram_mermaid": state_diagram_mermaid,
        }
        return success_response(data=result)
    except Exception as e:
        logger.exception("analyze_prd 失败")
        return error_response(str(e), status_code=500)

# ========== XMind 互转 API ==========

@test_case_bp.route('/api/import_xmind', methods=['POST'])
def api_import_xmind():
    """解析 XMind 文件，返回可导入的用例列表（供预览或调用 batch 接口创建）"""
    try:
        if 'file' not in request.files:
            return error_response('未上传文件', status_code=400)
        file = request.files['file']
        if not file.filename:
            return error_response('文件名为空', status_code=400)
        if not file.filename.lower().endswith('.xmind'):
            return error_response('请上传 .xmind 文件', status_code=400)
        from .xmind_conv import xmind_to_test_cases
        data = file.read()
        cases = xmind_to_test_cases(data)
        if not cases:
            return error_response('未从 XMind 中解析到测试用例（请确保根主题下有子节点）', status_code=400)
        return success_response(data={'cases': cases, 'count': len(cases)})
    except ImportError as e:
        return error_response(f'缺少依赖: 请执行 pip install xmindparser XMind', status_code=500)
    except Exception as e:
        logger.exception("XMind 导入失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/export_xmind', methods=['POST'])
def api_export_xmind():
    """导出测试用例为 XMind 文件"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        data = request.get_json() or {}
        case_ids = data.get('case_ids', [])
        export_all = data.get('export_all', False)
        if export_all:
            cases = test_case_storage.list_test_cases()
        elif case_ids:
            cases = [test_case_storage.get_test_case(cid) for cid in case_ids]
            cases = [c for c in cases if c]
        else:
            return error_response('请指定 case_ids 或 export_all=true', status_code=400)
        if not cases:
            return error_response('没有可导出的用例', status_code=400)
        from .xmind_conv import test_cases_to_xmind
        xmind_bytes = test_cases_to_xmind(cases)
        from io import BytesIO
        return send_file(
            BytesIO(xmind_bytes),
            mimetype='application/vnd.xmind.workbook',
            as_attachment=True,
            download_name='测试用例.xmind'
        )
    except ImportError as e:
        return error_response(f'缺少依赖: 请执行 pip install XMind', status_code=500)
    except Exception as e:
        logger.exception("XMind 导出失败")
        return error_response(str(e), status_code=500)

# ========== 测试用例 API ==========

@test_case_bp.route('/api/test-cases/batch', methods=['POST'])
def api_batch_create_test_cases():
    """批量创建测试用例"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        cases_data = data.get('cases', [])
        
        if not cases_data:
            return error_response('没有提供测试用例数据', status_code=400)
            
        created_cases = []
        errors = []
        
        for idx, case_data in enumerate(cases_data):
            try:
                case_id = test_case_storage.generate_case_id()
                
                # 处理步骤
                steps = []
                for step_data in case_data.get('steps', []):
                    steps.append(TestStep(
                        step_num=step_data.get('step_num', 0),
                        action=step_data.get('action', ''),
                        expected=step_data.get('expected', '')
                    ))
                
                test_case = TestCase(
                    id=case_id,
                    name=case_data.get('name', f'未命名用例_{idx+1}'),
                    description=case_data.get('description', ''),
                    category=case_data.get('category', '功能测试'),
                    tags=case_data.get('tags', []),
                    priority=case_data.get('priority', 'medium'),
                    status=case_data.get('status', 'active'),
                    steps=steps,
                    preconditions=case_data.get('preconditions', ''),
                    test_type=case_data.get('test_type', 'manual'),
                    related_package=case_data.get('related_package', ''),
                    created_by=case_data.get('created_by', 'system')
                )
                
                if test_case_storage.add_test_case(test_case):
                    created_cases.append(test_case.to_dict())
                else:
                    errors.append(f"用例 {case_data.get('name')} ID生成冲突")
            except Exception as e:
                logger.error(f"创建用例 {idx} 失败: {e}")
                errors.append(f"用例 {idx+1} 创建失败: {str(e)}")
        
        return success_response(data={
            'created_count': len(created_cases),
            'cases': created_cases,
            'errors': errors
        })
    except Exception as e:
        logger.exception("批量创建测试用例失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-cases', methods=['GET'])
def api_get_test_cases():
    """获取测试用例列表（支持筛选、分页、排序）"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        category = request.args.get('category', '').strip()
        tag = request.args.get('tag', '').strip()
        status = request.args.get('status', '').strip()
        search = request.args.get('search', '').strip()
        try:
            page = max(1, int(request.args.get('page', 1)))
            page_size = max(1, min(100, int(request.args.get('page_size', 20))))
        except (ValueError, TypeError):
            page, page_size = 1, 20
        sort_by = request.args.get('sort_by', 'updated_at')
        sort_order = request.args.get('sort_order', 'desc')
        
        cases = test_case_storage.list_test_cases(
            category=category if category else None,
            tag=tag if tag else None,
            status=status if status else None,
            search=search if search else None
        )
        
        reverse_order = (sort_order.lower() == 'desc')
        if sort_by == 'name':
            cases.sort(key=lambda x: x.name, reverse=reverse_order)
        elif sort_by == 'created_at':
            cases.sort(key=lambda x: x.created_at or datetime.min, reverse=reverse_order)
        elif sort_by == 'updated_at':
            cases.sort(key=lambda x: x.updated_at or datetime.min, reverse=reverse_order)
        
        total = len(cases)
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size
        paginated = cases[start_idx:end_idx]
        
        return success_response(data={
            'test_cases': [case.to_dict() for case in paginated],
            'pagination': {
                'page': page,
                'page_size': page_size,
                'total': total,
                'total_pages': (total + page_size - 1) // page_size
            }
        })
    except Exception as e:
        logger.exception("获取测试用例列表失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-cases', methods=['POST'])
def api_create_test_case():
    """创建测试用例"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        case_id = test_case_storage.generate_case_id()
        
        steps = []
        for step_data in data.get('steps', []):
            steps.append(TestStep(
                step_num=step_data.get('step_num', 0),
                action=step_data.get('action', ''),
                expected=step_data.get('expected', '')
            ))
        
        test_case = TestCase(
            id=case_id,
            name=data.get('name', ''),
            description=data.get('description', ''),
            category=data.get('category', '功能测试'),
            tags=data.get('tags', []),
            priority=data.get('priority', 'medium'),
            status=data.get('status', 'active'),
            steps=steps,
            preconditions=data.get('preconditions', ''),
            test_type=data.get('test_type', 'manual'),
            related_package=data.get('related_package', ''),
            created_by=data.get('created_by', 'system')
        )
        
        if test_case_storage.add_test_case(test_case):
            return success_response(data={'test_case': test_case.to_dict()})
        else:
            return error_response('用例ID已存在', status_code=400)
    except Exception as e:
        logger.exception("创建测试用例失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-cases/<case_id>', methods=['GET'])
def api_get_test_case(case_id):
    """获取测试用例详情"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        test_case = test_case_storage.get_test_case(case_id)
        if test_case:
            return success_response(data={'test_case': test_case.to_dict()})
        else:
            return error_response('用例不存在', status_code=404)
    except Exception as e:
        logger.exception("获取测试用例详情失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-cases/<case_id>', methods=['PUT'])
def api_update_test_case(case_id):
    """更新测试用例"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        test_case = test_case_storage.get_test_case(case_id)
        if not test_case:
            return error_response('用例不存在', status_code=404)
        
        if 'name' in data:
            test_case.name = data['name']
        if 'description' in data:
            test_case.description = data['description']
        if 'category' in data:
            test_case.category = data['category']
        if 'tags' in data:
            test_case.tags = data['tags']
        if 'priority' in data:
            test_case.priority = data['priority']
        if 'status' in data:
            test_case.status = data['status']
        if 'steps' in data:
            test_case.steps = [TestStep.from_dict(s) for s in data['steps']]
        if 'preconditions' in data:
            test_case.preconditions = data['preconditions']
        if 'test_type' in data:
            test_case.test_type = data['test_type']
        if 'related_package' in data:
            test_case.related_package = data['related_package']
        
        test_case.updated_at = datetime.now()
        
        if test_case_storage.update_test_case(test_case):
            return success_response(data={'test_case': test_case.to_dict()})
        else:
            return error_response('更新失败', status_code=500)
    except Exception as e:
        logger.exception("更新测试用例失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-cases/<case_id>', methods=['DELETE'])
def api_delete_test_case(case_id):
    """删除测试用例"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        if test_case_storage.delete_test_case(case_id):
            return success_response(message='删除成功')
        else:
            return error_response('用例不存在', status_code=404)
    except Exception as e:
        logger.exception("删除测试用例失败")
        return error_response(str(e), status_code=500)

# ========== 测试套件 API ==========

@test_case_bp.route('/api/test-suites', methods=['GET'])
def api_get_test_suites():
    """获取测试套件列表"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        suites = test_case_storage.list_test_suites()
        return success_response(data={'test_suites': [suite.to_dict() for suite in suites]})
    except Exception as e:
        logger.exception("获取测试套件列表失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-suites', methods=['POST'])
def api_create_test_suite():
    """创建测试套件"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        suite_id = test_case_storage.generate_suite_id()
        
        test_suite = TestSuite(
            id=suite_id,
            name=data.get('name', ''),
            description=data.get('description', ''),
            test_case_ids=data.get('test_case_ids', [])
        )
        
        if test_case_storage.add_test_suite(test_suite):
            return success_response(data={'test_suite': test_suite.to_dict()})
        else:
            return error_response('套件ID已存在', status_code=400)
    except Exception as e:
        logger.exception("创建测试套件失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-suites/<suite_id>', methods=['GET'])
def api_get_test_suite(suite_id):
    """获取测试套件详情"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        test_suite = test_case_storage.get_test_suite(suite_id)
        if test_suite:
            cases = []
            for case_id in test_suite.test_case_ids:
                case = test_case_storage.get_test_case(case_id)
                if case:
                    cases.append(case.to_dict())
            
            suite_dict = test_suite.to_dict()
            suite_dict['test_cases'] = cases
            return success_response(data={'test_suite': suite_dict})
        else:
            return error_response('套件不存在', status_code=404)
    except Exception as e:
        logger.exception("获取测试套件详情失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-suites/<suite_id>', methods=['PUT'])
def api_update_test_suite(suite_id):
    """更新测试套件"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        test_suite = test_case_storage.get_test_suite(suite_id)
        if not test_suite:
            return error_response('套件不存在', status_code=404)
        
        if 'name' in data:
            test_suite.name = data['name']
        if 'description' in data:
            test_suite.description = data['description']
        if 'test_case_ids' in data:
            test_suite.test_case_ids = data['test_case_ids']
        
        test_suite.updated_at = datetime.now()
        
        if test_case_storage.update_test_suite(test_suite):
            return success_response(data={'test_suite': test_suite.to_dict()})
        else:
            return error_response('更新失败', status_code=500)
    except Exception as e:
        logger.exception("更新测试套件失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-suites/<suite_id>', methods=['DELETE'])
def api_delete_test_suite(suite_id):
    """删除测试套件"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        if test_case_storage.delete_test_suite(suite_id):
            return success_response(message='删除成功')
        else:
            return error_response('套件不存在', status_code=404)
    except Exception as e:
        logger.exception("删除测试套件失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-suites/<suite_id>/execute', methods=['POST'])
def api_execute_test_suite(suite_id):
    """执行测试套件"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        device_id = data.get('device_id', '')
        package_name = data.get('package_name', '')
        
        test_suite = test_case_storage.get_test_suite(suite_id)
        if not test_suite:
            return error_response('套件不存在', status_code=404)
        
        execution_ids = []
        for case_id in test_suite.test_case_ids:
            exec_id = test_case_storage.generate_execution_id()
            execution = TestCaseExecution(
                id=exec_id,
                test_case_id=case_id,
                test_suite_id=suite_id,
                device_id=device_id,
                package_name=package_name,
                status='running',
                executor='system'
            )
            test_case_storage.add_execution(execution)
            execution_ids.append(exec_id)
        
        return success_response(
            message=f'已创建 {len(execution_ids)} 个执行记录',
            data={'execution_ids': execution_ids}
        )
    except Exception as e:
        logger.exception("执行测试套件失败")
        return error_response(str(e), status_code=500)

# ========== 执行记录 API ==========

@test_case_bp.route('/api/test-executions', methods=['GET'])
def api_get_test_executions():
    """获取执行记录列表（支持筛选、分页）"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        test_case_id = request.args.get('test_case_id', '').strip()
        test_suite_id = request.args.get('test_suite_id', '').strip()
        device_id = request.args.get('device_id', '').strip()
        status = request.args.get('status', '').strip()
        start_date = request.args.get('start_date', '').strip()
        end_date = request.args.get('end_date', '').strip()
        try:
            page = max(1, int(request.args.get('page', 1)))
            page_size = max(1, min(100, int(request.args.get('page_size', 20))))
        except (ValueError, TypeError):
            page, page_size = 1, 20
        
        start_dt = None
        end_dt = None
        if start_date:
            try:
                start_dt = datetime.fromisoformat(start_date.replace('Z', '+00:00'))
            except Exception:
                pass
        if end_date:
            try:
                end_dt = datetime.fromisoformat(end_date.replace('Z', '+00:00'))
            except Exception:
                pass
        
        executions = test_case_storage.list_executions(
            test_case_id=test_case_id if test_case_id else None,
            test_suite_id=test_suite_id if test_suite_id else None,
            device_id=device_id if device_id else None,
            status=status if status else None,
            start_date=start_dt,
            end_date=end_dt
        )
        
        total = len(executions)
        start_idx = (page - 1) * page_size
        end_idx = start_idx + page_size
        paginated = executions[start_idx:end_idx]
        
        return success_response(data={
            'executions': [execution.to_dict() for execution in paginated],
            'pagination': {
                'page': page,
                'page_size': page_size,
                'total': total,
                'total_pages': (total + page_size - 1) // page_size
            }
        })
    except Exception as e:
        logger.exception("获取执行记录列表失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-executions', methods=['POST'])
def api_create_test_execution():
    """创建执行记录"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        exec_id = test_case_storage.generate_execution_id()
        
        execution = TestCaseExecution(
            id=exec_id,
            test_case_id=data.get('test_case_id', ''),
            test_suite_id=data.get('test_suite_id'),
            device_id=data.get('device_id', ''),
            package_name=data.get('package_name', ''),
            status=data.get('status', 'running'),
            executor=data.get('executor', 'system')
        )
        
        if test_case_storage.add_execution(execution):
            return success_response(data={'execution': execution.to_dict()})
        else:
            return error_response('创建失败', status_code=500)
    except Exception as e:
        logger.exception("创建执行记录失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-executions/<exec_id>', methods=['GET'])
def api_get_test_execution(exec_id):
    """获取执行记录详情"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        execution = test_case_storage.get_execution(exec_id)
        if execution:
            test_case = test_case_storage.get_test_case(execution.test_case_id)
            execution_dict = execution.to_dict()
            if test_case:
                execution_dict['test_case'] = test_case.to_dict()
            
            return success_response(data={'execution': execution_dict})
        else:
            return error_response('执行记录不存在', status_code=404)
    except Exception as e:
        logger.exception("获取执行记录详情失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-executions/<exec_id>', methods=['PUT'])
def api_update_test_execution(exec_id):
    """更新执行记录"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        data = request.get_json() or {}
        execution = test_case_storage.get_execution(exec_id)
        if not execution:
            return error_response('执行记录不存在', status_code=404)
        
        if 'status' in data:
            execution.status = data['status']
            if execution.status in ['passed', 'failed', 'skipped'] and execution.end_time is None:
                execution.end_time = datetime.now()
                if execution.start_time:
                    execution.duration = (execution.end_time - execution.start_time).total_seconds()
        
        if 'step_results' in data:
            execution.step_results = data['step_results']
        if 'screenshots' in data:
            execution.screenshots = data['screenshots']
        if 'logs' in data:
            execution.logs = data['logs']
        if 'errors' in data:
            execution.errors = data['errors']
        if 'notes' in data:
            execution.notes = data['notes']
        if 'related_monkey_test' in data:
            execution.related_monkey_test = data['related_monkey_test']
        if 'related_performance_session' in data:
            execution.related_performance_session = data['related_performance_session']
        if 'related_log_monitor_session' in data:
            execution.related_log_monitor_session = data['related_log_monitor_session']
        
        if test_case_storage.update_execution(execution):
            return success_response(data={'execution': execution.to_dict()})
        else:
            return error_response('更新失败', status_code=500)
    except Exception as e:
        logger.exception("更新执行记录失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-executions/<exec_id>/generate_defect', methods=['POST'])
def api_generate_defect_from_execution(exec_id):
    """失败用例 → 缺陷描述：根据执行记录与用例步骤用 LLM 生成缺陷标题、复现步骤、预期/实际。"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        execution = test_case_storage.get_execution(exec_id)
        if not execution:
            return error_response('执行记录不存在', status_code=404)
        if execution.status != 'failed':
            return error_response('仅支持失败状态的执行记录', status_code=400)

        test_case = test_case_storage.get_test_case(execution.test_case_id)
        steps = []
        if test_case and test_case.steps:
            step_results_map = {s.get('step_num'): s for s in (execution.step_results or []) if isinstance(s, dict)}
            for step in test_case.steps:
                sr = step_results_map.get(step.step_num) or {}
                steps.append({
                    'step_num': step.step_num,
                    'action': step.action,
                    'expected': step.expected,
                    'actual': sr.get('actual') or step.actual or '',
                    'status': sr.get('status') or step.status or '',
                })
        else:
            steps = execution.step_results or []

        context_parts = [
            f"用例名称: {test_case.name if test_case else execution.test_case_id}",
            f"用例描述: {(test_case.description or '')[:500]}",
            f"设备: {execution.device_id}  包名: {execution.package_name}",
            "步骤与结果:",
        ]
        for s in steps:
            if isinstance(s, dict):
                context_parts.append(f"  步骤{s.get('step_num', '')}: {s.get('action', '')} | 预期: {s.get('expected', '')} | 实际: {s.get('actual', '')} | 状态: {s.get('status', '')}")
            else:
                context_parts.append(f"  {s}")
        if execution.errors:
            context_parts.append("执行错误: " + "\n".join(execution.errors[:10]))
        if execution.logs:
            context_parts.append("日志片段: " + "\n".join(execution.logs[-15:]))
        if execution.notes:
            context_parts.append("备注: " + execution.notes[:300])
        context = "\n".join(context_parts)[:6000]

        try:
            from utils.llm_client import call_llm
        except Exception:
            return error_response('LLM 未配置，请先在用例管理中配置 LLM', status_code=503)

        prompt = (
            "根据以下失败用例执行记录，生成缺陷描述，用中文按以下格式输出（不要 markdown 标题符号）：\n"
            "缺陷标题：<一句话概括问题>\n"
            "复现步骤：<按步骤列出操作>\n"
            "预期结果：<应出现的结果>\n"
            "实际结果：<实际出现的现象或报错>\n"
            "环境/备注：<设备、包名等可选>\n\n"
            "---\n\n" + context
        )
        result = call_llm([{"role": "user", "content": prompt}], config_path=LLM_CONFIG_FILE, timeout=60)
        text = (result or "").strip()
        if not text:
            return error_response('LLM 未返回内容', status_code=500)
        return success_response(data={'defect_description': text})
    except Exception as e:
        logger.exception("生成缺陷描述失败")
        return error_response(str(e), status_code=500)

@test_case_bp.route('/api/test-executions/statistics', methods=['GET'])
def api_get_execution_statistics():
    """获取执行统计"""
    try:
        if not test_case_storage:
            return error_response('测试用例存储未初始化', status_code=500)
        
        test_case_id = request.args.get('test_case_id', '').strip()
        test_suite_id = request.args.get('test_suite_id', '').strip()
        
        stats = test_case_storage.get_execution_statistics(
            test_case_id=test_case_id if test_case_id else None,
            test_suite_id=test_suite_id if test_suite_id else None
        )
        
        return success_response(data={'statistics': stats})
    except Exception as e:
        logger.exception("获取执行统计失败")
        return error_response(str(e), status_code=500)
