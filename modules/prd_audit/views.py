"""
PRD 审计独立模块：自包含实现，不依赖 test_case。
提供独立页与全部 API（生成、分析、解析 PDF/DOCX、导出 Word、默认提示词、LLM 配置、保存为测试用例）。
"""

import os
import json
import io
import re
import csv
import zipfile
import uuid
import threading
import subprocess
from datetime import datetime, timezone
import time
from collections import Counter
from typing import Any, Dict, List
from flask import render_template, request, session, url_for, Response, stream_with_context, send_file
from utils.response import success_response, error_response
from utils.logger import setup_logger
from utils.llm_client import call_llm
from . import prd_audit_bp
from . import pipeline
from .rule_plugin_engine import load_rule_plugin_profiles, save_rule_plugin_profiles, get_plugin_usage_stats
from .prompt_center import load_prompt_center, save_prompt_center, get_prompt_evaluation_stats
from .outline_engine import run_outline_engine
from .outline_llm import run_outline_llm, merge_llm_with_local
from .feishu_client import is_feishu_doc_url, fetch_feishu_doc_content
from .audit_learning import (
    get_learning_status,
    get_learning_lane_stats,
    get_learning_quality_dashboard,
    build_rule_draft_from_snapshots,
    load_rule_candidates,
    apply_selected_candidates,
    publish_applied_rules,
    list_rule_backups,
    rollback_rules_from_backup,
    save_outline_owner_correction,
)
from . import review_meeting_v2 as rmv2
from . import prd_rewrite_engine as prd_rewrite

logger = setup_logger("prd_audit_module")

# V2 async task store (in-memory, MVP; V2 only)
_RMV2_TASKS_LOCK = threading.Lock()
_RMV2_TASKS: Dict[str, Dict[str, Any]] = {}

def _rmv2_task_create(meta: Dict[str, Any]) -> str:
    tid = f"rmv2task_{uuid.uuid4().hex[:10]}"
    with _RMV2_TASKS_LOCK:
        _RMV2_TASKS[tid] = {
            "task_id": tid,
            "state": "RUNNING",  # RUNNING|DONE|ERROR
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "updated_at": datetime.now().isoformat(timespec="seconds"),
            "progress": [],
            "result": None,
            "error": "",
            "meta": meta or {},
        }
    return tid

def _rmv2_task_append(tid: str, text: str) -> None:
    msg = str(text or "").strip()
    if not msg:
        return
    with _RMV2_TASKS_LOCK:
        t = _RMV2_TASKS.get(tid)
        if not isinstance(t, dict):
            return
        t["updated_at"] = datetime.now().isoformat(timespec="seconds")
        t.setdefault("progress", [])
        t["progress"].append(msg)
        if len(t["progress"]) > 60:
            t["progress"] = t["progress"][-60:]

def _rmv2_task_finish(tid: str, *, result: Dict[str, Any]) -> None:
    with _RMV2_TASKS_LOCK:
        t = _RMV2_TASKS.get(tid)
        if not isinstance(t, dict):
            return
        t["state"] = "DONE"
        t["updated_at"] = datetime.now().isoformat(timespec="seconds")
        t["result"] = result

def _rmv2_task_fail(tid: str, err: str) -> None:
    with _RMV2_TASKS_LOCK:
        t = _RMV2_TASKS.get(tid)
        if not isinstance(t, dict):
            return
        t["state"] = "ERROR"
        t["updated_at"] = datetime.now().isoformat(timespec="seconds")
        t["error"] = str(err or "启动失败")

def _rmv2_task_set_partial(tid: str, partial: Dict[str, Any]) -> None:
    with _RMV2_TASKS_LOCK:
        t = _RMV2_TASKS.get(tid)
        if not isinstance(t, dict):
            return
        t["updated_at"] = datetime.now().isoformat(timespec="seconds")
        t["partial"] = partial

STORAGE_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(STORAGE_DIR, exist_ok=True)
DEFAULT_PRD_AUDIT_PROMPT_FILE = os.path.join(STORAGE_DIR, "prd_audit_prompt_default.txt")
# 全平台共用的大模型配置（统一入口）
# 平台级 LLM 配置单一真源（与 utils.llm_client.DEFAULT_LLM_CONFIG 保持一致）
PLATFORM_LLM_CONFIG_FILE = os.path.abspath(os.path.join(
    os.path.dirname(__file__), '..', '..', 'config', 'llm_config.json'
))
LLM_CONFIG_FILE = PLATFORM_LLM_CONFIG_FILE

FALLBACK_PRD_PROMPT = (
    "你是 PRD 审计专家。请基于输入 PRD 输出结构化风险报告，至少包含："
    "总体结论、漏洞与风险清单、测试重点、研发重点、计划建议。"
    "仅输出报告正文，不要开场白。\n\nPRD 内容：\n{content}"
)

KTV_CAPABILITY_CARDS = [
    {
        "capability_id": "KTV_PLAY_CTRL_001",
        "name": "切歌控制",
        "domain": "KTV点歌系统",
        "category": "播放控制",
        "priority": "P0",
        "trigger": "当前歌曲播放中，管理员点击“切歌”",
        "preconditions": [
            "管理员权限=是",
            "当前队列长度>=1",
        ],
        "system_behaviors": [
            "当前歌曲在500ms内停止播放",
            "队列下一首在2s内开始播放",
            "所有终端同步更新“当前播放歌曲”",
        ],
        "exceptions": [
            {"condition": "下一首解码失败", "action": "自动跳过并播放再下一首，记录错误码"},
            {"condition": "队列为空", "action": "进入待机页并提示“暂无歌曲”"},
        ],
        "logging": [
            "操作者ID",
            "房间ID",
            "歌曲ID",
            "操作时间",
            "结果码",
        ],
        "acceptance_criteria": [
            "100次切歌中，95次在2s内成功切换",
        ],
        "required_evidence": [
            "PRD中明确管理员切歌权限",
            "PRD中明确切歌时延SLA（500ms停止、2s起播）",
            "PRD中明确多终端同步规则",
            "PRD中明确异常分支（解码失败、队列为空）",
            "PRD中明确日志字段与验收口径",
        ],
        "bad_patterns": [
            "只写切歌功能，未写权限约束",
            "只写正常流程，未写异常回退",
            "未定义时延指标或验收标准",
            "未定义终端状态同步",
        ],
        "suggestion_template": "补充切歌控制时序图：鉴权校验→停止当前→拉起下一首→多端同步；并补充异常回退与日志字段定义。",
    }
]
KNOWLEDGE_CARDS_FILE = os.path.join(STORAGE_DIR, "knowledge_cards.json")
PRD_RULES_V2_FILE = os.path.join(STORAGE_DIR, "prd_scan_rules_v2.json")
PRD_RULES_V1_FILE = os.path.join(STORAGE_DIR, "prd_scan_rules.json")
BUG_RAW_FILE = os.path.join(STORAGE_DIR, "bug_raw.json")
BUG_PATTERN_FILE = os.path.join(STORAGE_DIR, "bug_pattern.json")
BUG_RULE_FILE = os.path.join(STORAGE_DIR, "bug_rule.json")
VECTOR_DATA_FILE = os.path.join(STORAGE_DIR, "vector_data.json")
GATE_CONFIG_FILE = os.path.join(STORAGE_DIR, "gate_config.json")
REVIEW_STATE_FILE = os.path.join(STORAGE_DIR, "learning_repo", "review_state.json")
FEISHU_WATCH_FILE = os.path.join(STORAGE_DIR, "learning_repo", "feishu_watch.json")

# -------- 虚拟评审会（独立于静态审计）--------
# 说明：会议态只保存在内存。输入可为 learning_repo/snapshots（snapshot 模式），或 brainstorm（无 PRD 推断草案，默认 BLOCKED）。
_REVIEW_MEETINGS: Dict[str, Dict[str, Any]] = {}

@prd_audit_bp.route("/api/review_meeting_v2/build_info", methods=["GET"])
def api_review_meeting_v2_build_info():
    """
    Debug helper: show which code is actually running (avoid "edited but not effective").
    Platform-generic; no secrets.
    """
    try:
        here = os.path.abspath(__file__)
        rmv2_path = os.path.abspath(getattr(rmv2, "__file__", "") or "")
        tpl_path = os.path.abspath(os.path.join(os.path.dirname(__file__), "templates", "prd_audit_review_meeting_v2.html"))
        now = datetime.now(timezone.utc).isoformat()

        def _mtime(p: str) -> str:
            try:
                ts = os.path.getmtime(p)
                return datetime.fromtimestamp(ts, tz=timezone.utc).isoformat()
            except Exception:
                return ""

        commit = ""
        try:
            repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
            out = subprocess.check_output(["git", "rev-parse", "--short", "HEAD"], cwd=repo_root, stderr=subprocess.DEVNULL)
            commit = out.decode("utf-8", errors="ignore").strip()
        except Exception:
            commit = ""

        return success_response(
            data={
                "now_utc": now,
                "git_commit": commit,
                "paths": {
                    "views_py": here,
                    "views_mtime_utc": _mtime(here),
                    "review_meeting_v2_py": rmv2_path,
                    "review_meeting_v2_mtime_utc": _mtime(rmv2_path) if rmv2_path else "",
                    "template_v2_html": tpl_path,
                    "template_v2_mtime_utc": _mtime(tpl_path),
                },
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_v2_build_info failed")
        return error_response(str(e), status_code=500)

def _normalize_review_meeting_mode(data: Dict[str, Any]) -> str:
    m = str((data or {}).get("mode") or "snapshot").strip().lower()
    if m in ("brainstorm", "sandbox", "no_doc", "no_document", "nodoc"):
        return "brainstorm"
    return "snapshot"

def _brainstorm_agenda_cards(idea: str, use_llm: bool) -> List[Dict[str, Any]]:
    """
    无 PRD 文档的沙盘议题：证据一律 inference，议题卡用于引导多角色讨论；
    裁决侧默认 BLOCKED，直到补齐文档锚点。
    """
    idea = (idea or "").strip()
    agenda_cards: List[Dict[str, Any]] = []

    def _one_card(idx: int, topic: str) -> Dict[str, Any]:
        return {
            "session_id": "brainstorm",
            "meeting_mode": "brainstorm",
            "issue_id": f"brainstorm-P0-{idx:02d}",
            "topic": topic,
            "risk_level": "P0",
            "one_liner": idea[:120] if idx == 1 else f"{idea[:80]}（{topic}）",
            "evidence": [{"type": "inference", "ref": "", "text": "无 PRD 原文输入：本卡为推断草案"}],
            "decision_needed": "请拍板阈值/终态/观测字段；无文档前仅记录草案，不关闭 P0。",
            "candidate_ac": ["Given... When... Then...(含量化项，数字为暂定，需文档固化)"],
            "constraints": ["不得将推断结论当作已发布 PRD 事实"],
            "open_questions": ["缺少 PRD 章节/锚点：需补齐后方可关闭 P0"],
            "source": {"idea": idea},
        }

    if use_llm:
        try:
            sys_p = (
                "你是测试主导的评审会秘书。请仅输出 JSON：{\"cards\":[...]}。"
                "每个 card 含字段：topic（exception/state/security/conflict/cross_end/data_contract）、one_liner、decision_needed。"
                "不要引用不存在的 PRD 原文或行号。"
            )
            user_p = (
                "输入是一句话需求/议题：\n"
                f"{idea}\n\n"
                "请生成最多 3 条、且 topic 互不重复的议题卡。"
            )
            obj = _llm_call_json(sys_p, user_p, timeout=60, max_tokens=1200)
            cards = obj.get("cards") if isinstance(obj.get("cards"), list) else []
            used_topics: set = set()
            for c in cards:
                if not isinstance(c, dict) or len(agenda_cards) >= 3:
                    break
                topic = str(c.get("topic") or "generic").strip()
                if topic in used_topics:
                    continue
                used_topics.add(topic)
                agenda_cards.append(
                    {
                        "session_id": "brainstorm",
                        "meeting_mode": "brainstorm",
                        "issue_id": f"brainstorm-P0-{len(agenda_cards)+1:02d}",
                        "topic": topic,
                        "risk_level": "P0",
                        "one_liner": str(c.get("one_liner") or idea)[:120],
                        "evidence": [{"type": "inference", "ref": "", "text": "无 PRD 原文输入：推断草案"}],
                        "decision_needed": str(c.get("decision_needed") or "请拍板阈值/终态/观测字段。")[:300],
                        "candidate_ac": ["Given... When... Then...(含量化项，数字为暂定，需文档固化)"],
                        "constraints": ["不得将推断结论当作已发布 PRD 事实"],
                        "open_questions": ["缺少 PRD 章节/锚点：需补齐后方可关闭 P0"],
                        "source": {"idea": idea},
                    }
                )
        except Exception:
            agenda_cards = []

    if len(agenda_cards) < 3:
        used = {str(c.get("topic") or "") for c in agenda_cards}
        fill_topics = ["exception", "state", "conflict", "security", "cross_end", "data_contract"]
        for tp in fill_topics:
            if len(agenda_cards) >= 3:
                break
            if tp in used:
                continue
            used.add(tp)
            agenda_cards.append(_one_card(len(agenda_cards) + 1, tp))
    return agenda_cards[:3]

def _local_scenarios_for_card(card: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Scenario Generator（本地模板兜底）：
    为单个议题卡生成可直接用于 QA 追问的 Given/When/Then 场景。
    说明：这里只服务“虚拟评审会”，不参与静态审计的报告生成。
    """
    topic = str(card.get("topic") or "generic").strip().lower()
    one_liner = str(card.get("one_liner") or "").strip()
    subject = one_liner[:26] if one_liner else "该能力"
    ev0 = ""
    evidence = card.get("evidence") if isinstance(card.get("evidence"), list) else []
    if evidence and isinstance(evidence[0], dict):
        ev0 = str(evidence[0].get("ref") or "") or str(evidence[0].get("type") or "")

    common_obs = ["session_id", "request_id", "state", "error_code"]
    common_required = ["PRD 明确阈值/终态/错误码/提示文案", "PRD 明确最小观测字段（日志/指标）"]

    def s(given: str, when: str, then: str, risk: str = "P1", required=None, obs=None, tags=None):
        return {
            "title": f"{subject}场景推演",
            "risk_level": risk,
            "given": given,
            "when": when,
            "then": then,
            "required_evidence": (required or []) + common_required,
            "observability_min": {"logs": (obs or []) + common_obs, "metrics": [], "traces": []},
            "tags": tags or [],
            "anchors_hint": ev0,
        }

    out: List[Dict[str, Any]] = []
    if topic in ("exception", "exception_flow", "cloud_degrade"):
        out.append(
            s(
                "Given 弱网/高延迟",
                "When 关键请求超时并触发重试",
                "Then 必须进入明确终态（失败/降级/待上传），并对用户有可见提示；不得静默失败。",
                risk="P0",
                required=["PRD 定义超时阈值T、重试次数N、退避Δ、失败提示口径"],
                obs=["retry_count", "timeout_ms", "degrade_mode"],
                tags=["弱网", "超时", "重试", "降级"],
            )
        )
        out.append(
            s(
                "Given 结果需要上传/同步",
                "When 上传失败或服务返回 5xx",
                "Then 本地缓存策略/重试边界/最终一致性窗口必须明确；否则手机端与 TV 端列表会不一致。",
                risk="P0",
                required=["PRD 明确缓存保留时长、重试窗口、失败后用户可见状态"],
                obs=["upload_state", "http_status", "queue_size"],
                tags=["上传", "一致性"],
            )
        )
    elif topic in ("state", "device_lifecycle", "transfer_cleanup"):
        out.append(
            s(
                "Given 产生了未认领数据/中间态结果",
                "When 发生重启/关台/重开台/断电上电",
                "Then 必须按 PRD 定义清空或保留；并给出“转台不清空”等例外的对偶规则。",
                risk="P0",
                required=["PRD 明确清空触发条件清单", "PRD 明确不清空例外（如转台）"],
                obs=["exit_reason", "cleanup_items_applied", "final_state"],
                tags=["重启", "清空", "转台"],
            )
        )
        out.append(
            s(
                "Given 用户中途退出/切后台/重进",
                "When 状态机处于处理中/上传中/保存中",
                "Then 必须定义是否允许恢复、恢复入口、超时收敛与最终状态；否则出现“卡死/重复执行”。",
                risk="P1",
                required=["PRD 提供状态机转移表或等价描述"],
                obs=["state_before", "state_after", "state_updated_at"],
                tags=["状态机", "恢复"],
            )
        )
    elif topic in ("conflict", "rating_switch", "dispatch"):
        out.append(
            s(
                "Given 同时存在自动规则与手动开关",
                "When 用户在播控栏临时关闭，但自动规则再次触发",
                "Then 必须按优先级裁决表决策：谁生效、何时生效、UI 图标与真实能力是否一致。",
                risk="P0",
                required=["PRD 给出优先级裁决表（总开关/播控栏/自动规则）"],
                obs=["global_switch", "session_switch", "effective_switch", "conflict_reason"],
                tags=["优先级", "开关冲突"],
            )
        )
        out.append(
            s(
                "Given 设置页总开关关闭",
                "When 手机端触发“快唱副歌”等自动入口",
                "Then 必须明确是否允许强制开启；若不允许，应提示原因并记录审计日志。",
                risk="P0",
                required=["PRD 明确总开关是否具备最高否决权"],
                obs=["actor", "entry", "auth_result"],
                tags=["总开关", "自动入口"],
            )
        )
    elif topic in ("security", "security_access", "qr_security"):
        out.append(
            s(
                "Given 外链/扫码/分享入口存在",
                "When 未授权用户访问或 token 过期",
                "Then 必须拦截并给出明确提示与审计日志，禁止越权读取。",
                risk="P0",
                required=["PRD 明确鉴权对象、资源范围、token 过期策略与错误码"],
                obs=["subject_id", "resource_id", "auth_result", "audit_log_id"],
                tags=["鉴权", "越权"],
            )
        )
    else:
        out.append(
            s(
                "Given 核心流程存在关键阈值",
                "When 达到/未达到阈值边界（例如 10s）",
                "Then 必须定义保存/丢弃/提示的可验收口径，并写清最小观测字段。",
                risk="P1",
                required=["PRD 明确阈值与边界处理策略"],
                obs=["threshold", "final_state"],
                tags=["阈值", "边界"],
            )
        )

    # 控制数量 3-6 条：补齐通用“并发”场景
    out.append(
        s(
            "Given 多入口/多终端可同时操作",
            "When 两个终端同时触发同一能力（并发）",
            "Then 必须定义幂等/去重/锁/最终态；否则出现状态竞争与结果错乱。",
            risk="P0" if topic in ("conflict", "state", "exception") else "P1",
            required=["PRD 明确并发裁决与幂等策略"],
            obs=["idempotency_key", "lock_acquired", "dedup_hit"],
            tags=["并发", "幂等"],
        )
    )
    # 去掉过长/空项
    clean = []
    for it in out:
        if not isinstance(it, dict):
            continue
        if str(it.get("given") or "").strip() and str(it.get("when") or "").strip() and str(it.get("then") or "").strip():
            clean.append(it)
    return clean[:6]

def _llm_scenarios_for_card(card: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Scenario Generator（LLM增强）：输出 JSON {"scenarios":[...]}，失败抛异常，上层降级到本地模板。
    """
    one_liner = str(card.get("one_liner") or "").strip()
    topic = str(card.get("topic") or "generic").strip()
    meeting_mode = str(card.get("meeting_mode") or "snapshot").strip().lower()
    evidence = card.get("evidence") or []
    ev_text = "\n".join(
        [f"- {e.get('type')} {e.get('ref')}: {e.get('text') or ''}".strip() for e in evidence if isinstance(e, dict)]
    ) or "- 无"

    sys_p = (
        "你是测试组长，专门负责制造“麻烦场景”，让 PRD 的缺口暴露出来。"
        "请严格输出 JSON：{\"scenarios\":[...]}，禁止输出任何多余文字。必须中文。"
    )
    if meeting_mode == "brainstorm":
        sys_p += "【无文档沙盘】禁止捏造原文锚点/章节名；required_evidence 必须写满。"
    user_p = (
        f"议题：{one_liner}\n"
        f"Topic：{topic}\n"
        f"证据：\n{ev_text}\n\n"
        "请输出 4-6 条场景，每条场景结构：\n"
        "{\n"
        '  "title": "一句话场景名",\n'
        '  "risk_level": "P0|P1|P2",\n'
        '  "given": "Given ...",\n'
        '  "when": "When ...",\n'
        '  "then": "Then ...（必须是可验收口径，含阈值/终态/提示/观测字段至少其一）",\n'
        '  "required_evidence": ["需要PRD补充的证据/定义"],\n'
        '  "observability_min": {"logs":["..."],"metrics":["..."],"traces":["..."]},\n'
        '  "tags": ["弱网","超时","并发","清空","优先级"]\n'
        "}\n"
        "要求：优先覆盖 弱网/超时/并发/清空与保留/多开关优先级。"
    )
    obj = _llm_call_json(sys_p, user_p, timeout=60, max_tokens=1600)
    arr = obj.get("scenarios") if isinstance(obj, dict) else None
    arr = arr if isinstance(arr, list) else []
    out = []
    for it in arr[:8]:
        if not isinstance(it, dict):
            continue
        # 最小字段校验
        if not str(it.get("given") or "").strip():
            continue
        if not str(it.get("when") or "").strip():
            continue
        if not str(it.get("then") or "").strip():
            continue
        out.append(it)
    return out[:6]

def _attach_scenarios_to_cards(agenda_cards: List[Dict[str, Any]], use_llm: bool) -> List[Dict[str, Any]]:
    if not isinstance(agenda_cards, list) or not agenda_cards:
        return agenda_cards
    for c in agenda_cards:
        if not isinstance(c, dict):
            continue
        scenarios = []
        if use_llm:
            try:
                scenarios = _llm_scenarios_for_card(c)
            except Exception:
                scenarios = []
        if not scenarios:
            scenarios = _local_scenarios_for_card(c)
        c["scenarios"] = scenarios
    return agenda_cards

def _now_str() -> str:
    import time
    return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime())

def _extract_first_json_obj(text: str) -> Dict[str, Any]:
    """
    从模型输出中尽力抽取第一个 JSON 对象。
    允许模型输出前后带解释文字/代码块；只要正文里有一个 JSON 对象即可。
    """
    raw = (text or "").strip()
    if not raw:
        return {}
    # 清理常见 markdown fence
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    start = raw.find("{")
    if start < 0:
        return {}
    depth = 0
    in_str = None
    esc = False
    for i in range(start, len(raw)):
        c = raw[i]
        if esc:
            esc = False
            continue
        if in_str:
            if c == "\\":
                esc = True
                continue
            if c == in_str:
                in_str = None
            continue
        if c in ('"', "'"):
            in_str = c
            continue
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                seg = raw[start : i + 1]
                try:
                    obj = json.loads(seg)
                    return obj if isinstance(obj, dict) else {}
                except Exception:
                    return {}
    return {}

def _llm_call_json(system_prompt: str, user_prompt: str, timeout: int = 60, max_tokens: int = 1400) -> Dict[str, Any]:
    """
    调用 LLM 并要求输出 JSON。失败则抛异常，由上层决定是否降级到模板会议。
    """
    from utils.llm_client import call_llm, load_llm_config

    config_path = _pick_llm_config_path()
    cfg = load_llm_config(config_path)
    if not (cfg.get("api_key") or "").strip():
        raise ValueError("LLM API Key 未配置")
    text = call_llm(
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        config_path=config_path,
        timeout=int(timeout or 60),
        max_tokens=max_tokens,
    )
    obj = _extract_first_json_obj(str(text or ""))
    if not obj:
        raise ValueError("LLM 未输出可解析 JSON")
    return obj

def _llm_build_rounds_for_card(card: Dict[str, Any]) -> Dict[str, Any]:
    """
    为单个议题卡生成 QA→Dev→PM 三轮结构化输出（JSON）。
    """
    topic = str(card.get("topic") or "generic")
    meeting_mode = str(card.get("meeting_mode") or "snapshot").strip().lower()
    evidence = card.get("evidence") or []
    ev_text = "\n".join([f"- {e.get('type')} {e.get('ref')}: {e.get('text') or ''}".strip() for e in evidence if isinstance(e, dict)]) or "- 无"

    common_rules = (
        "你在虚拟评审会中扮演指定角色。请严格输出一个 JSON 对象，禁止输出任何额外文字。"
        "必须用中文。不要编造 PRD 原文；若缺信息请列为 question/open_questions。"
    )
    if meeting_mode == "brainstorm":
        common_rules += (
            "【无文档沙盘】当前没有可引用的 PRD 原文：禁止捏造锚点/行号/章节名。"
            "所有结论必须可被证伪：在 PM 输出中必须把 if_blocked 写满（缺失证据/待确认口径）；"
            "ac_final 只能是“待文档固化”的草案条款，且必须标注“暂定”。"
        )

    qa_schema = (
        "{\n"
        '  "qa_challenges": [\n'
        '    {"gap": "...", "trigger": "...", "user_visible": "...", "system_visible": "...", "worst_case": "...", "question_for_pm": "..."}\n'
        "  ]\n"
        "}"
    )
    dev_schema = (
        "{\n"
        '  "dev_responses": [\n'
        '    {"to_gap": "...", "feasibility": "can_do|need_constraint|reject", "proposal": "...", "cost_boundary": "...", "residual_risk": "...", "needs_pm_decision": ["..."]}\n'
        "  ]\n"
        "}"
    )
    pm_schema = (
        "{\n"
        '  "pm_decision": {\n'
        '    "option": "A|B|C",\n'
        '    "decision_summary": "...",\n'
        '    "ac_final": ["Given... When... Then...(含量化项)"],\n'
        '    "acceptance_scope": "..." \n'
        "  },\n"
        '  "if_blocked": ["若无法裁决，列出缺失信息"]\n'
        "}"
    )

    base_context = (
        f"议题ID：{card.get('issue_id')}\n"
        f"风险等级：{card.get('risk_level')}\n"
        f"Topic：{topic}\n"
        f"一句话问题：{card.get('one_liner')}\n"
        f"证据（锚点/引用/推断）：\n{ev_text}\n"
        f"需要拍板：{card.get('decision_needed')}\n"
    )

    qa_sys = f"{common_rules}\n你的角色：测试（QA），目标是把问题变成可验收条款（阈值/终态/观测字段），并指出不可验收点。"
    qa_user = base_context + f"\n请输出 QA 质疑清单 JSON，结构如下：\n{qa_schema}"
    qa = _llm_call_json(qa_sys, qa_user, timeout=60, max_tokens=1400)

    # 将 QA gap 列表串给 Dev
    gaps = qa.get("qa_challenges") if isinstance(qa.get("qa_challenges"), list) else []
    gap_text = "\n".join([f"- {g.get('gap')}" for g in gaps if isinstance(g, dict)]) or "- （无）"

    dev_sys = f"{common_rules}\n你的角色：研发（Dev），目标是给出最小可实现方案、成本边界，并指出必须由 PM 裁决的点。"
    dev_user = base_context + f"\nQA质疑点：\n{gap_text}\n\n请输出 Dev 回应 JSON，结构如下：\n{dev_schema}"
    dev = _llm_call_json(dev_sys, dev_user, timeout=70, max_tokens=1600)

    needs = []
    for r in (dev.get("dev_responses") or []):
        if isinstance(r, dict):
            for x in (r.get("needs_pm_decision") or []):
                xs = str(x or "").strip()
                if xs and xs not in needs:
                    needs.append(xs)
    needs_text = "\n".join([f"- {x}" for x in needs]) or "- （无）"

    pm_sys = f"{common_rules}\n你的角色：产品（PM），必须做裁决并把 AC 定稿（Given-When-Then，含量化阈值）。不允许“视情况而定”。"
    if meeting_mode == "brainstorm":
        pm_sys += "【无文档】你不得宣布“已决定可上线”；必须把 if_blocked 列为阻塞项（缺 PRD 证据），ac_final 仅作草案。"
    pm_user = base_context + f"\nDev 需要 PM 决策点：\n{needs_text}\n\n请输出 PM 裁决 JSON，结构如下：\n{pm_schema}"
    pm = _llm_call_json(pm_sys, pm_user, timeout=70, max_tokens=1700)

    return {"qa": qa, "dev": dev, "pm": pm}

def _decision_log_from_llm(card: Dict[str, Any], rounds: Dict[str, Any]) -> Dict[str, Any]:
    meeting_mode = str(card.get("meeting_mode") or "snapshot").strip().lower()
    pm = rounds.get("pm") if isinstance(rounds.get("pm"), dict) else {}
    pm_dec = pm.get("pm_decision") if isinstance(pm.get("pm_decision"), dict) else {}
    blocked = pm.get("if_blocked") if isinstance(pm.get("if_blocked"), list) else []
    ac_final = pm_dec.get("ac_final") if isinstance(pm_dec.get("ac_final"), list) else []
    status = "blocked" if blocked or not ac_final else "decided"
    summary = str(pm_dec.get("decision_summary") or "").strip()
    blocked_by = [str(x) for x in blocked if str(x or "").strip()]
    entry: Dict[str, Any] = {
        "issue_id": card.get("issue_id"),
        "status": status,
        "decision": {
            "option": pm_dec.get("option") or ("A" if status == "decided" else ""),
            "summary": summary,
            "rationale": "由多角色评审会裁决生成",
        },
        "ac_final": [str(x) for x in ac_final if str(x or "").strip()][:6],
        "observability_min": {
            "logs": ["session_id", "request_id", "state", "error_code"],
            "metrics": [],
            "traces": [],
        },
        "owners": {"pm": "PM", "dev": "Dev", "qa": "QA"},
        "follow_ups": [
            {"type": "prd_update", "content": "将裁决口径写入 PRD 附录/补丁包并发布。"},
            {"type": "test_case", "content": "按 ac_final 与观测字段补齐验收用例。"},
        ],
        "blocked_by": blocked_by,
    }
    if meeting_mode == "brainstorm":
        entry["status"] = "blocked"
        entry["evidence_class"] = "inference"
        hard = [
            "无 PRD 原文锚点：无法在文档层关闭 P0",
            "推断草案不得作为生产就绪的唯一验收依据",
        ]
        merged = list(blocked_by)
        for h in hard:
            if h not in merged:
                merged.append(h)
        entry["blocked_by"] = merged
        if summary and not summary.startswith("（推断草案"):
            entry["decision"]["summary"] = "（推断草案，待文档固化）" + summary
        elif not summary:
            entry["decision"]["summary"] = "（推断草案，待文档固化）待 PM 在 PRD 中落地阈值/终态/观测字段。"
        entry["decision"]["option"] = ""
        fu = [{"type": "evidence", "content": "补齐 PRD 章节/锚点/量化口径后重新开会或跑静态审计关闭 P0。"}]
        fu.extend(entry.get("follow_ups") or [])
        entry["follow_ups"] = fu
    return entry

def _topic_from_merged_issue(m: Dict[str, Any]) -> str:
    t = " ".join([str(m.get("name") or ""), str(m.get("description") or ""), " ".join(m.get("types") or [])])
    if re.search(r"(矛盾|冲突|互斥|抢占|优先级|总开关)", t):
        return "conflict"
    if re.search(r"(越权|鉴权|权限|未授权|二维码|扫码|外链|受控|凭证)", t):
        return "security"
    if re.search(r"(退出|中断|切后台|重进|回滚|恢复|清理|状态)", t):
        return "state"
    if re.search(r"(失败|超时|弱网|断网|降级|重试|兜底|错误码|不可用)", t):
        return "exception"
    if re.search(r"(多端|跨端|同步|对账|列表一致)", t):
        return "cross_end"
    if re.search(r"(字段|返回|接口|错误码|枚举)", t):
        return "data_contract"
    return "generic"

def _build_agenda_cards_from_snapshot(snapshot: Dict[str, Any], limit: int = 3) -> List[Dict[str, Any]]:
    """
    从学习快照中重建 Stage3（不依赖 L3 文本），然后按 P0 优先 + topic 去重挑 3 条议题卡。
    """
    stage1 = snapshot.get("stage1_output") if isinstance(snapshot.get("stage1_output"), dict) else {}
    stage2 = snapshot.get("stage2_output") if isinstance(snapshot.get("stage2_output"), dict) else {}
    offline_mode = bool(snapshot.get("offline_mode"))
    s3 = pipeline._build_stage3_report(stage1, stage2, offline_mode=offline_mode)
    merged = s3.get("merged_issues") if isinstance(s3.get("merged_issues"), list) else []
    p0 = [m for m in merged if isinstance(m, dict) and str(m.get("risk_level") or "").upper() == "P0"]
    # topic 去重：保证 3 条不是同一类
    picked: List[Dict[str, Any]] = []
    used_topics = set()
    for m in p0:
        topic = _topic_from_merged_issue(m)
        if topic in used_topics:
            continue
        used_topics.add(topic)
        picked.append(m)
        if len(picked) >= max(1, int(limit or 1)):
            break
    if len(picked) < max(1, int(limit or 1)):
        for m in p0:
            if m not in picked:
                picked.append(m)
            if len(picked) >= max(1, int(limit or 1)):
                break

    cards: List[Dict[str, Any]] = []
    for idx, m in enumerate(picked, start=1):
        topic = _topic_from_merged_issue(m)
        anchors = m.get("anchors") or []
        evidence = []
        if anchors:
            evidence.append({"type": "anchor", "ref": str(anchors[0]), "text": ""})
        else:
            evidence.append({"type": "inference", "ref": "", "text": "缺少可引用锚点，本条来自合并问题推断"})
        one_liner = str(m.get("description") or m.get("name") or f"P0议题{idx}").strip()
        if len(one_liner) > 120:
            one_liner = one_liner[:118] + "…"
        cards.append(
            {
                "session_id": str(snapshot.get("snapshot_id") or ""),
                "meeting_mode": "snapshot",
                "issue_id": f"{str(snapshot.get('snapshot_id') or 'snap')[:16]}-P0-{idx:02d}",
                "topic": topic,
                "risk_level": "P0",
                "one_liner": one_liner,
                "evidence": evidence,
                "decision_needed": "为保证可测与可复盘：需要拍板阈值/优先级/终态落点/错误码与最小可观测字段。",
                "candidate_ac": [
                    "Given 触发关键动作 When 达到阈值/条件 Then 系统进入明确终态并提示用户下一步（可量化）。",
                ],
                "constraints": [],
                "open_questions": [],
                "source": {"merged_issue": m},
            }
        )
    return cards

def _meeting_system_message(mode: str) -> str:
    if str(mode or "").strip().lower() == "brainstorm":
        return (
            "会议已启动：无 PRD「无文档沙盘」模式。证据为推断（inference），"
            "Decision Log 默认 BLOCKED，直至补齐文档锚点/量化口径。本页对话仅展示过程，最终以裁决表为准。"
        )
    return "会议已启动：输入来自静态审计快照（P0议题）。本页对话仅展示过程，最终以裁决表为准。"

def _seed_meeting_messages(agenda_cards: List[Dict[str, Any]], mode: str = "snapshot") -> List[Dict[str, Any]]:
    """模板版三轮发言（不含系统行，由上层统一插入一条 system）。"""
    msgs: List[Dict[str, Any]] = []
    is_brain = str(mode or "").strip().lower() == "brainstorm"
    for c in agenda_cards:
        if is_brain:
            msgs.append(
                {
                    "role": "assistant",
                    "speaker": "测试（QA）",
                    "ts": _now_str(),
                    "text": f"议题：{c.get('one_liner')}\n无原文：我只记录可测缺口与风险假设；任何数字阈值都标「暂定」，不当作已发布需求。",
                }
            )
            msgs.append(
                {
                    "role": "assistant",
                    "speaker": "研发（Dev）",
                    "ts": _now_str(),
                    "text": "可做技术预案（状态机/重试/超时），但无 PRD 锚点无法承诺接口契约；需 PM 在文档里固化后再估排期。",
                }
            )
            msgs.append(
                {
                    "role": "assistant",
                    "speaker": "产品（PM）",
                    "ts": _now_str(),
                    "text": "接受「先草案后文档」：本议题在裁决表里保持 BLOCKED，补齐章节后再开一次会关闭。",
                }
            )
        else:
            msgs.append({"role": "assistant", "speaker": "测试（QA）", "ts": _now_str(), "text": f"议题：{c.get('one_liner')}\n我先卡三件事：阈值/终态/观测字段。缺任何一个都不可验收。"})
            msgs.append({"role": "assistant", "speaker": "研发（Dev）", "ts": _now_str(), "text": "可做最小方案：补状态枚举与错误码，增加幂等/防抖；成本可控。需要 PM 拍板阈值与提示强度。"})
            msgs.append({"role": "assistant", "speaker": "产品（PM）", "ts": _now_str(), "text": "同意以“可验收”为硬门槛：给出阈值与最小字段集；失败不允许静默。具体数值可在裁决表里落地。"})
    return msgs

def _build_decision_log_for_cards(agenda_cards: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    out = []
    for c in agenda_cards:
        topic = c.get("topic") or "generic"
        if str(c.get("meeting_mode") or "").strip().lower() == "brainstorm":
            out.append(
                {
                    "issue_id": c.get("issue_id"),
                    "status": "blocked",
                    "evidence_class": "inference",
                    "decision": {
                        "option": "",
                        "summary": "（推断草案，待文档固化）无 PRD 原文：仅记录讨论方向，不关闭 P0。",
                        "rationale": "无文档沙盘：模板会议无法替代文档证据。",
                    },
                    "ac_final": [
                        str(x) for x in (c.get("candidate_ac") or ["Given... When... Then...(含量化项，数字为暂定)"]) if str(x or "").strip()
                    ][:4],
                    "observability_min": {"logs": ["session_id", "request_id", "state", "error_code"], "metrics": [], "traces": []},
                    "owners": {"pm": "PM", "dev": "Dev", "qa": "QA"},
                    "follow_ups": [
                        {"type": "evidence", "content": "补齐 PRD 章节/锚点/量化口径后重新开会或跑静态审计关闭 P0。"},
                        {"type": "prd_update", "content": "将草案口径写入 PRD 附录并标注“已确认/已废弃”。"},
                    ],
                    "blocked_by": [
                        "无 PRD 原文锚点：无法在文档层关闭 P0",
                        "推断草案不得作为生产就绪的唯一验收依据",
                    ],
                }
            )
            continue
        # 给一个“可落地”的默认门槛，后续可在 UI 中编辑或二次会议调整
        if topic == "exception":
            ac = ["T=1.0s 超时判失败；N=2 次重试；Δ=2s；失败必须提示且记录 error_code。"]
            obs = {"logs": ["session_id", "request_id", "business_id", "state", "error_code", "retry_count", "start_latency_ms"], "metrics": ["success_rate", "p95_latency_ms"], "traces": []}
        elif topic == "state":
            ac = ["W=2.0s 内收敛到 READY/RESUME_ALLOWED/FAILED；清理清单逐项可对账。"]
            obs = {"logs": ["session_id", "exit_reason", "cleanup_items_applied", "final_state", "state_updated_at"], "metrics": ["state_converge_p95_ms"], "traces": []}
        elif topic == "conflict":
            ac = ["总开关最高否决权；任务启动锁定 config_snapshot；冲突必须提示且各入口一致。"]
            obs = {"logs": ["config_snapshot_id", "global_switch", "session_switch", "effective_switch", "conflict_reason"], "metrics": ["conflict_rate"], "traces": []}
        elif topic == "security":
            ac = ["外显/扫码/外链访问必须鉴权；过期/无资源/越权有明确提示与审计日志。"]
            obs = {"logs": ["subject_id", "actor", "resource_id", "auth_result", "error_code", "audit_log_id"], "metrics": ["unauthorized_block_rate"], "traces": []}
        else:
            ac = ["给出可验收口径：阈值/终态/观测字段齐全。"]
            obs = {"logs": ["session_id", "state", "error_code"], "metrics": [], "traces": []}

        out.append(
            {
                "issue_id": c.get("issue_id"),
                "status": "decided",
                "decision": {
                    "option": "A",
                    "summary": "以测试可验收为硬门槛：阈值+终态+观测字段齐全，失败不静默。",
                    "rationale": "没有阈值/终态/观测字段，线上无法复盘，测试无法对结果负责。",
                },
                "ac_final": ac,
                "observability_min": obs,
                "owners": {"pm": "PM", "dev": "Dev", "qa": "QA"},
                "follow_ups": [
                    {"type": "prd_update", "content": "补齐裁决口径（阈值、终态、错误码/字段）并作为 PRD 附录发布。"},
                    {"type": "test_case", "content": "按 error_code / 退出原因 / 冲突矩阵分桶补齐用例并回归。"},
                ],
                "blocked_by": [],
            }
        )
    return out

def _build_prd_patch(decision_log: List[Dict[str, Any]]) -> str:
    lines = ["## 附录：审计裁决版口径（自动生成）", ""]
    for item in decision_log or []:
        if not isinstance(item, dict):
            continue
        issue_id = str(item.get("issue_id") or "")
        lines.append(f"### {issue_id}")
        st = str(item.get("status") or "").strip().lower()
        if st == "blocked":
            lines.append("- **状态**：BLOCKED（待 PRD 证据后再冻结口径）")
            bb = item.get("blocked_by") if isinstance(item.get("blocked_by"), list) else []
            for b in bb[:8]:
                bs = str(b or "").strip()
                if bs:
                    lines.append(f"- **阻塞**：{bs}")
        evc = str(item.get("evidence_class") or "").strip().lower()
        if evc == "inference":
            lines.append("- **证据等级**：inference（推断草案，不得单独作为上线验收依据）")
        for ac in item.get("ac_final") or []:
            lines.append(f"- {str(ac)}")
        obs = item.get("observability_min") if isinstance(item.get("observability_min"), dict) else {}
        logs = obs.get("logs") if isinstance(obs.get("logs"), list) else []
        if logs:
            lines.append(f"- 最小观测字段：{', '.join(str(x) for x in logs[:24])}")
        lines.append("")
    return "\n".join(lines).strip()

DEFAULT_BUG_PATTERNS = [
    {"pattern_id": "P001", "keywords": ["黑屏", "无画面", "无视频"], "category": "媒体异常", "design_gap": ["异常处理缺失", "资源释放缺失"], "rule": "媒体流程必须定义失败兜底策略", "weight": 0.9, "enabled": True},
    {"pattern_id": "P002", "keywords": ["卡死", "无响应", "假死"], "category": "稳定性", "design_gap": ["超时机制缺失", "状态回收缺失"], "rule": "关键交互必须定义超时与恢复策略", "weight": 0.88, "enabled": True},
    {"pattern_id": "P003", "keywords": ["崩溃", "闪退", "重启"], "category": "稳定性", "design_gap": ["异常隔离缺失"], "rule": "核心模块必须定义异常隔离与重试边界", "weight": 0.95, "enabled": True},
    {"pattern_id": "P004", "keywords": ["不同步", "状态错乱", "状态不一致"], "category": "状态管理", "design_gap": ["状态机缺失"], "rule": "多状态系统必须定义状态机与转移条件", "weight": 0.92, "enabled": True},
    {"pattern_id": "P005", "keywords": ["越权", "权限", "未授权"], "category": "权限安全", "design_gap": ["权限规则缺失"], "rule": "高风险操作必须定义角色鉴权和越权拦截", "weight": 0.9, "enabled": True},
    {"pattern_id": "P006", "keywords": ["并发", "抢占", "竞态"], "category": "并发冲突", "design_gap": ["并发裁决缺失"], "rule": "并发场景必须定义裁决优先级与幂等策略", "weight": 0.86, "enabled": True},
    {"pattern_id": "P007", "keywords": ["恢复失败", "无法恢复", "回退失败"], "category": "恢复机制", "design_gap": ["回退策略缺失"], "rule": "中断流程必须定义恢复或回退策略", "weight": 0.87, "enabled": True},
    {"pattern_id": "P008", "keywords": ["主板", "芯片", "平台差异"], "category": "平台兼容", "design_gap": ["平台差异约束缺失"], "rule": "涉及硬件能力的功能必须定义平台差异策略", "weight": 0.82, "enabled": True},
]

def _pick_llm_config_path() -> str:
    """
    返回平台级单一配置路径（config/llm_config.json）。
    所有模块（含 PRD 审计）应共用这一份配置。
    """
    return PLATFORM_LLM_CONFIG_FILE

def _normalize_knowledge_items(items):
    if not isinstance(items, list):
        return []
    out = []
    for it in items:
        if not isinstance(it, dict):
            continue
        capability_id = str(it.get("capability_id") or "").strip()
        name = str(it.get("name") or "").strip()
        if not capability_id or not name:
            continue
        out.append(it)
    return out

import threading
import tempfile
import time

_FILE_LOCK = threading.Lock()

def _load_json_file(path, default_value):
    if os.path.exists(path):
        try:
            with _FILE_LOCK:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception as e:
            logger.warning(f"读取 JSON 文件失败 ({path}): {e}")
            return default_value
    return default_value

def _save_json_file(path, data):
    """
    带锁且原子性的保存 JSON 文件，防止多线程/多进程写入冲突及文件损坏。
    """
    with _FILE_LOCK:
        try:
            # 使用临时文件写入，然后原子替换，防止写入中断导致文件损坏
            fd, temp_path = tempfile.mkstemp(dir=os.path.dirname(path), suffix=".tmp")
            try:
                with os.fdopen(fd, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                
                # Windows 下 os.replace 是原子的（前提是同一驱动器）
                os.replace(temp_path, path)
            except Exception as e:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                raise e
        except Exception as e:
            logger.error(f"保存 JSON 文件失败 ({path}): {e}")
            raise e

def _djb2_hash(s: str) -> str:
    h = 5381
    for ch in s:
        h = ((h << 5) + h) + ord(ch)
        h &= 0xFFFFFFFF
    return format(h, "08x")

def _defect_key(d: Dict[str, Any]) -> str:
    if not isinstance(d, dict):
        return "d00000000"
    base = "|".join([
        str(d.get("type") or "").strip(),
        str(d.get("module") or "").strip(),
        str(d.get("description") or "").strip(),
        str(d.get("anchor") or "").strip(),
    ])
    return "d" + _djb2_hash(base)

def _load_watch_state() -> Dict[str, Any]:
    data = _load_json_file(FEISHU_WATCH_FILE, {"items": []})
    if not isinstance(data, dict):
        data = {"items": []}
    if not isinstance(data.get("items"), list):
        data["items"] = []
    return data

def _save_watch_state(data: Dict[str, Any]) -> None:
    if not isinstance(data, dict):
        data = {"items": []}
    if not isinstance(data.get("items"), list):
        data["items"] = []
    os.makedirs(os.path.dirname(FEISHU_WATCH_FILE), exist_ok=True)
    _save_json_file(FEISHU_WATCH_FILE, data)

from .job_queue import JobQueue

_job_queue = JobQueue(STORAGE_DIR)

def _match_watch_item(doc_url: str) -> str:
    try:
        from .feishu_client import extract_document_id_from_url
        parsed = extract_document_id_from_url(doc_url or "")
        if not parsed:
            return ""
        doc_id, doc_type = parsed
        return f"{doc_type}:{doc_id}"
    except Exception:
        return ""

def _feishu_doc_audit_handler(payload: Dict[str, Any]) -> Dict[str, Any]:
    from .feishu_client import fetch_feishu_doc_content
    import hashlib

    doc_url = str(payload.get("doc_url") or "").strip()
    name = str(payload.get("name") or "").strip()
    use_llm = bool(payload.get("use_llm", True))
    timeout = int(payload.get("timeout") or 120)

    ok, prd_text = fetch_feishu_doc_content(doc_url)
    if not ok:
        raise ValueError(prd_text or "飞书文档拉取失败")

    h = hashlib.sha1((prd_text or "").encode("utf-8")).hexdigest()

    state = _load_watch_state()
    items = state.get("items") if isinstance(state.get("items"), list) else []
    key = _match_watch_item(doc_url)
    target = None
    for it in items:
        if not isinstance(it, dict):
            continue
        if (it.get("key") or "") == key:
            target = it
            break
    if not target:
        target = {"key": key, "name": name or key, "doc_url": doc_url, "enabled": True}
        items.append(target)
        state["items"] = items

    last_hash = str(target.get("last_hash") or "")
    if last_hash and last_hash == h and bool(payload.get("skip_if_unchanged", True)):
        target["last_checked_at"] = int(time.time())
        _save_watch_state(state)
        return {"doc_url": doc_url, "unchanged": True}

    config_path = _pick_llm_config_path()
    if not use_llm:
        config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

    report_md, stage1_output, stage2_output, stage3_output = pipeline.run_prd_audit_sync(
        prd_text, llm_config_path=config_path, timeout=timeout
    )

    snapshot_id = ""
    if isinstance(stage3_output, dict):
        snapshot_id = str(stage3_output.get("snapshot_id") or "")

    target["name"] = name or target.get("name") or key
    target["doc_url"] = doc_url
    target["key"] = key
    target["last_hash"] = h
    target["last_audit_at"] = int(time.time())
    target["last_snapshot_id"] = snapshot_id
    target["last_checked_at"] = int(time.time())
    _save_watch_state(state)

    return {
        "doc_url": doc_url,
        "name": target.get("name"),
        "snapshot_id": snapshot_id,
        "score": (stage3_output.get("summary", {}).get("score") if isinstance(stage3_output, dict) else None),
    }

_job_queue.register_handler("feishu_doc_audit", _feishu_doc_audit_handler)

def _guardrail_regression_handler(payload: Dict[str, Any]) -> Dict[str, Any]:
    from .audit_learning import load_adversarial_cases, append_regression_run
    from .guardrail_engine import evaluate_guardrail

    req = payload if isinstance(payload, dict) else {}
    timeout = int(req.get("timeout") or 120)
    default_use_llm = bool(req.get("use_llm", True))
    only_ids = req.get("case_ids")
    only_ids = only_ids if isinstance(only_ids, list) else []

    cases_obj = load_adversarial_cases()
    items = cases_obj.get("items") if isinstance(cases_obj, dict) else []
    items = items if isinstance(items, list) else []

    picked = []
    for it in items:
        if not isinstance(it, dict):
            continue
        cid = str(it.get("case_id") or "").strip()
        if not cid:
            continue
        if only_ids and cid not in only_ids:
            continue
        picked.append(it)

    picked = picked[:30]
    if not picked:
        run = {"ts": int(time.time()), "total": 0, "passed": 0, "cases": []}
        append_regression_run(run)
        return run

    total = 0
    passed = 0
    attack_total = 0
    blocked = 0
    breach = 0
    normal_total = 0
    false_positive = 0
    results = []

    for it in picked:
        cid = str(it.get("case_id") or "").strip()
        name = str(it.get("name") or cid).strip()
        kind = str(it.get("kind") or "normal").strip()
        prd_text = str(it.get("prd_text") or "").strip()
        min_score = it.get("min_guardrail_score")
        try:
            min_score = int(float(min_score)) if min_score is not None else 70
        except Exception:
            min_score = 70
        use_llm = it.get("use_llm")
        use_llm = default_use_llm if use_llm is None else bool(use_llm)
        config_path = _pick_llm_config_path()
        if not use_llm:
            config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

        ok = False
        gr = {}
        err = ""
        try:
            report_md, stage1_output, stage2_output, stage3_output = pipeline.run_prd_audit_sync(
                prd_text, llm_config_path=config_path, timeout=timeout
            )
            gr = evaluate_guardrail(
                prd_text=prd_text,
                stage1_output=stage1_output,
                stage2_output=stage2_output,
                report_md=report_md,
                test_cases=[],
            )
            score = int(float(gr.get("score") or 0))
            fatal = any(
                isinstance(c, dict)
                and str(c.get("severity") or "").upper() == "P0"
                and not bool(c.get("ok"))
                for c in (gr.get("checks") or [])
            )
            ok = (score >= min_score) and (not fatal)
        except Exception as e:
            err = str(e)
            ok = False

        total += 1
        passed += 1 if ok else 0

        checks = gr.get("checks") if isinstance(gr, dict) else []
        checks = checks if isinstance(checks, list) else []
        has_inj_signal = any(isinstance(c, dict) and c.get("id") == "prompt_injection_signal" and not bool(c.get("ok")) for c in checks)
        inj_followed = any(isinstance(c, dict) and c.get("id") == "prompt_injection_followed" and not bool(c.get("ok")) for c in checks)

        if kind in ("prompt_injection", "format_attack"):
            attack_total += 1
            if inj_followed:
                breach += 1
            else:
                blocked += 1 if ok else 0
        else:
            normal_total += 1
            if has_inj_signal and not inj_followed:
                false_positive += 1 if (not ok) else 0

        results.append({
            "case_id": cid,
            "name": name,
            "kind": kind,
            "ok": ok,
            "min_guardrail_score": min_score,
            "guardrail_score": int(float(gr.get("score") or 0)) if isinstance(gr, dict) else 0,
            "breach": bool(inj_followed),
            "error": err,
        })

    run = {
        "ts": int(time.time()),
        "total": total,
        "passed": passed,
        "pass_rate": 0 if total <= 0 else round(float(passed) / float(total), 4),
        "attack_total": attack_total,
        "blocked": blocked,
        "breach": breach,
        "block_rate": 0 if attack_total <= 0 else round(float(blocked) / float(attack_total), 4),
        "normal_total": normal_total,
        "false_positive": false_positive,
        "false_positive_rate": 0 if normal_total <= 0 else round(float(false_positive) / float(normal_total), 4),
        "cases": results,
    }
    append_regression_run(run)
    return run

_job_queue.register_handler("guardrail_regression", _guardrail_regression_handler)

def _ensure_bug_assets():
    if not os.path.exists(BUG_PATTERN_FILE):
        _save_json_file(BUG_PATTERN_FILE, {"items": DEFAULT_BUG_PATTERNS})
    if not os.path.exists(BUG_RAW_FILE):
        _save_json_file(BUG_RAW_FILE, {"items": []})
    if not os.path.exists(BUG_RULE_FILE):
        _save_json_file(BUG_RULE_FILE, {"items": []})
    if not os.path.exists(VECTOR_DATA_FILE):
        _save_json_file(VECTOR_DATA_FILE, {"items": []})

def _normalize_severity(s):
    v = str(s or "").strip().upper()
    if v in ["P0", "P1", "P2"]:
        return v
    if "P0" in v:
        return "P0"
    if "P1" in v:
        return "P1"
    if "P2" in v:
        return "P2"
    raw = str(s or "").strip()
    if any(x in raw for x in ["A-严重", "A严重", "严重", "高危"]):
        return "P0"
    if any(x in raw for x in ["B-适中", "B适中", "中等", "中危"]):
        return "P1"
    if any(x in raw for x in ["C-微小", "C微小", "轻微", "低危"]):
        return "P2"
    return "P2"

def _normalize_frequency(s):
    v = str(s or "").strip()
    return v or "中"

def _tokenize_text(text):
    t = str(text or "").lower()
    parts = re.split(r"[^\w\u4e00-\u9fff]+", t)
    return [p for p in parts if p]

def _text_score(query, content):
    q = set(_tokenize_text(query))
    c = set(_tokenize_text(content))
    if not q or not c:
        return 0.0
    inter = len(q.intersection(c))
    union = len(q.union(c))
    s = inter / max(1, union)
    if str(query or "").strip() and str(query).strip() in str(content or ""):
        s += 0.2
    return round(min(1.0, s), 4)

def _next_rule_id(existing):
    mx = 0
    for it in existing:
        rid = str((it or {}).get("rule_id") or "")
        m = re.match(r"RBUG_(\d+)$", rid)
        if m:
            mx = max(mx, int(m.group(1)))
    return "RBUG_{:03d}".format(mx + 1)

def _load_bug_patterns():
    _ensure_bug_assets()
    data = _load_json_file(BUG_PATTERN_FILE, {"items": []})
    items = data.get("items") if isinstance(data, dict) else []
    return items if isinstance(items, list) else []

def _analyze_bug_desc(desc):
    text = str(desc or "").strip()
    if not text:
        return None
    patterns = _load_bug_patterns()
    hits = []
    for p in patterns:
        if not isinstance(p, dict):
            continue
        if p.get("enabled") is False:
            continue
        kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
        matched = [k for k in kws if str(k) and str(k) in text]
        if matched:
            hits.append((float(p.get("weight") or 0.5), p, matched))
    hits = sorted(hits, key=lambda x: x[0], reverse=True)
    if hits:
        _, top, matched = hits[0]
        return {
            "category": top.get("category") or "通用缺陷",
            "pattern_id": top.get("pattern_id") or "",
            "pattern": "、".join([str(x) for x in matched[:3]]),
            "design_gap": top.get("design_gap") if isinstance(top.get("design_gap"), list) else [],
            "rule": top.get("rule") or "关键流程必须定义异常兜底策略",
            "weight": float(top.get("weight") or 0.5),
        }
    return {
        "category": "通用缺陷",
        "pattern_id": "",
        "pattern": "未命中",
        "design_gap": ["异常处理缺失"],
        "rule": "关键流程必须定义异常兜底策略",
        "weight": 0.4,
    }

def _vector_search_local(query, top_k=5, type_filter=None, board_type=None):
    _ensure_bug_assets()
    data = _load_json_file(VECTOR_DATA_FILE, {"items": []})
    items = data.get("items") if isinstance(data, dict) else []
    if not isinstance(items, list):
        items = []
    out = []
    for it in items:
        if not isinstance(it, dict):
            continue
        if type_filter and str(it.get("type") or "") != str(type_filter):
            continue
        if board_type and str(it.get("board_type") or "") != str(board_type):
            continue
        content = str(it.get("content") or "")
        score = _text_score(query, content)
        if score <= 0:
            continue
        row = dict(it)
        row["score"] = score
        out.append(row)
    out.sort(key=lambda x: float(x.get("score") or 0), reverse=True)
    return out[: max(1, int(top_k or 5))]

def _load_knowledge_cards():
    if os.path.exists(KNOWLEDGE_CARDS_FILE):
        try:
            with open(KNOWLEDGE_CARDS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                items = _normalize_knowledge_items(data.get("items"))
            else:
                items = _normalize_knowledge_items(data)
            if items:
                return items
        except Exception as e:
            logger.warning("读取能力库文件失败: %s", e)
    return _normalize_knowledge_items(KTV_CAPABILITY_CARDS)

def _save_knowledge_cards(items):
    payload = {"items": _normalize_knowledge_items(items)}
    with open(KNOWLEDGE_CARDS_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return payload["items"]

def _match_knowledge_base(prd_text: str, top_k: int = 5):
    """
    本地规则匹配：从PRD文本中匹配知识库，返回相关建议
    不需要大模型，纯关键词匹配
    """
    if not prd_text:
        return []
    
    prd_text_lower = prd_text.lower()
    knowledge_cards = _load_knowledge_cards()
    
    matches = []
    for card in knowledge_cards:
        if not isinstance(card, dict):
            continue
        
        # 提取知识卡片的关键词
        keywords = []
        # 从 name, domain, category, trigger 中提取关键词
        for field in ['name', 'domain', 'category', 'trigger', 'capability_id']:
            value = str(card.get(field) or "").strip()
            if value:
                keywords.append(value)
        
        # 匹配关键词
        matched_keywords = []
        for kw in keywords:
            if kw.lower() in prd_text_lower:
                matched_keywords.append(kw)
        
        # 如果有关键词匹配，计算相关度分数
        if matched_keywords:
            score = len(matched_keywords) / max(len(keywords), 1)
            # 优先级的权重
            priority = card.get('priority', 'P3')
            priority_weight = {'P0': 1.5, 'P1': 1.2, 'P2': 1.0, 'P3': 0.8}.get(priority, 1.0)
            score *= priority_weight
            
            matches.append({
                'capability_id': card.get('capability_id', ''),
                'name': card.get('name', ''),
                'domain': card.get('domain', ''),
                'priority': priority,
                'trigger': card.get('trigger', ''),
                'suggestion': card.get('suggestion_template', ''),
                'matched_keywords': matched_keywords,
                'score': round(score, 2)
            })
    
    # 按分数排序，返回 top_k 个
    matches.sort(key=lambda x: x['score'], reverse=True)
    return matches[:top_k]

def _load_rule_items(path):
    data = _load_json_file(path, {"rules": []})
    if isinstance(data, dict):
        items = data.get("rules")
    elif isinstance(data, list):
        items = data
    else:
        items = []
    return [it for it in items if isinstance(it, dict)]

def _build_platform_center_context():
    knowledge_cards = _load_knowledge_cards()
    knowledge_domains = sorted(
        {
            str((it or {}).get("domain") or "").strip()
            for it in knowledge_cards
            if str((it or {}).get("domain") or "").strip()
        }
    )

    published_rules = _load_rule_items(PRD_RULES_V2_FILE)
    legacy_rules = _load_rule_items(PRD_RULES_V1_FILE)

    _ensure_bug_assets()
    bug_patterns = _load_bug_patterns()
    bug_raw_items = _load_json_file(BUG_RAW_FILE, {"items": []}).get("items", [])
    bug_raw_items = bug_raw_items if isinstance(bug_raw_items, list) else []
    bug_rule_items = _load_json_file(BUG_RULE_FILE, {"items": []}).get("items", [])
    bug_rule_items = bug_rule_items if isinstance(bug_rule_items, list) else []

    learning_status = get_learning_status()
    learning_dashboard = get_learning_quality_dashboard(limit=5000)
    candidate_payload = load_rule_candidates()
    candidate_items = candidate_payload.get("candidates") if isinstance(candidate_payload, dict) else []
    candidate_items = candidate_items if isinstance(candidate_items, list) else []

    applied_rules = _load_rule_items(str(learning_status.get("applied_file") or ""))
    latest_snapshot = learning_status.get("latest_snapshot") if isinstance(learning_status, dict) else {}
    latest_snapshot = latest_snapshot if isinstance(latest_snapshot, dict) else {}
    kpis = learning_dashboard.get("kpis") if isinstance(learning_dashboard, dict) else {}
    kpis = kpis if isinstance(kpis, dict) else {}

    published_enabled = sum(1 for r in published_rules if r.get("enabled") is not False)
    published_categories = len(
        {
            str((r or {}).get("category") or "").strip()
            for r in published_rules
            if str((r or {}).get("category") or "").strip()
        }
    )
    bug_patterns_enabled = sum(1 for p in bug_patterns if p.get("enabled") is not False)
    bug_pattern_categories = len(
        {
            str((p or {}).get("category") or "").strip()
            for p in bug_patterns
            if str((p or {}).get("category") or "").strip()
        }
    )

    groups = [
        {
            "title": "业务认知资产",
            "icon": "fa-database",
            "summary": "沉淀业务能力、角色职责、前置条件、异常与反模式，帮助系统理解“这类需求本来应该长什么样”。",
            "items": [
                {
                    "name": "PRD知识库",
                    "count": len(knowledge_cards),
                    "meta": f"覆盖领域 {len(knowledge_domains)} 个",
                    "desc": "适合录入能力卡片、业务约束、异常处理要求和建议模板。",
                    "href": "/prd_audit/knowledge",
                    "action": "进入知识库",
                }
            ],
        },
        {
            "title": "机器审计资产",
            "icon": "fa-shield-alt",
            "summary": "把组织经验变成可执行的检测规则，决定系统会自动卡什么问题、按什么风险级别提示。",
            "items": [
                {
                    "name": "正式规则库（v2）",
                    "count": len(published_rules),
                    "meta": f"启用 {published_enabled} 条 / 分类 {published_categories} 个",
                    "desc": "当前真正参与 PRD 审计命中的正式规则资产。",
                    "href": "/prd_audit/learning_mvp",
                    "action": "去学习MVP查看发布链路",
                },
                {
                    "name": "映射规则页",
                    "count": len(bug_rule_items),
                    "meta": f"旧版规则 {len(legacy_rules)} 条",
                    "desc": "当前“PRD规则库”页面展示的是可维护的缺陷映射规则，用于补充经验型规则入口。",
                    "href": "/prd_audit/rules",
                    "action": "进入规则页",
                },
            ],
        },
        {
            "title": "缺陷经验资产",
            "icon": "fa-bug",
            "summary": "把历史 Bug 从原始记录抽象为可复用模式，再反向喂给审计系统，避免同类问题重复发生。",
            "items": [
                {
                    "name": "Bug模式库",
                    "count": len(bug_patterns),
                    "meta": f"启用 {bug_patterns_enabled} 条 / 分类 {bug_pattern_categories} 个",
                    "desc": "沉淀“踩过哪些坑、应该防什么坑”。",
                    "href": "/prd_audit/bug_patterns",
                    "action": "进入Bug模式库",
                },
                {
                    "name": "Bug原始库",
                    "count": len(bug_raw_items),
                    "meta": f"映射规则 {len(bug_rule_items)} 条",
                    "desc": "保存原始缺陷样本，供模式提炼、关键词归纳和向量检索使用。",
                    "href": "/prd_audit/bug_patterns",
                    "action": "查看Bug资产",
                },
            ],
        },
        {
            "title": "学习生长资产",
            "icon": "fa-brain",
            "summary": "把每次审计留下的快照、候选规则、待发布规则组织成一条“先学习、再人工审核、再发布”的生长链路。",
            "items": [
                {
                    "name": "历史快照",
                    "count": int(learning_status.get("snapshot_count", 0) or 0),
                    "meta": "来源于每次“开始分析”后的自动落盘",
                    "desc": "是学习 MVP 的数据底座，也是后续回放和复盘的依据。",
                    "href": "/prd_audit/",
                    "action": "回到审计页查看历史记录",
                },
                {
                    "name": "候选规则",
                    "count": len(candidate_items),
                    "meta": f"待发布规则 {len(applied_rules)} 条 / 备份 {int(kpis.get('backup_count', 0) or 0)} 份",
                    "desc": "系统从历史快照中提炼出的候选项，需人工勾选后才能进入正式规则库。",
                    "href": "/prd_audit/learning_mvp",
                    "action": "进入学习MVP",
                },
            ],
        },
    ]

    summary = {
        "knowledge_count": len(knowledge_cards),
        "published_rule_count": len(published_rules),
        "bug_pattern_count": len(bug_patterns),
        "snapshot_count": int(learning_status.get("snapshot_count", 0) or 0),
        "candidate_count": len(candidate_items),
        "applied_count": len(applied_rules),
        "latest_snapshot": latest_snapshot,
    }

    return {
        "summary": summary,
        "groups": groups,
    }

def _build_xmind_tree(nodes):
    root = {
        "id": str(uuid.uuid4()),
        "title": "PRD功能清单",
        "children": {"attached": []},
    }
    stack = [root]
    for raw in nodes:
        title = str((raw or {}).get("title") or "").strip()
        if not title:
            continue
        level = int((raw or {}).get("level") or 1)
        if level < 1:
            level = 1
        if level > 6:
            level = 6
        while len(stack) > level:
            stack.pop()
        parent = stack[-1]
        topic = {"id": str(uuid.uuid4()), "title": title}
        parent.setdefault("children", {}).setdefault("attached", []).append(topic)
        stack.append(topic)
    return root

def _build_xmind_file(nodes):
    tree_root = _build_xmind_tree(nodes)
    content = [
        {
            "id": str(uuid.uuid4()),
            "class": "sheet",
            "title": "PRD功能导图",
            "rootTopic": tree_root,
        }
    ]
    metadata = {"creator": {"name": "Trae PRD Audit"}, "activeSheetId": content[0]["id"]}
    manifest = {"file-entries": {"content.json": {}, "metadata.json": {}}}

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content.json", json.dumps(content, ensure_ascii=False))
        zf.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False))
        zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
    buf.seek(0)
    return buf

@prd_audit_bp.route("/")
def index():
    """独立 PRD 审计页"""
    return render_template("prd_audit_index.html")

@prd_audit_bp.route("/review_meeting")
def review_meeting_page():
    """虚拟评审会控制台（带对话框）"""
    return render_template("prd_audit_review_meeting.html")

@prd_audit_bp.route("/review_meeting_v2")
def review_meeting_v2_page():
    """虚拟评审会 V2：QA 主导竖版工作台（独立于静态审计输出）"""
    return render_template("prd_audit_review_meeting_v2.html")

@prd_audit_bp.route("/api/review_meeting_v2/start", methods=["POST"])
def api_review_meeting_v2_start():
    """
    V2：启动会议（QA 主导）。
    - mode=snapshot: 从 learning_repo/snapshots 读取静态审计快照（只读消费）
    - mode=prd: 传入 prd_text，先跑一次静态审计（不修改静态审计页），再用 stage3_output 生成会议
    """
    try:
        data = request.get_json() or {}
        mode = str(data.get("mode") or "snapshot").strip().lower()
        enforce_gate = bool(data.get("enforce_gate", True))
        # V2 默认智能优先：尽量使用 LLM 讨论；失败则可降级，但会明牌提示。
        use_llm = bool(data.get("use_llm", True))
        # 若为 true：LLM 失败则直接终止/报错，不允许兜底模板。
        # 默认开启：确保“人话摘要”只在 LLM 真成功时出现，避免静默降级造成误判。
        llm_required = bool(data.get("llm_required", True))
        # 静态审计约束：默认开启（强绑定锚点/缺口/required_evidence，不跑偏）。
        audit_constraint = bool(data.get("audit_constraint", True))
        snapshot_id = str(data.get("snapshot_id") or "").strip()
        prd_text = str(data.get("prd_text") or "").strip()

        stage3_output = {}
        snap_id_for_meeting = ""
        llm_config_path = ""
        llm_config_error = ""

        if mode == "prd":
            if not prd_text:
                return error_response("prd_text 不能为空", status_code=400)
            # Run audit sync to get Stage3 output. This does not alter static audit page behavior.
            from utils.llm_client import load_llm_config

            config_path = _pick_llm_config_path()
            try:
                llm_config = load_llm_config(config_path)
            except FileNotFoundError:
                return error_response("API Key 未配置", status_code=400)
            if not (llm_config.get("api_key") or "").strip():
                return error_response("API Key 未配置", status_code=400)
            _, _, _, stage3_output = pipeline.run_prd_audit_sync(prd_text, llm_config_path=config_path, timeout=90)
            llm_config_path = config_path
            # meeting snapshot id (not static snapshot)
            snap_id_for_meeting = "prd_input_" + uuid.uuid4().hex[:8]
        else:
            from .audit_learning import SNAPSHOT_DIR, load_all_snapshots

            snapshot = None
            if snapshot_id:
                fp = os.path.join(SNAPSHOT_DIR, f"{snapshot_id}.json")
                if not os.path.exists(fp):
                    return error_response("snapshot_id 不存在", status_code=404)
                with open(fp, "r", encoding="utf-8") as f:
                    snapshot = json.load(f)
            else:
                snaps = load_all_snapshots(limit=1)
                snapshot = snaps[0] if snaps else None
            if not isinstance(snapshot, dict):
                return error_response("未找到可用快照，请先执行一次 PRD 审计生成快照", status_code=400)
            snap_id_for_meeting = str(snapshot.get("snapshot_id") or snapshot_id or "")
            stage3_output = snapshot.get("stage3_output") if isinstance(snapshot.get("stage3_output"), dict) else {}

            # Optional: enforce gate by P0 in stage2 defects (same as V1 semantics), but only blocks meeting start.
            if enforce_gate:
                stage2 = snapshot.get("stage2_output") if isinstance(snapshot.get("stage2_output"), dict) else {}
                defects = stage2.get("defects") if isinstance(stage2.get("defects"), list) else []
                p0_cnt = sum(1 for d in defects if isinstance(d, dict) and str(d.get("risk_level") or "").upper() == "P0")
                if p0_cnt > 0:
                    return error_response(
                        "门禁未通过：存在 P0 风险缺陷，建议先修正 PRD 或关闭 enforce_gate 用于复盘讨论。",
                        status_code=412,
                        data={"p0_defects_count": p0_cnt, "snapshot_id": snap_id_for_meeting},
                    )
            # LLM config is optional for snapshot mode; if missing we will fall back to template discussion.
            try:
                from utils.llm_client import load_llm_config

                config_path = _pick_llm_config_path()
                llm_config = load_llm_config(config_path)
                if (llm_config.get("api_key") or "").strip():
                    llm_config_path = config_path
                else:
                    llm_config_error = "LLM 配置存在但 api_key 为空，已降级为模板讨论"
            except Exception as e:
                llm_config_path = ""
                llm_config_error = f"LLM 配置读取失败，已降级为模板讨论：{e}"

        if not isinstance(stage3_output, dict) or not stage3_output:
            return error_response("缺少 stage3_output，无法生成会议", status_code=400)

        # Pass PRD content to V2 so it can extract quotes by anchor (strengthen PRD binding in discussion/UI).
        prd_for_v2 = ""
        if mode == "prd":
            prd_for_v2 = prd_text
        else:
            try:
                prd_for_v2 = str(snapshot.get("stage2_output", {}).get("prd_content") or "") if isinstance(snapshot, dict) else ""
            except Exception:
                prd_for_v2 = ""
        meeting = rmv2.create_meeting_from_stage3(stage3_output, snapshot_id=snap_id_for_meeting, mode=mode, prd_content=prd_for_v2)
        llm_error = ""
        if use_llm and llm_config_path:
            meeting, llm_error = rmv2.run_llm_discussion_for_meeting(
                meeting=meeting,
                llm_config_path=llm_config_path,
                max_rounds=2,
                audit_constraint=audit_constraint,
            )
        # 明牌降级：用户选择 use_llm 但无法加载配置时，也要把原因回传到前端。
        if use_llm and (not llm_config_path) and (not llm_error) and llm_config_error:
            llm_error = llm_config_error
        if llm_required and (not use_llm or not llm_config_path or llm_error):
            return error_response(
                "LLM 必须参与，但当前未成功调用 LLM：" + (llm_error or llm_config_error or "未知原因"),
                status_code=412,
                data={"llm_error": llm_error or llm_config_error},
            )
        return success_response(
            data={
                "meeting_id": meeting.meeting_id,
                "meeting": meeting.to_dict(),
                "use_llm": bool(use_llm and bool(llm_config_path) and not bool(llm_error)),
                "llm_error": llm_error,
                "llm_required": llm_required,
                "audit_constraint": audit_constraint,
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_v2_start failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/export_markdown", methods=["POST"])
def api_review_meeting_v2_export_markdown():
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        m = rmv2.get_meeting(mid)
        if not m:
            return error_response("meeting_id 不存在或已过期", status_code=404)
        md = rmv2.export_meeting_markdown(m.to_dict())
        return success_response(data={"meeting_id": mid, "markdown": md})
    except Exception as e:
        logger.exception("api_review_meeting_v2_export_markdown failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/export_prd_v2", methods=["POST"])
def api_review_meeting_v2_export_prd_v2():
    """
    V2：导出“重构后的 PRD v2.0（Executable Version）”，用于一键交付可开工规范。
    """
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        # Default to LLM rewrite export to produce a more "human-like" PRD.
        # Callers can explicitly turn it off via use_rewrite_engine=false.
        use_rewrite_engine = bool(data.get("use_rewrite_engine", True))
        export_kind = str(data.get("export_kind") or "draft").strip().lower()
        if export_kind not in ("draft", "frozen"):
            export_kind = "draft"
        m = rmv2.get_meeting(mid)
        if not m:
            return error_response("meeting_id 不存在或已过期", status_code=404)

        gate = m.gate if isinstance(getattr(m, "gate", None), dict) else {}
        gate_state = str(gate.get("gate") or "").strip().upper()
        if export_kind == "frozen" and gate_state != "PASS":
            reason = str(gate.get("reason") or "门禁未通过").strip()
            return error_response(
                "无法导出“正式冻结版”：当前门禁未 PASS。请先闭环 P0/必补证据并重新开会生成 PASS。\n"
                + (f"当前结论：{gate_state or '—'}；原因：{reason}" if reason else f"当前结论：{gate_state or '—'}"),
                status_code=412,
            )
        llm_error = ""
        mode_used = "native"
        if use_rewrite_engine:
            cfg = _pick_llm_config_path()
            md, mode_used, llm_error = prd_rewrite.rewrite_prd(
                original_prd=str(m.prd_source or ""),
                blockers=m.blockers if isinstance(m.blockers, list) else [],
                decisions=m.decisions if isinstance(m.decisions, list) else [],
                prd_patch=str(m.prd_patch or ""),
                gate=gate,
                export_kind=export_kind,
                now_str=rmv2._now_str() if hasattr(rmv2, "_now_str") else "",
                llm_config_path=cfg,
                use_llm=True,
                timeout=120,
            )
        else:
            md = rmv2.export_prd_v2_markdown(m.to_dict())
        return success_response(
            data={
                "meeting_id": mid,
                "markdown": md,
                "mode_used": mode_used,
                "llm_error": llm_error,
                "use_rewrite_engine": use_rewrite_engine,
                "export_kind": export_kind,
                "gate": gate,
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_v2_export_prd_v2 failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/export_llm_only_summary", methods=["POST"])
def api_review_meeting_v2_export_llm_only_summary():
    """
    V2: Export an LLM-only, forwardable summary (Gate + top P0) with strict traceability.
    This does NOT change meeting state and does NOT replace the structured Gate.
    """
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        llm_required = bool(data.get("llm_required", False))
        # strict_mode=True: Scheme A first; if it fails, auto-fallback to LLM direct (still PRD-only, no hard refuse)
        # strict_mode=False: LLM direct only
        strict_mode = bool(data.get("strict_mode", True))
        if not mid:
            return error_response("meeting_id 不能为空", status_code=400)
        m = rmv2.get_meeting(mid)
        if not m:
            return error_response("meeting_id 不存在或已过期", status_code=404)

        gate = m.gate if isinstance(getattr(m, "gate", None), dict) else {}
        blockers = m.blockers if isinstance(getattr(m, "blockers", None), list) else []
        decisions = m.decisions if isinstance(getattr(m, "decisions", None), list) else []

        # Structured fallback (always available)
        gate_state = str((gate or {}).get("gate") or "—").strip().upper() or "—"
        gate_reason = str((gate or {}).get("reason") or "").strip()
        req_top = (gate or {}).get("required_evidence_top_llm")
        if not isinstance(req_top, list) or not req_top:
            req_top = (gate or {}).get("required_evidence_top")
        if not isinstance(req_top, list):
            req_top = []
        p0 = [
            b
            for b in blockers
            if isinstance(b, dict) and str(b.get("risk_level") or "").upper() == "P0"
        ]
        p0_blocked = [
            b
            for b in p0
            if str(b.get("status") or "").upper() == "BLOCKED"
        ]
        fb_lines = []
        fb_lines.append("## 可转发版结论（结构化兜底）")
        fb_lines.append(f"- 发布门禁/开工门禁结论：{gate_state}")
        if gate_reason:
            fb_lines.append(f"- 原因：{gate_reason}")
        if req_top:
            fb_lines.append("- Top 待补齐：")
            for x in req_top[:6]:
                xs = str(x or "").strip()
                if xs:
                    fb_lines.append(f"  - {xs}")
        if p0_blocked:
            fb_lines.append("- P0（卡住的前2条）：")
            for b in p0_blocked[:2]:
                title = str(b.get("title") or "未命名").strip()
                hs = str(b.get("human_summary") or "").strip()
                fb_lines.append(f"  - {title}" + (f"：{hs}" if hs else ""))
        fallback_text = "\n".join(fb_lines).strip()

        cfg = _pick_llm_config_path()
        if not cfg:
            if llm_required:
                return error_response("LLM 未配置（llm_config_path 为空）", status_code=412, data={"fallback": fallback_text})
            return success_response(data={"meeting_id": mid, "text": fallback_text, "mode_used": "fallback", "llm_error": "LLM 未配置", "fallback": fallback_text})

        # Build LLM prompt with strict constraints (no new facts; evidence required).
        # Scheme A: 2-pass LLM
        # - Pass 1: extract evidence-only facts (strict JSON, quotes must exist in PRD or known blocker quotes)
        # - Pass 2: write forwardable summary using ONLY extracted facts
        from utils.llm_client import call_llm_with_retry
        import re

        def _pick(x: dict, k: str, default: str = "") -> str:
            return str((x or {}).get(k) or default).strip()

        def _enforce_llm_only_safety(out: str) -> str:
            """
            Hard safety rules for forwardable text: no invented specifics, no code tokens.
            This is intentionally conservative: prefer "待补/未覆盖" over any concrete value.
            """
            t = str(out or "").strip()
            if not t:
                return t

            # Never allow markdown inline-code (commonly used to sneak in fields/codes).
            t = t.replace("`", "")

            # Gate is not PRD; re-label citation if model misuses it.
            t = t.replace("[SRC:PRD]", "[SRC:GATE]")

            # Normalize "TBD" into Chinese.
            t = re.sub(r"\bTBD\b", "待补齐", t, flags=re.IGNORECASE)

            # Remove typical "oral confirmed/decided" hallucination phrasing (keep sentence fluent).
            t = re.sub(
                r"PM\s*已?\s*口头\s*确(?:认|定)[^。\n]*",
                "该口径尚未在 PRD 中形成可追溯文本（需 PM 补齐并冻结）",
                t,
            )
            t = re.sub(r"已决定采用[^。\n]*", "冲突处理策略尚未形成可追溯文本（需 PM 冻结）", t)
            t = re.sub(r"已确定[^。\n]*", "相关规则尚未形成可追溯文本（需 PM 冻结）", t)

            # Mask list markers / Top3 to avoid over-sanitizing numbering.
            # e.g. "1. xxx", "2) xxx", "1、xxx", "Top3"
            t = re.sub(r"(?m)^\s*(\d+)\s*([.)、])", r"__LIST__\1\2", t)
            t = t.replace("Top3", "__TOP3__").replace("TOP3", "__TOP3__").replace("top3", "__TOP3__")

            # Strip / neutralize high-risk specifics ONLY (do not nuke all digits).
            # - error code tokens
            t = re.sub(r"\bERR_[A-Z0-9_]+\b", "（错误码名称待 PM 冻结）", t)
            # - explicit durations
            t = re.sub(r"\b\d+\s*(ms|毫秒|秒|s|分钟|min)\b", "（时长待 PM 冻结）", t, flags=re.IGNORECASE)
            # - retry/backoff counts/intervals (common patterns)
            t = re.sub(r"(重试|退避|间隔)\s*\d+\s*(次|轮|秒|s)?", r"\1（规则待 PM 冻结）", t, flags=re.IGNORECASE)
            t = re.sub(r"\b\d+\s*次\b", "（次数待 PM 冻结）", t)
            # - request method / explicit paths
            t = re.sub(r"\b(GET|POST|HTTP)\b", "（接口细节待 PM 冻结）", t, flags=re.IGNORECASE)
            t = re.sub(r"/[A-Za-z0-9_\-/{}/\.]+", "（路径待 PM 冻结）", t)

            # Prevent over-strong language when gate isn't PASS.
            if gate_state != "PASS":
                t = t.replace("正式冻结PRD", "正式冻结版 PRD")
                t = t.replace("可作为最终开测依据", "可作为最终依据")

            # Restore masks.
            t = t.replace("__TOP3__", "Top3")
            t = re.sub(r"__LIST__(\d+)([.)、])", r"\1\2", t)

            # Light "make sentences read better" cleanup (no new facts).
            t = re.sub(r"[ 　]+\n", "\n", t)
            t = re.sub(r"\n{3,}", "\n\n", t)
            t = re.sub(r"（\s*）", "", t)
            t = re.sub(r"：\s*（", "：（", t)
            t = re.sub(r"\s+（", "（", t)
            t = re.sub(r"）\s+", "）", t)
            return t.strip()

        def _contains_quote_in_sources(q: str, prd_text: str, extra_sources) -> bool:
            qq = str(q or "").strip()
            if not qq:
                return False
            if prd_text and qq in prd_text:
                return True
            for s in extra_sources or []:
                if s and qq in s:
                    return True
            return False

        def _validate_forwardable_output(out: str):
            t = str(out or "").strip()
            if not t:
                return False, "输出为空"
            # Must contain required three headers
            for h in ("## 评审结论（可转发版）", "## Top3 待补齐（按优先级）", "## P0 两条（各1-2行，面向PM）"):
                if h not in t:
                    return False, f"缺少小节标题：{h}"
            # Must contain at least 3 evidence lines
            if "依据：" not in t:
                return False, "缺少“依据：”行"
            # Forbid code-ish tokens and inline-code
            if "`" in t:
                return False, "包含反引号代码格式"
            if re.search(r"\bERR_[A-Z0-9_]+\b", t):
                return False, "包含错误码 token（ERR_...）"
            if re.search(r"\b(GET|POST|HTTP)\b", t, flags=re.IGNORECASE):
                return False, "包含请求方法/HTTP 字样"
            if re.search(r"/[A-Za-z0-9_\-/{}/\.]+", t):
                return False, "包含疑似路径"
            # Forbid explicit durations / retry numbers (high risk)
            if re.search(r"\b\d+\s*(ms|毫秒|秒|s|分钟|min)\b", t, flags=re.IGNORECASE):
                return False, "包含具体时长数值"
            if re.search(r"\b\d+\s*次\b", t):
                return False, "包含具体次数数值"
            return True, ""

        # Keep only a small, high-signal package to avoid drift.
        blocked_pack = []
        for b in p0_blocked[:6]:
            blocked_pack.append(
                {
                    "title": _pick(b, "title", "未命名"),
                    "module": _pick(b, "module", ""),
                    "summary": _pick(b, "human_summary", ""),
                    "required": (b.get("required_evidence") if isinstance(b.get("required_evidence"), list) else [])[:6],
                    "quotes": (b.get("evidence_quotes") if isinstance(b.get("evidence_quotes"), list) else [])[:2],
                }
            )
        decided_pack = []
        for d in decisions[:12]:
            if not isinstance(d, dict):
                continue
            if str(d.get("status") or "").upper() != "DECIDED":
                continue
            decided_pack.append(
                {
                    "title": _pick(d, "title", "未命名"),
                    "owner": _pick(d, "owner", "PM"),
                    "summary": _pick(d, "human_summary", ""),
                    "decision": _pick(d, "decision", ""),
                }
            )

        import json as _json

        prd_text = str(getattr(m, "prd_source", "") or "")

        def _llm_direct() -> dict:
            direct_prompt = f"""你是资深产品经理，请仅基于【原始PRD原文】写一段“可直接转发”的评审结论摘要（人话、可执行）。

要求：
1) 禁止编造：不要发明具体数值/接口路径/字段名/错误码 token/状态枚举/次数/间隔；若原文没有，就写“PRD 原文未覆盖，需要 PM 补齐”。
2) 结构固定（只输出这 3 个小节）：
   - ## 评审结论（可转发版）
   - ## Top3 待补齐（按优先级）
   - ## P0 两条（各1-2行，面向PM）
3) 尽量引用原文短摘录作为依据：每条关键结论后写“依据：”并引用一句原文（用引号），没有就写“PRD原文未覆盖该点”。
4) 不要工程黑话，不要用反引号。

原始PRD原文：
\"\"\"{prd_text[:18000]}\"\"\"

辅助信息（门禁与P0标题，仅用于帮助你定位重点，不可当作 PRD 事实来写死）：
{_json.dumps({
  "gate": {"gate": gate_state, "reason": gate_reason, "required_evidence_top": req_top[:8]},
  "p0_titles": [str(b.get("title") or "").strip() for b in p0_blocked[:6] if isinstance(b, dict)],
}, ensure_ascii=False)}

请直接输出 Markdown，不要解释。"""
            try:
                text = call_llm_with_retry(
                    messages=[{"role": "user", "content": direct_prompt}],
                    config_path=cfg,
                    timeout=90,
                    max_retries=3,
                    retry_delay=3,
                )
                text = str(text or "").strip()
                if not text:
                    raise RuntimeError("LLM 返回为空")
                # Best-effort safety (do not reject; just sanitize risky tokens)
                text = _enforce_llm_only_safety(text)
                return {
                    "meeting_id": mid,
                    "text": text,
                    "mode_used": "llm_direct",
                    "llm_error": "",
                    "fallback": fallback_text,
                }
            except Exception as e:
                raise

        # Non-strict: direct LLM output only
        if not strict_mode:
            try:
                return success_response(data=_llm_direct())
            except Exception as e:
                if llm_required:
                    return error_response(f"LLM 生成失败：{e}", status_code=412, data={"fallback": fallback_text})
                return success_response(
                    data={"meeting_id": mid, "text": fallback_text, "mode_used": "fallback", "llm_error": str(e), "fallback": fallback_text}
                )

        # Pass 1: evidence-only extraction (strict JSON)
        extract_prompt = f"""你是“证据抽取器”。你的唯一任务是：从【原始PRD原文】中抽取可引用的原句证据（短摘录），并整理成严格 JSON。

硬约束（必须遵守）：
1) 你只能输出 JSON（不要 Markdown、不要解释、不要多余文字）。
2) quote 必须是【原始PRD原文】里的原句短摘录（必须能在原文中逐字找到），长度 12-80 字。
3) 如果 PRD 原文完全没有覆盖某个点，不要编；用 covered=false，quote 为空字符串，并在 reason 写“PRD原文未覆盖该点”，并给出 src=[SRC:B#]（对应 blocker 序号）。
4) 禁止输出任何具体数值/接口路径/字段名/错误码 token；如果 PRD 原句里出现这些，也不要摘录那部分（换一段不包含这些的原句；实在没有就按未覆盖处理）。

输出 JSON schema（必须完全匹配）：
{{
  "gate": {{"gate":"PASS|FAIL|—","reason":"...","src":"[SRC:GATE]"}},
  "top3_gaps":[
    {{"title":"...","impact":"...","pm_questions":["..."],"covered":true|false,"quote":"...","src":"[SRC:PRD]|[SRC:B1]|[SRC:B2]"}}
  ],
  "p0_two":[
    {{"title":"...","ask":"...","covered":true|false,"quote":"...","src":"[SRC:PRD]|[SRC:B1]"}}
  ]
}}

原始PRD原文：
\"\"\"{prd_text[:12000]}\"\"\"

门禁与P0事实包（辅助你定位缺口标题，不可当作 PRD 证据）：
{_json.dumps({
  "gate": {"gate": gate_state, "reason": gate_reason, "required_evidence_top": req_top[:8]},
  "p0_blocked": blocked_pack,
}, ensure_ascii=False)}
"""

        # Pass 2: forwardable writing using ONLY extracted facts
        prompt = f"""你是资深产品经理，请基于【事实 JSON】写一段可直接转发给产品/研发的短文（人话、可执行）。

硬约束（必须遵守）：
1) 不得新增事实：禁止发明任何具体数值/接口路径/字段名/错误码 token/状态枚举/次数/间隔。
2) 只能使用【事实 JSON】里的字段与 quote；禁止输出事实 JSON 之外的信息。
3) 强制可追溯：每条关键结论后必须带“依据：…”，规则：
   - 若 covered=true：必须引用 quote（用引号包起来）并保留对应 src；
   - 若 covered=false：必须写“PRD原文未覆盖该点”并保留对应 src（[SRC:B#]）。
4) 输出只允许包含 3 个小节（标题必须一致）：
   - ## 评审结论（可转发版）
   - ## Top3 待补齐（按优先级）
   - ## P0 两条（各1-2行，面向PM）
5) 禁止出现代码符号/格式：不要使用反引号，不要输出路径/请求方法。

事实 JSON：
{{FACTS_JSON}}

请直接输出 Markdown，不要解释。"""

        try:
            # Pass 1 call
            facts_raw = call_llm_with_retry(
                messages=[{"role": "user", "content": extract_prompt}],
                config_path=cfg,
                timeout=90,
                max_retries=3,
                retry_delay=3,
            )
            facts_raw = str(facts_raw or "").strip()
            # Robust JSON extraction (model may wrap in ```json fences)
            if facts_raw.startswith("```"):
                facts_raw = re.sub(r"^```[a-zA-Z]*\s*", "", facts_raw).strip()
                facts_raw = re.sub(r"\s*```$", "", facts_raw).strip()
            facts = {}
            if facts_raw:
                try:
                    facts = _json.loads(facts_raw)
                except Exception:
                    mobj = re.search(r"\{[\s\S]*\}\s*$", facts_raw)
                    if mobj:
                        facts = _json.loads(mobj.group(0))
                    else:
                        raise
            if not isinstance(facts, dict):
                raise RuntimeError("Pass1 输出不是 JSON object")

            # Validate extracted quotes really exist in PRD (or known evidence quotes as extra source)
            extra_sources = []
            for b in blocked_pack:
                if isinstance(b, dict):
                    qs = b.get("quotes")
                    if isinstance(qs, list):
                        for q in qs:
                            if q:
                                extra_sources.append(str(q))

            def _check_items(items: list, label: str) -> None:
                if not isinstance(items, list):
                    raise RuntimeError(f"Pass1 字段 {label} 非 list")
                for it in items:
                    if not isinstance(it, dict):
                        raise RuntimeError(f"Pass1 {label} item 非 dict")
                    covered = bool(it.get("covered", False))
                    quote = str(it.get("quote") or "").strip()
                    if covered:
                        if not quote:
                            raise RuntimeError(f"Pass1 {label} 有 covered=true 但 quote 为空")
                        if not _contains_quote_in_sources(quote, prd_text, extra_sources):
                            raise RuntimeError(f"Pass1 {label} quote 不在原文/证据中：{quote[:30]}…")
                    else:
                        # covered=false must not provide quote
                        if quote:
                            raise RuntimeError(f"Pass1 {label} covered=false 但 quote 非空")

            _check_items(facts.get("top3_gaps"), "top3_gaps")
            _check_items(facts.get("p0_two"), "p0_two")

            # Pass 2 call
            prompt2 = prompt.replace("{FACTS_JSON}", _json.dumps(facts, ensure_ascii=False))
            text = call_llm_with_retry(
                messages=[{"role": "user", "content": prompt2}],
                config_path=cfg,
                timeout=90,
                max_retries=3,
                retry_delay=3,
            )
            text = str(text or "").strip()
            if not text:
                raise RuntimeError("LLM 返回为空")
            ok, why = _validate_forwardable_output(text)
            if not ok:
                raise RuntimeError(f"Pass2 输出校验失败：{why}")
            text = _enforce_llm_only_safety(text)
            return success_response(
                data={
                    "meeting_id": mid,
                    "text": text,
                    "mode_used": "llm_only",
                    "llm_error": "",
                    "fallback": fallback_text,
                }
            )
        except Exception as e:
            # Strict mode auto-fallback: still PRD-only, but don't block delivery
            try:
                d = _llm_direct()
                d["llm_error"] = f"证据版失败，已自动降级：{e}"
                return success_response(data=d)
            except Exception as e2:
                if llm_required:
                    return error_response(f"LLM-only 生成失败：{e}", status_code=412, data={"fallback": fallback_text})
                return success_response(
                    data={
                        "meeting_id": mid,
                        "text": fallback_text,
                        "mode_used": "fallback",
                        "llm_error": f"{e}；降级也失败：{e2}",
                        "fallback": fallback_text,
                    }
                )
    except Exception as e:
        logger.exception("api_review_meeting_v2_export_llm_only_summary failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/send_feishu", methods=["POST"])
def api_review_meeting_v2_send_feishu():
    """
    Send exported PRD v2.0 markdown to Feishu group bot webhook.
    This is a UI convenience API; it does not modify meeting state.
    """
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        webhook_url = str(data.get("webhook_url") or "").strip()
        title = str(data.get("title") or "PRD v2.0").strip()
        text = str(data.get("text") or "").strip()
        if not mid:
            return error_response("meeting_id 不能为空", status_code=400)
        if not webhook_url:
            return error_response("webhook_url 不能为空", status_code=400)
        if not text:
            return error_response("text 不能为空", status_code=400)

        # Basic allowlist for safety (Feishu/Lark webhooks)
        if ("open.feishu.cn" not in webhook_url) and ("open.larksuite.com" not in webhook_url):
            return error_response("webhook_url 域名不合法（仅支持飞书/ Lark 机器人 Webhook）", status_code=400)

        import json as _json
        import urllib.request as _urlreq

        def _post(payload: dict) -> None:
            body = _json.dumps(payload, ensure_ascii=False).encode("utf-8")
            req = _urlreq.Request(
                webhook_url,
                data=body,
                headers={"Content-Type": "application/json; charset=utf-8"},
                method="POST",
            )
            with _urlreq.urlopen(req, timeout=10) as resp:
                _ = resp.read()

        # Split long text into chunks (Feishu has message length limits)
        max_len = 3200
        chunks = []
        t = text
        while t:
            chunks.append(t[:max_len])
            t = t[max_len:]
            if len(chunks) > 12:
                # avoid runaway
                break

        sent = 0
        # First message includes title + meeting id
        head = f"{title}\nmeeting_id={mid}\n\n"
        if chunks:
            chunks[0] = head + chunks[0]
        for c in chunks:
            payload = {"msg_type": "text", "content": {"text": c}}
            _post(payload)
            sent += 1

        return success_response(data={"meeting_id": mid, "sent_count": sent})
    except Exception as e:
        logger.exception("api_review_meeting_v2_send_feishu failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/apply_freeze_rules", methods=["POST"])
def api_review_meeting_v2_apply_freeze_rules():
    """
    Apply PM-decided concrete rules to meeting (project-specific values).
    This does NOT "force PASS". It updates meeting.decisions/blockers and recomputes gate+prd_v2.
    """
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        payload = data.get("freeze_rules") if isinstance(data.get("freeze_rules"), dict) else {}
        m = rmv2.get_meeting(mid)
        if not m:
            return error_response("meeting_id 不存在或已过期", status_code=404)
        updated = rmv2.apply_freeze_rules(meeting=m, freeze_rules=payload)
        return success_response(
            data={
                "meeting_id": updated.meeting_id,
                "meeting": updated.to_dict(),
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_v2_apply_freeze_rules failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/start_async", methods=["POST"])
def api_review_meeting_v2_start_async():
    """
    V2：异步启动会议（用于 mode=prd 的耗时步骤），前端可轮询 progress 展示“正在做什么”。
    """
    try:
        data = request.get_json() or {}
        mode = str(data.get("mode") or "prd").strip().lower()
        if mode != "prd":
            return error_response("start_async 目前仅支持 mode=prd", status_code=400)
        prd_text = str(data.get("prd_text") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)
        scan_level = str(data.get("scan_level") or "quick").strip().lower()
        if scan_level not in ("quick", "full"):
            scan_level = "quick"
        enforce_gate = bool(data.get("enforce_gate", True))
        use_llm = bool(data.get("use_llm", True))
        llm_required = bool(data.get("llm_required", True))
        audit_constraint = bool(data.get("audit_constraint", True))

        tid = _rmv2_task_create({"mode": "prd"})
        _rmv2_task_append(tid, "已接收 PRD 输入，开始生成会议用快照…")

        def _runner():
            try:
                config_path = ""
                # Only require LLM config when we will call LLM (meeting discussion) or run full audit.
                if use_llm or scan_level == "full":
                    from utils.llm_client import load_llm_config

                    _rmv2_task_append(tid, "读取 LLM 配置…")
                    config_path = _pick_llm_config_path()
                    llm_config = load_llm_config(config_path)
                    if not (llm_config.get("api_key") or "").strip():
                        raise RuntimeError("API Key 未配置")
                else:
                    _rmv2_task_append(tid, "跳过 LLM 配置（use_llm=false 且 scan_level=quick）…")

                if scan_level == "full":
                    _rmv2_task_append(tid, "运行静态审计（Stage1~Stage3/L3 母表）…")
                else:
                    _rmv2_task_append(tid, "快速扫描生成会议（不跑六段式静态审计）…")
                hb_stop = {"stop": False}
                audit_timeout_s = int(data.get("audit_timeout_s") or 150)  # hard wall-clock cap for async task
                audit_started_at = time.time()

                def _hb():
                    start_ts = time.time()
                    while not hb_stop.get("stop"):
                        time.sleep(8)
                        if hb_stop.get("stop"):
                            break
                        waited = int(max(0, time.time() - start_ts))
                        total = int(max(0, time.time() - audit_started_at))
                        _rmv2_task_append(tid, f"静态审计仍在运行中…（已等待 {waited}s）")
                        if total >= audit_timeout_s:
                            # hard stop signal; runner will fail shortly
                            hb_stop["stop"] = True
                            _rmv2_task_append(tid, f"静态审计超过硬超时 {audit_timeout_s}s，准备终止并返回失败（避免一直卡住）。")

                snap_id_for_meeting = "prd_input_" + uuid.uuid4().hex[:8]
                if scan_level == "full":
                    threading.Thread(target=_hb, daemon=True).start()
                    # Run audit in a sub-thread so we can enforce a hard wall-clock timeout.
                    audit_out = {"stage3": None, "err": None}

                    def _run_audit():
                        try:
                            _, _, _, s3 = pipeline.run_prd_audit_sync(prd_text, llm_config_path=config_path, timeout=90)
                            audit_out["stage3"] = s3
                        except Exception as e:
                            audit_out["err"] = str(e)

                    th = threading.Thread(target=_run_audit, daemon=True)
                    th.start()
                    th.join(timeout=max(1, audit_timeout_s))
                    hb_stop["stop"] = True
                    if th.is_alive():
                        raise RuntimeError(f"静态审计执行超时（>{audit_timeout_s}s）。建议缩短 PRD 文本或检查 LLM/网络状态。")
                    if audit_out.get("err"):
                        raise RuntimeError("静态审计失败：" + str(audit_out.get("err")))
                    stage3_output = audit_out.get("stage3")
                    if not isinstance(stage3_output, dict) or not stage3_output:
                        raise RuntimeError("缺少 stage3_output，无法生成会议")
                    _rmv2_task_append(tid, "生成 V2 会议（完整版：基于审计输出）…")
                    meeting = rmv2.create_meeting_from_stage3(stage3_output, snapshot_id=snap_id_for_meeting, mode="prd", prd_content=prd_text)
                else:
                    _rmv2_task_append(tid, "生成 V2 会议（轻量版：不跑六段式审计）…")
                    meeting = rmv2.create_meeting_from_quick_scan(prd_text=prd_text, snapshot_id=snap_id_for_meeting, mode="prd")
                _rmv2_task_set_partial(tid, {"meeting_id": meeting.meeting_id, "meeting": meeting.to_dict()})

                llm_error = ""
                if use_llm and config_path:
                    _rmv2_task_append(tid, "启动多角色 LLM 讨论（QA→Dev→PM，限轮收敛）…")
                    meeting, llm_error = rmv2.run_llm_discussion_for_meeting(
                        meeting=meeting,
                        llm_config_path=config_path,
                        max_rounds=2,
                        audit_constraint=audit_constraint,
                        on_update=lambda m, note="": _rmv2_task_set_partial(
                            tid,
                            {"meeting_id": m.meeting_id, "meeting": m.to_dict(), "note": note},
                        ),
                    )
                if llm_required and (not use_llm or llm_error):
                    raise RuntimeError("LLM 必须参与，但本次 LLM 讨论失败：" + str(llm_error or "未知原因"))

                if enforce_gate:
                    _rmv2_task_append(tid, "门禁评估完成。")
                _rmv2_task_append(tid, "完成：已生成会议结果与 PRD v2.0。")
                _rmv2_task_finish(
                    tid,
                    result={
                        "meeting_id": meeting.meeting_id,
                        "meeting": meeting.to_dict(),
                        "use_llm": bool(use_llm and bool(config_path) and not bool(llm_error)),
                        "llm_error": llm_error,
                        "llm_required": llm_required,
                        "audit_constraint": audit_constraint,
                        "scan_level": scan_level,
                    },
                )
            except Exception as e:
                _rmv2_task_append(tid, "失败：" + str(e))
                _rmv2_task_fail(tid, str(e))

        threading.Thread(target=_runner, daemon=True).start()
        return success_response(data={"task_id": tid})
    except Exception as e:
        logger.exception("api_review_meeting_v2_start_async failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/task_status", methods=["GET", "POST"])
def api_review_meeting_v2_task_status():
    try:
        data = request.get_json(silent=True) or {}
        tid = str((data.get("task_id") or request.args.get("task_id") or "")).strip()
        if not tid:
            return error_response("task_id 不能为空", status_code=400)
        with _RMV2_TASKS_LOCK:
            t = _RMV2_TASKS.get(tid)
            if not isinstance(t, dict):
                return error_response("task_id 不存在或已过期", status_code=404)
            out = {
                "task_id": t.get("task_id"),
                "state": t.get("state"),
                "created_at": t.get("created_at"),
                "updated_at": t.get("updated_at"),
                "progress": t.get("progress") if isinstance(t.get("progress"), list) else [],
                "error": t.get("error") or "",
                "result": t.get("result"),
                "partial": t.get("partial"),
            }
        return success_response(data=out)
    except Exception as e:
        logger.exception("api_review_meeting_v2_task_status failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting_v2/save_snapshot", methods=["POST"])
def api_review_meeting_v2_save_snapshot():
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        m = rmv2.get_meeting(mid)
        if not m:
            return error_response("meeting_id 不存在或已过期", status_code=404)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        info = rmv2.save_meeting_snapshot(base_dir, m)
        return success_response(data={"meeting_id": mid, **info})
    except Exception as e:
        logger.exception("api_review_meeting_v2_save_snapshot failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting/start", methods=["POST"])
def api_review_meeting_start():
    """启动一场虚拟评审会：快照模式消费 Stage3；brainstorm 模式可无 PRD（推断+默认 BLOCKED）。"""
    try:
        data = request.get_json() or {}
        mode = _normalize_review_meeting_mode(data)
        sid = str(data.get("snapshot_id") or "").strip()
        use_llm = bool(data.get("use_llm", True))
        enforce_gate = bool(data.get("enforce_gate", True))
        from .audit_learning import SNAPSHOT_DIR, load_all_snapshots

        snapshot: Any = None
        agenda_cards: List[Dict[str, Any]] = []
        if mode == "brainstorm":
            idea = str(data.get("idea") or "").strip()
            if not idea:
                return error_response("brainstorm 模式需要提供 idea", status_code=400)
            agenda_cards = _brainstorm_agenda_cards(idea, use_llm=use_llm)
            agenda_cards = _attach_scenarios_to_cards(agenda_cards, use_llm=use_llm)
        else:
            if sid:
                file_path = os.path.join(SNAPSHOT_DIR, f"{sid}.json")
                if not os.path.exists(file_path):
                    return error_response("snapshot_id 不存在", status_code=404)
                with open(file_path, "r", encoding="utf-8") as f:
                    snapshot = json.load(f)
            else:
                snaps = load_all_snapshots(limit=1)
                snapshot = snaps[0] if snaps else None
            if not isinstance(snapshot, dict):
                return error_response("未找到可用快照，请先执行一次 PRD 审计生成快照", status_code=400)

            # 门禁：快照模式默认强制要求“可直接同步的共识摘要通过最小事实校验”且 P0 已清零。
            # 仅在 enforce_gate=False 时允许绕过（例如内部沙盘复盘/演示）。
            if enforce_gate:
                extras = snapshot.get("extras") if isinstance(snapshot.get("extras"), dict) else {}
                shared_summary = extras.get("shared_summary") if isinstance(extras.get("shared_summary"), dict) else {}
                gen_mode = str(shared_summary.get("generation_mode") or "").strip()
                failures = shared_summary.get("validation_failures") if isinstance(shared_summary.get("validation_failures"), list) else []
                p0_cnt = 0
                try:
                    stage2 = snapshot.get("stage2_output") if isinstance(snapshot.get("stage2_output"), dict) else {}
                    defects = stage2.get("defects") if isinstance(stage2.get("defects"), list) else []
                    p0_cnt = sum(
                        1
                        for d in defects
                        if isinstance(d, dict) and str(d.get("risk_level") or "").upper() == "P0"
                    )
                except Exception:
                    p0_cnt = 0
                gate_failures = []
                if p0_cnt > 0:
                    gate_failures.append(f"存在 P0 风险缺陷：{p0_cnt} 项")
                if gen_mode != "llm_validated":
                    gate_failures.append("全员共识摘要未通过最小事实校验（未使用 llm_validated 直写版本）")
                for f in failures[:6]:
                    fs = str(f or "").strip()
                    if fs:
                        gate_failures.append("摘要校验失败：" + fs)
                if gate_failures:
                    required = []
                    for f in gate_failures:
                        if "型号缺失" in f or "范围收缩" in f:
                            required.append("补齐《目标与范围》：明确盒子型号/版本/端，并确保枚举不被缩窄。")
                        if "红线缺失" in f:
                            required.append("补齐《数据生命周期/清空规则》：清空触发条件 + 不清空例外（如转台不清空）。")
                        if "冲突缺失" in f or "优先级" in f:
                            required.append("补齐《开关优先级裁决表》：设置总开关/播控栏开关/自动规则的抢占顺序。")
                        if "P0 风险缺陷" in f:
                            required.append("逐条关闭 P0：为每条 P0 给出可验收 AC + 最小观测字段 + 锚点证据。")
                    # 去重
                    dedup = []
                    seen = set()
                    for x in required:
                        xs = str(x or "").strip()
                        if not xs or xs in seen:
                            continue
                        seen.add(xs)
                        dedup.append(xs)
                    return error_response(
                        "门禁未通过：请先修正 PRD 再开评审会（或传 enforce_gate=false 绕过）。",
                        status_code=412,
                        data={
                            "gate_result": {
                                "mode": "GATE",
                                "pass": False,
                                "failures": gate_failures[:12],
                                "required_evidence": dedup[:12],
                                "signals": {
                                    "p0_defects_count": p0_cnt,
                                    "summary_generation_mode": gen_mode,
                                },
                            }
                        },
                    )
            agenda_cards = _build_agenda_cards_from_snapshot(snapshot, limit=3)
            agenda_cards = _attach_scenarios_to_cards(agenda_cards, use_llm=use_llm)

        meeting_id = f"meet_{uuid.uuid4().hex[:10]}"
        msgs = [{"role": "system", "speaker": "系统", "ts": _now_str(), "text": _meeting_system_message(mode)}]

        decision_log = []
        llm_ok = False
        llm_error = ""
        if use_llm:
            try:
                for c in agenda_cards:
                    try:
                        rounds = _llm_build_rounds_for_card(c)
                        # 对话区：展示三轮摘要（不直接把 JSON 长段塞进对话）
                        qa = rounds.get("qa") if isinstance(rounds.get("qa"), dict) else {}
                        dev = rounds.get("dev") if isinstance(rounds.get("dev"), dict) else {}
                        pm = rounds.get("pm") if isinstance(rounds.get("pm"), dict) else {}
                        qa_gaps = qa.get("qa_challenges") if isinstance(qa.get("qa_challenges"), list) else []
                        qa_lines = []
                        for g in qa_gaps[:4]:
                            if isinstance(g, dict) and str(g.get("gap") or "").strip():
                                qa_lines.append(f"- {str(g.get('gap')).strip()}")
                        msgs.append({"role": "assistant", "speaker": "测试（QA）", "ts": _now_str(), "text": f"议题：{c.get('one_liner')}\n不可验收点：\n" + ("\n".join(qa_lines) if qa_lines else "（无）")})
                        dev_rs = dev.get("dev_responses") if isinstance(dev.get("dev_responses"), list) else []
                        dev_lines = []
                        for r in dev_rs[:3]:
                            if isinstance(r, dict):
                                feas = str(r.get("feasibility") or "").strip()
                                prop = str(r.get("proposal") or "").strip()
                                if prop:
                                    dev_lines.append(f"- [{feas or 'n/a'}] {prop[:80]}")
                        msgs.append({"role": "assistant", "speaker": "研发（Dev）", "ts": _now_str(), "text": "最小实现回应：\n" + ("\n".join(dev_lines) if dev_lines else "（无）")})
                        pm_dec = pm.get("pm_decision") if isinstance(pm.get("pm_decision"), dict) else {}
                        ac_final = pm_dec.get("ac_final") if isinstance(pm_dec.get("ac_final"), list) else []
                        ac_lines = [f"- {str(x)[:120]}" for x in ac_final[:3] if str(x or "").strip()]
                        pm_msg = f"裁决：{str(pm_dec.get('decision_summary') or '').strip()}\nAC（节选）：\n" + ("\n".join(ac_lines) if ac_lines else "（无）")
                        if str(c.get("meeting_mode") or "").strip().lower() == "brainstorm":
                            pm_msg += "\n（无文档沙盘：Decision Log 侧默认 BLOCKED，直至补齐 PRD 锚点/量化口径）"
                        msgs.append({"role": "assistant", "speaker": "产品（PM）", "ts": _now_str(), "text": pm_msg})
                        decision_log.append(_decision_log_from_llm(c, rounds))
                        llm_ok = True
                    except Exception as e:
                        # 单条议题 LLM 失败：降级为模板，保证会议仍可用
                        err = str(e)
                        llm_error = err or llm_error
                        logger.warning("LLM meeting card failed, fallback to local. issue_id=%s err=%s", c.get("issue_id"), err)
                        # 追加模板版三句，避免 UI 空白
                        msgs.append({"role": "assistant", "speaker": "系统", "ts": _now_str(), "text": f"该议题 LLM 生成失败，已降级为本地模板。原因：{err}"})
                        msgs.extend(_seed_meeting_messages([c], mode))
                        decision_log.append(_build_decision_log_for_cards([c])[0])
            except Exception as e:
                llm_error = str(e)
                logger.warning("LLM meeting start failed, fallback to local: %s", llm_error)

        if not llm_ok and not decision_log:
            # 降级：模板会议（不影响能用）
            if use_llm and llm_error:
                msgs.append({"role": "assistant", "speaker": "系统", "ts": _now_str(), "text": f"LLM 评审会调用失败，已降级到本地模板会议。原因：{llm_error}"})
            msgs.extend(_seed_meeting_messages(agenda_cards, mode))
            decision_log = _build_decision_log_for_cards(agenda_cards)

        patch = _build_prd_patch(decision_log)
        snap_id_out = str(snapshot.get("snapshot_id") or "") if isinstance(snapshot, dict) else ""
        _REVIEW_MEETINGS[meeting_id] = {
            "meeting_id": meeting_id,
            "meeting_mode": mode,
            "snapshot_id": snap_id_out,
            "agenda_cards": agenda_cards,
            "decision_log": decision_log,
            "prd_patch": patch,
            "messages": msgs,
            "created_at": _now_str(),
            "use_llm": bool(use_llm and llm_ok),
        }
        return success_response(
            data={
                "meeting_id": meeting_id,
                "meeting_mode": mode,
                "snapshot_id": snap_id_out,
                "agenda_cards": agenda_cards,
                "decision_log": decision_log,
                "prd_patch": patch,
                "messages": msgs,
                "use_llm": bool(use_llm and llm_ok),
                "llm_error": llm_error if (use_llm and not llm_ok and llm_error) else "",
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_start failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review_meeting/send", methods=["POST"])
def api_review_meeting_send():
    """对话追加（测试补充/追问）。本地 MVP：记录对话并把内容沉淀到补丁末尾。"""
    try:
        data = request.get_json() or {}
        mid = str(data.get("meeting_id") or "").strip()
        text = str(data.get("text") or "").strip()
        if not mid or mid not in _REVIEW_MEETINGS:
            return error_response("meeting_id 不存在或已过期，请重新开始会议", status_code=404)
        if not text:
            return error_response("text 不能为空", status_code=400)
        st = _REVIEW_MEETINGS[mid]
        st["messages"].append({"role": "user", "speaker": "你", "ts": _now_str(), "text": text})
        # 系统回声：把用户补充收敛为“待确认/补丁增补”，保持测试主导风格
        st["messages"].append(
            {
                "role": "assistant",
                "speaker": "秘书（收敛）",
                "ts": _now_str(),
                "text": "已记录为补充关注点：将进入会议纪要/PRD补丁的“待确认/验收补充”段落。",
            }
        )
        # 追加到补丁末尾（不改原裁决表，只做补充）
        st["prd_patch"] = (st.get("prd_patch") or "") + "\n\n### 会议补充（测试关注点）\n- " + text.replace("\n", "\n- ")
        return success_response(
            data={
                "meeting_id": mid,
                "messages": st.get("messages") or [],
                "decision_log": st.get("decision_log") or [],
                "prd_patch": st.get("prd_patch") or "",
            }
        )
    except Exception as e:
        logger.exception("api_review_meeting_send failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/knowledge")
def knowledge_page():
    return render_template("prd_audit_knowledge.html")

@prd_audit_bp.route("/learning_mvp")
def learning_mvp_page():
    return render_template("prd_audit_learning_mvp.html")

@prd_audit_bp.route("/rules")
def rules_page():
    return render_template("prd_audit_rules.html")

@prd_audit_bp.route("/bug_patterns")
def bug_patterns_page():
    return render_template("prd_audit_bug_patterns.html")

@prd_audit_bp.route("/matrix_view")
def matrix_view_page():
    return render_template("prd_audit_matrix_view.html")

@prd_audit_bp.route("/platform_center")
def platform_center_page():
    context = _build_platform_center_context()
    return render_template(
        "prd_audit_platform_center.html",
        asset_summary=context.get("summary") or {},
        asset_groups=context.get("groups") or [],
    )

# ---------- API：生成（流式） ----------

@prd_audit_bp.route("/api/generate", methods=["POST"])
def api_generate():
    """PRD 三段式或单次调用生成（流式 NDJSON）"""
    try:
        data = request.get_json() or {}
        gen_type = data.get("type")
        content = data.get("content")
        use_llm = bool(data.get("use_llm", True))
        custom_prompt = data.get("prompt")
        report_level = (data.get("report_level") or "L3").upper()

        if not gen_type or gen_type != "prd_review":
            return error_response("type 必须为 prd_review", status_code=400)
        if not content or (isinstance(content, str) and not content.strip()):
            return error_response("content 不能为空", status_code=400)

        content = content.strip() if isinstance(content, str) else content

        if isinstance(content, str) and is_feishu_doc_url(content):
            ok, result = fetch_feishu_doc_content(content)
            if not ok:
                return error_response(
                    result if result else "无法拉取飞书文档，请复制内容粘贴或导出后上传",
                    status_code=400,
                )
            content = result
            logger.info("Feishu doc content fetched, length=%s", len(content))

        # 是否启用大模型：前端可选。use_llm = False 时强制使用本地规则体检模式。
        from utils.llm_client import load_llm_config
        config_path = _pick_llm_config_path()
        if use_llm:
            # 用户显式选择 use_llm=true 时，不允许静默降级，否则会出现“我明明用大模型但结果全是本地套话”的错觉。
            try:
                llm_config = load_llm_config(config_path)
                api_key = (llm_config.get("api_key") or "").strip()
                if not api_key:
                    return error_response(
                        f"LLM API Key 未配置，请先在「LLM 配置」中填写并保存（当前配置文件：{config_path}）",
                        status_code=400,
                    )
            except FileNotFoundError:
                return error_response(
                    f"LLM 配置文件不存在：{config_path}。请先在「LLM 配置」中保存一次。",
                    status_code=400,
                )
            except Exception as e:
                return error_response(
                    f"加载 LLM 配置失败：{e}",
                    status_code=400,
                )
        else:
            # 显式禁用 LLM：使用一个真实存在的 disabled 配置文件，避免“配置不存在”噪音
            logger.info("PRD 审计本次调用被配置为不使用大模型（use_llm=false），将仅使用本地规则体检模式。")
            disabled_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")
            try:
                os.makedirs(STORAGE_DIR, exist_ok=True)
                if not os.path.exists(disabled_path):
                    with open(disabled_path, "w", encoding="utf-8") as f:
                        json.dump(
                            {
                                "llm_provider": "deepseek",
                                "base_url": "https://api.deepseek.com/v1",
                                "api_key": "",
                                "model": "deepseek-chat",
                                "fallback_enabled": False,
                            },
                            f,
                            ensure_ascii=False,
                            indent=2,
                        )
            except Exception as e:
                logger.warning("写入 __llm_disabled__.json 失败（将继续走本地解析兜底）：%s", e)
            config_path = disabled_path

        def generate():
            try:
                for chunk in pipeline.run_prd_audit_stream(
                    content,
                    llm_config_path=config_path,
                    timeout=180,
                    custom_prompt=custom_prompt.strip() if custom_prompt and custom_prompt.strip() else None,
                    report_level=report_level,
                ):
                    yield chunk
            except Exception as e:
                logger.exception("PRD generate failed")
                yield json.dumps({"type": "error", "text": str(e)}, ensure_ascii=False) + "\n"

        return Response(
            stream_with_context(generate()),
            content_type="application/x-ndjson; charset=utf-8",
        )
    except Exception as e:
        logger.exception("api_generate failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/analyze_impact", methods=["POST"])
def api_analyze_impact():
    """架构级变更影响分析"""
    try:
        data = request.get_json() or {}
        change_desc = (data.get("change_desc") or "").strip()
        scan_result = data.get("scan_result")

        if not change_desc:
            return error_response("需求变更描述不能为空", status_code=400)
        if not scan_result or not isinstance(scan_result, dict):
            return error_response("缺少架构透视扫描数据", status_code=400)

        from .architecture_scanner import analyze_impact
        impact_result = analyze_impact(change_desc, scan_result)
        
        # 将 impact_result 转换为 Markdown 报告返回前端
        # （这里的生成逻辑可以放在后端，也可以让前端渲染 JSON。由于原先有 generate_html_report 的思路，这里生成简单的 Markdown）
        
        affected_modules = impact_result.get("affected_modules", [])
        affected_states = impact_result.get("affected_states", [])
        affected_apis = impact_result.get("affected_apis", [])
        affected_entities = impact_result.get("affected_entities", [])
        risk_assessment = impact_result.get("risk_assessment", {})

        # --- 将本地正则分析结果交由 LLM 进行二次润色与智能推理 ---
        from utils.llm_client import call_llm
        import json
        
        # 提取本地分析的上下文，喂给大模型
        context_data = {
            "change_desc": change_desc,
            "local_analysis": impact_result
        }
        
        prompt = f"""你是一个资深的软件架构师和测试专家。
用户提出了一个需求变更（一句话描述）。我已经通过代码/文档解析引擎，提取了系统当前的模块、状态机、API和实体数据，并做了一次初步的影响分析。

【用户变更需求】：
{change_desc}

【本地初步分析结果】：
{json.dumps(context_data['local_analysis'], ensure_ascii=False, indent=2)}

请你基于以上信息，结合你对软件工程的理解，输出一份**更智能、更具可读性、且具有深度业务推理**的《变更影响分析报告》。
注意：
1. 本地分析可能会漏掉一些隐含的业务逻辑关联（比如“星耀屏广告”可能与“投屏”、“点歌”存在抢占关系），请发挥你的推理能力进行补充。
2. 即使本地分析显示“None”或“无明显影响”，如果你认为在实际业务中可能会有影响，请明确指出并解释原因。
3. 请直接输出 Markdown 格式的报告。包含以下部分：
   - 📌 变更核心逻辑解读
   - 🎯 综合风险评估（整体影响、风险等级）
   - 🧩 受影响的模块与接口（说明影响原因）
   - 🔄 受影响的状态机与数据流
   - 🧪 回归测试策略建议与高风险点

输出报告："""

        try:
            llm_report = call_llm(messages=[{"role": "user", "content": prompt}])
            return success_response({"impact_report": llm_report})
        except Exception as e:
            logger.exception("LLM impact analysis failed")
            # 降级：如果大模型调用失败，退回到本地拼接的简单报告
            md = [f"### 变更影响分析报告：{change_desc}\n"]
            
            md.append("#### 1. 综合风险评估")
            md.append(f"- **整体影响范围**: {risk_assessment.get('overall_impact', '未知')}")
            md.append(f"- **回归测试建议级别**: {risk_assessment.get('regression_level', '未知')}")
            md.append(f"- **高风险点**: {risk_assessment.get('high_risk_points', 0)} 个")
            md.append("\n**测试策略建议**:")
            for advice in risk_assessment.get("testing_advice", []):
                md.append(f"- {advice}")
                
            md.append("\n#### 2. 受影响的模块")
            if affected_modules:
                for m in affected_modules:
                    md.append(f"- **{m.get('name')}** (影响: {m.get('impact_scope')}, 风险: {m.get('risk')})")
                    if m.get('dependencies'):
                        md.append(f"  - 依赖方: {', '.join(m.get('dependencies'))}")
            else:
                md.append("- 无明显受影响的模块")

            md.append("\n#### 3. 受影响的状态/流程")
            if affected_states:
                for s in affected_states:
                    md.append(f"- **{s.get('state')}** (模块: {s.get('module')}, 影响: {s.get('impact_scope')})")
                    for t in s.get('related_transitions', []):
                        md.append(f"  - 关联跃迁: {t}")
            else:
                md.append("- 无明显受影响的状态")

            md.append("\n#### 4. 受影响的 API 接口")
            if affected_apis:
                for a in affected_apis:
                    md.append(f"- **{a.get('api')}** (影响: {a.get('impact_scope')})")
                    md.append(f"  - 调用关系: {a.get('caller')} -> {a.get('callee')}")
            else:
                md.append("- 无明显受影响的 API")

            md.append("\n#### 5. 受影响的数据实体")
            if affected_entities:
                for e in affected_entities:
                    md.append(f"- **{e.get('entity')}** (影响: {e.get('impact_scope')})")
                    md.append(f"  - 变更动作: {e.get('action')}")
                    if e.get('related_fields'):
                        md.append(f"  - 关联字段: {', '.join(e.get('related_fields'))}")
            else:
                md.append("- 无明显受影响的数据实体")

            return success_response({
                "impact_report": "\n".join(md),
                "raw_data": impact_result
            })
    except Exception as e:
        logger.exception("api_analyze_impact failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/chat", methods=["POST"])
def api_chat():
    """AI智能助手对话接口"""
    try:
        data = request.get_json() or {}
        user_message = (data.get("message") or "").strip()
        context = data.get("context") or {}
        
        if not user_message:
            return error_response("消息不能为空", status_code=400)
            
        # 组装 Prompt
        system_prompt = "你是一个专门负责辅助 PRD 审计和架构分析的资深产品研发专家。请根据提供的上下文回答用户的问题。回答要专业、简明扼要。"
        
        context_str = "【当前上下文】\n"
        if context.get("prd_content"):
            # 如果 PRD 太长，截断一下以防 Token 超限
            prd_text = str(context["prd_content"])
            context_str += f"PRD 内容摘要: {prd_text[:2000]}...\n"
        
        if context.get("architecture_scan"):
            context_str += f"架构分析结果: {json.dumps(context['architecture_scan'], ensure_ascii=False)[:1000]}...\n"
            
        if context.get("audit_report"):
            # 只取核心摘要
            context_str += f"当前审计报告状态: 已生成报告\n"
            
        full_prompt = f"{system_prompt}\n\n{context_str}\n\n【用户问题】\n{user_message}"
        
        # 调用大模型
        response_text = call_llm(
            prompt=full_prompt,
            config_file=LLM_CONFIG_FILE,
            system_prompt=system_prompt,
            max_tokens=1500
        )
        
        if not response_text:
            response_text = "抱歉，我暂时无法回答这个问题，请检查 LLM 配置是否正确。"
            
        return success_response({"reply": response_text})
        
    except Exception as e:
        logger.exception("api_chat failed")
        return error_response(str(e), status_code=500)

# ---------- API：大纲专用（大模型，通用 PRD） ----------

@prd_audit_bp.route("/api/outline_llm", methods=["POST"])
def api_outline_llm():
    """仅生成「四支柱」认知大纲 JSON（prd_outline_llm_prompt.txt），适配任意 PRD；可选合并本地 outline_engine。"""
    try:
        data = request.get_json() or {}
        prd_text = (data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        merge_local = bool(data.get("merge_local", False))
        run_stage1 = data.get("run_stage1", True)
        if isinstance(run_stage1, str):
            run_stage1 = run_stage1.strip().lower() not in ("0", "false", "no", "")

        stage1_override = data.get("stage1_output")
        timeout = int(data.get("timeout", 120) or 120)
        if timeout < 30:
            timeout = 30
        if timeout > 600:
            timeout = 600

        from utils.llm_client import load_llm_config

        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
        except FileNotFoundError:
            return error_response("API Key 未配置（prd_audit 与 test_case 均未配置）", status_code=400)
        if not (llm_config.get("api_key") or "").strip():
            return error_response("API Key 未配置（prd_audit 与 test_case 均未配置）", status_code=400)

        from .system_model import extract_prd_structure

        if run_stage1:
            stage1 = extract_prd_structure(
                prd_text,
                llm_config_path=config_path,
                timeout=min(90, timeout),
            )
        else:
            stage1 = stage1_override if isinstance(stage1_override, dict) else {}
            if not stage1:
                return error_response("run_stage1=false 时需传入非空 stage1_output", status_code=400)

        llm_result = run_outline_llm(
            prd_text,
            stage1,
            config_path,
            timeout=timeout,
        )

        payload = {
            "llm": llm_result,
            "stage1_output": stage1,
        }
        if merge_local:
            local_engine = run_outline_engine(prd_text, stage1, {})
            payload["local_outline_engine"] = local_engine
            payload["merged"] = merge_llm_with_local(
                llm_result.get("llm_outline") or {},
                local_engine,
            )

        return success_response(
            data=payload,
            message="大纲生成完成" if llm_result.get("ok") else (llm_result.get("error") or "大纲生成未完全成功"),
        )
    except Exception as e:
        logger.exception("api_outline_llm failed")
        return error_response(str(e), status_code=500)

# ---------- API：分析（同步，返回完整 JSON） ----------

@prd_audit_bp.route("/api/analyze_prd", methods=["POST"])
def api_analyze_prd():
    """非流式 PRD 分析，返回 raw_report_markdown、stage1/2/3、summary 等"""
    try:
        data = request.get_json() or {}
        prd_text = (data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        from utils.llm_client import load_llm_config
        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
        except FileNotFoundError:
            return error_response("API Key 未配置（prd_audit 与 test_case 均未配置）", status_code=400)
        if not (llm_config.get("api_key") or "").strip():
            return error_response("API Key 未配置（prd_audit 与 test_case 均未配置）", status_code=400)

        merged_report, stage1_output, stage2_output, stage3_output = pipeline.run_prd_audit_sync(
            prd_text, llm_config_path=config_path, timeout=90
        )
        # 纯 PRD 分析模式下跳过测试资产生成（避免与“仅分析 PRD”意图冲突）
        if getattr(pipeline, "PRD_ANALYSIS_ONLY_MODE", False):
            test_matrix = {}
            test_points_obj = {}
            validation_outline = {}
        else:
            try:
                from .test_matrix_generator import TestMatrixGenerator
                test_matrix = TestMatrixGenerator(stage1_output, stage2_output).generate()
            except Exception:
                logger.exception("api_analyze_prd: generate test_matrix failed")
                test_matrix = {}
            try:
                from .test_points_engine import run_test_points_engine, generate_validation_outline
                test_points_obj = run_test_points_engine(
                    prd_text=prd_text,
                    stage1_output=stage1_output if isinstance(stage1_output, dict) else {},
                    stage2_output=stage2_output if isinstance(stage2_output, dict) else {},
                    outline_engine={},
                    platform_impact={},
                    dependency_analysis={},
                    test_matrix=test_matrix if isinstance(test_matrix, dict) else {},
                )
                # 生成验证大纲
                validation_outline = generate_validation_outline(test_points_obj)
            except Exception:
                logger.exception("api_analyze_prd: generate test_points failed")
                test_points_obj = {}
                validation_outline = {}

        score = stage3_output.get("summary", {}).get("quality_score")
        # run_prd_audit_sync 已构建 shared_summary/reader_guide 并挂到 stage3_output，避免重复调用浪费 token
        shared_summary = stage3_output.get("shared_summary") if isinstance(stage3_output, dict) else {}
        reader_guide = stage3_output.get("reader_guide") if isinstance(stage3_output, dict) else {}
        if not isinstance(shared_summary, dict):
            shared_summary = {}
        if not isinstance(reader_guide, dict):
            reader_guide = {}
        quality_summary = {
            "overall": score,
            "dimensions": {},
            "rule_engine_score": score,
        }
        states = pipeline._ensure_list(stage1_output.get("states"))
        state_diagram_mermaid = ""
        if states:
            cleaned = [s.replace(" ", "_") for s in states]
            lines = ["stateDiagram-v2", f"  [*] --> {cleaned[0]}"]
            for i in range(len(cleaned) - 1):
                lines.append(f"  {cleaned[i]} --> {cleaned[i + 1]}")
            state_diagram_mermaid = "\n".join(lines)
        defects = stage3_output.get("defects") or []
        prd_gaps = "\n".join([f"- {d.get('module', '【PRD未说明】')}：{d.get('description', '【PRD未说明】')}" for d in defects[:20]])
        risks = "\n".join([f"- {x}" for x in (stage3_output.get("risks") or [])])
        rd_focus = "\n".join([f"- {x}" for x in (stage3_output.get("dev_focus") or [])])
        test_focus = "\n".join([f"- {x}" for x in (stage3_output.get("test_focus") or [])])
        plan_list = stage3_output.get("plan") or []
        plan_level = "P1"
        if any((d.get("risk_level") or "").upper() == "P0" for d in defects):
            plan_level = "P0"
        elif defects and all((d.get("risk_level") or "").upper() == "P2" for d in defects):
            plan_level = "P2"
        plan = {
            "level": plan_level,
            "order_advice": plan_list[:2],
            "prep_advice": plan_list[2:4],
        }
        section_1 = "## 一、总体结论\n\n- 质量评分：{}/10\n- 风险等级：{}\n- 主要问题：{}".format(
            stage3_output.get("summary", {}).get("quality_score", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("risk_level", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("main_problem", "【PRD未说明】"),
        )
        section_2 = "## 二、漏洞与风险清单\n\n" + (prd_gaps or "- 未发现漏洞")
        section_4 = "## 四、测试重点\n\n" + (test_focus or "- 【PRD未说明】")
        section_5 = "## 五、研发重点\n\n" + (rd_focus or "- 【PRD未说明】")
        section_6 = "## 六、计划建议\n\n" + ("\n".join([f"- {x}" for x in plan_list]) if plan_list else "- 【PRD未说明】")
        result = {
            "raw_report_markdown": merged_report,
            "status": ["Stage1完成", "Stage2完成", "Stage3完成", "Stage4完成"],
            "summary": section_1,
            "modules": stage1_output.get("modules") or [],
            "prd_gaps": prd_gaps,
            "risks": risks,
            "rd_focus": rd_focus,
            "test_focus": test_focus,
            "test_points": test_points_obj,
            "validation_outline": validation_outline,  # 新增：验证大纲
            "plan": plan,
            "quality": quality_summary,
            "stage1_output": stage1_output,
            "stage2_output": stage2_output,
            "stage3_output": stage3_output,
            "shared_summary": shared_summary if isinstance(shared_summary, dict) else {},
            "reader_guide": reader_guide if isinstance(reader_guide, dict) else {},
            "test_matrix": test_matrix if isinstance(test_matrix, dict) else {},
            "system_model": {
                "states": stage1_output.get("states") or [],
                "events": [],
                "rules": stage1_output.get("business_rules") or [],
                "data_rules": stage1_output.get("exceptions") or [],
                "data_structures": stage1_output.get("data_structures") or [],
                "permissions": stage1_output.get("permissions") or [],
                "edge_cases": stage1_output.get("edge_cases") or [],
                "dependencies": stage1_output.get("dependencies") or [],
                "transitions": [],
                "high_priority_events": [],
            },
            "rule_analysis": stage2_output,
            "test_point_matrix": (test_points_obj.get("test_point_matrix") if isinstance(test_points_obj, dict) else {}) or {"version": "v1", "items": [], "stats": {"total": 0, "p0": 0, "p1": 0, "p2": 0, "by_module": {}}},
            "sections": {
                "1": section_1,
                "2": section_2,
                "4": section_4,
                "5": section_5,
                "6": section_6,
            },
            "state_diagram_mermaid": state_diagram_mermaid,
        }
        return success_response(data=result)
    except Exception as e:
        logger.exception("api_analyze_prd failed")
        return error_response(str(e), status_code=500)

def _evaluate_bugs_for_preview(items: list) -> list:
    """对即将导入的 Bug 列表进行 AI 预审"""
    if not items:
        return items

    # 准备 LLM 评估的 Prompt
    bug_texts = []
    for i, it in enumerate(items):
        desc = str(it.get("bug_desc", "")).strip()
        sev = str(it.get("severity", "P2")).strip()
        bug_texts.append(f"[{i}] 描述: {desc} | 严重度: {sev}")

    prompt = f"""你是一个高级测试架构师，正在审查一批准备导入到“PRD 审计系统”的 Bug 缺陷记录。
我们的目标是通过这些历史 Bug 来训练系统的 PRD 审计规则，使其能发现 PRD 中“缺失的业务逻辑、异常处理、边界条件和状态流转等设计缺陷”。

请逐一评估以下 Bug 是否有导入价值，并严格按以下标准分类：
1. "推荐导入": 逻辑漏洞、状态流转错误、并发冲突、异常流程未闭环、越权等设计/场景缺失类 Bug。
2. "建议修正": 有价值，但描述过于简略，缺少上下文或具体表现，建议补充完整场景后导入。
3. "建议废弃": 纯 UI/样式细节偏差、拼写错误、环境配置问题、明显的代码手误等（对 PRD 约束意义不大的低级错误）。

输入 Bug 列表：
{chr(10).join(bug_texts)}

请返回 JSON 数组，每个元素包含：
- "index": 输入的序号（整数）
- "ai_status": "推荐导入" 或 "建议修正" 或 "建议废弃"
- "ai_reason": 简短的评估理由（15字以内）

只输出合法的 JSON 数组，不要包含任何 markdown 代码块标记，不要多余说明。"""

    # 回退到本地规则的辅助函数
    def _fallback_eval(bug_item):
        desc = bug_item.get("bug_desc", "").lower()
        if any(kw in desc for kw in ["ui", "颜色", "对其", "拼写", "环境", "配置", "错别字", "字体", "间距", "样式"]):
            return {"ai_status": "建议废弃", "ai_reason": "UI/环境类无关PRD"}
        elif len(desc) < 6:
            return {"ai_status": "建议修正", "ai_reason": "描述过于简略"}
        else:
            return {"ai_status": "推荐导入", "ai_reason": "本地规则通过"}

    try:
        config_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', 'config', 'llm_config.json'))
        resp = call_llm([{"role": "user", "content": prompt}], config_path=config_path, timeout=30)
        resp_text = str(resp or "").strip()
        # 清理可能存在的 markdown 代码块
        if resp_text.startswith("```json"):
            resp_text = resp_text[7:]
        if resp_text.startswith("```"):
            resp_text = resp_text[3:]
        if resp_text.endswith("```"):
            resp_text = resp_text[:-3]
        resp_text = resp_text.strip()
        
        eval_list = json.loads(resp_text)
        if isinstance(eval_list, list):
            eval_map = {int(x.get("index", -1)): x for x in eval_list if isinstance(x, dict)}
            for i, it in enumerate(items):
                ev = eval_map.get(i)
                if ev:
                    it["ai_status"] = ev.get("ai_status", "推荐导入")
                    it["ai_reason"] = ev.get("ai_reason", "AI认为有价值")
                else:
                    fb = _fallback_eval(it)
                    it.update(fb)
        else:
            raise ValueError("LLM 返回的不是 JSON 数组")
    except Exception as e:
        logger.warning(f"AI 预审 Bug 失败，回退到本地启发式规则: {e}")
        for it in items:
            fb = _fallback_eval(it)
            it.update(fb)
            
    return items

@prd_audit_bp.route("/api/bug/import", methods=["POST"])
def api_bug_import():
    try:
        _ensure_bug_assets()
        is_preview = request.args.get("preview", "false").lower() == "true"
        items = []
        if "file" in request.files:
            f = request.files["file"]
            raw = f.read()
            txt = raw.decode("utf-8-sig", errors="ignore")
            sample = "\n".join(txt.splitlines()[:3])
            delimiter = "\t" if ("\t" in sample and sample.count("\t") >= sample.count(",")) else ","
            reader = csv.DictReader(io.StringIO(txt), delimiter=delimiter)
            for row in reader:
                if not isinstance(row, dict):
                    continue
                bug_desc = str(
                    row.get("bug_desc")
                    or row.get("desc")
                    or row.get("概要")
                    or row.get("问题描述")
                    or row.get("summary")
                    or ""
                ).strip()
                if not bug_desc:
                    continue
                sev = (
                    row.get("severity")
                    or row.get("自定义字段(严重程度)")
                    or row.get("严重程度")
                    or row.get("等级")
                    or row.get("自定义字段(优先级)")
                    or row.get("优先级")
                )
                items.append(
                    {
                        "bug_desc": bug_desc,
                        "severity": _normalize_severity(sev),
                        "frequency": _normalize_frequency(row.get("frequency")),
                    }
                )
        else:
            data = request.get_json() or {}
            if isinstance(data.get("items"), list):
                for row in data.get("items"):
                    if not isinstance(row, dict):
                        continue
                    bug_desc = str(row.get("bug_desc") or "").strip()
                    if not bug_desc:
                        continue
                    items.append(
                        {
                            "bug_desc": bug_desc,
                            "severity": _normalize_severity(row.get("severity")),
                            "frequency": _normalize_frequency(row.get("frequency")),
                        }
                    )
            text_blob = str(data.get("text") or "").strip()
            if text_blob:
                # 尝试检测是否为 Jira 样式的 TSV/CSV 文本
                # 特征：第一行包含 "问题关键字" 或 "概要" 且含有 tab 或 多个空格
                first_line = text_blob.splitlines()[0] if text_blob else ""
                is_jira_format = False
                delimiter = None
                
                if "问题关键字" in first_line or "概要" in first_line:
                    if "\t" in first_line:
                        delimiter = "\t"
                        is_jira_format = True
                    elif "  " in first_line: # 可能是多个空格分隔
                        # 尝试用正则替换多个空格为 tab，或者直接解析
                        is_jira_format = True
                        delimiter = "regex_spaces"

                if is_jira_format:
                    lines = text_blob.splitlines()
                    # 如果是正则空格分隔，先处理一下 header
                    if delimiter == "regex_spaces":
                        # 简单策略：每一行都用正则 split，然后取对应 index
                        # 这是一个比较脆弱的策略，但对于复制粘贴的文本可能有效
                        # 更好的方式是看 header 的列位置
                        header_parts = re.split(r'\s{2,}|\t', lines[0].strip())
                        # 寻找 "概要" 和 "自定义字段(严重程度)" / "自定义字段(优先级)" 的 index
                        summary_idx = -1
                        severity_idx = -1
                        priority_idx = -1
                        
                        for idx, col in enumerate(header_parts):
                            c = col.strip()
                            if c == "概要": summary_idx = idx
                            if "严重程度" in c: severity_idx = idx
                            if "优先级" in c: priority_idx = idx
                        
                        if summary_idx != -1:
                            for line in lines[1:]:
                                if not line.strip(): continue
                                parts = re.split(r'\s{2,}|\t', line.strip())
                                # Parts 长度可能不一致，因为有些列可能为空导致合并
                                # 这里只能尽力而为。如果 parts 长度不够，可能无法准确提取
                                # 但用户给的例子看起来对齐得还不错
                                if len(parts) > summary_idx:
                                    bug_desc = parts[summary_idx].strip()
                                    sev = "P2"
                                    if severity_idx != -1 and len(parts) > severity_idx:
                                        sev = parts[severity_idx].strip()
                                    if priority_idx != -1 and len(parts) > priority_idx:
                                        # 优先级可能覆盖严重程度
                                        p = parts[priority_idx].strip()
                                        if p: sev = p
                                    
                                    if bug_desc:
                                        items.append({
                                            "bug_desc": bug_desc, 
                                            "severity": _normalize_severity(sev), 
                                            "frequency": "中"
                                        })
                    else:
                        # 标准 TSV 解析
                        try:
                            reader = csv.DictReader(io.StringIO(text_blob), delimiter=delimiter)
                            for row in reader:
                                bug_desc = str(
                                    row.get("bug_desc")
                                    or row.get("desc")
                                    or row.get("概要")
                                    or row.get("问题描述")
                                    or row.get("summary")
                                    or ""
                                ).strip()
                                if not bug_desc: continue
                                sev = (
                                    row.get("severity")
                                    or row.get("自定义字段(严重程度)")
                                    or row.get("严重程度")
                                    or row.get("等级")
                                    or row.get("自定义字段(优先级)")
                                    or row.get("优先级")
                                )
                                items.append({
                                    "bug_desc": bug_desc,
                                    "severity": _normalize_severity(sev),
                                    "frequency": "中"
                                })
                        except Exception:
                            # 解析失败回退到普通按行处理
                            pass

                # 如果没有被识别为 Jira 格式，或者解析后 items 仍为空（解析失败），则回退到逐行处理
                if not items:
                    for line in text_blob.splitlines():
                        d = str(line or "").strip()
                        if d:
                            items.append({"bug_desc": d, "severity": "P2", "frequency": "中"})

        if not items:
            return error_response("未解析到可导入 Bug", status_code=400)

        if is_preview:
            evaluated_items = _evaluate_bugs_for_preview(items)
            return success_response(data={"items": evaluated_items})

        raw_data = _load_json_file(BUG_RAW_FILE, {"items": []})
        raw_items = raw_data.get("items") if isinstance(raw_data, dict) else []
        if not isinstance(raw_items, list):
            raw_items = []
        vec_data = _load_json_file(VECTOR_DATA_FILE, {"items": []})
        vec_items = vec_data.get("items") if isinstance(vec_data, dict) else []
        if not isinstance(vec_items, list):
            vec_items = []
        imported = []
        for it in items:
            row = {
                "id": "B_" + uuid.uuid4().hex[:8],
                "bug_desc": it["bug_desc"],
                "severity": _normalize_severity(it.get("severity")),
                "frequency": _normalize_frequency(it.get("frequency")),
            }
            raw_items.append(row)
            imported.append(row)
            vec_items.append(
                {
                    "id": "V_" + uuid.uuid4().hex[:8],
                    "content": row["bug_desc"],
                    "type": "bug",
                    "board_type": "",
                    "ref_id": row["id"],
                }
            )
        _save_json_file(BUG_RAW_FILE, {"items": raw_items})
        _save_json_file(VECTOR_DATA_FILE, {"items": vec_items})
        return success_response(data={"imported_count": len(imported), "items": imported})
    except Exception as e:
        logger.exception("api_bug_import failed")
        return error_response(str(e), status_code=500)

# ==================== JIRA INTEGRATION ====================
@prd_audit_bp.route("/api/bug/import/jira", methods=["POST"])
def api_bug_import_jira():
    """从 Jira 直接拉取 Bug"""
    try:
        data = request.get_json() or {}
        jira_url = data.get("jira_url", "").strip().rstrip("/")
        username = data.get("username", "").strip()
        api_token = data.get("api_token", "").strip()
        jql = data.get("jql", "").strip()
        is_preview = data.get("preview", False)

        if not all([jira_url, username, api_token, jql]):
            return error_response("Jira 配置参数缺失 (jira_url, username, api_token, jql 均为必填)", status_code=400)

        import requests
        from requests.auth import HTTPBasicAuth

        # 构造 Jira 搜索请求
        search_url = f"{jira_url}/rest/api/2/search"
        payload = {
            "jql": jql,
            "maxResults": 100, # 默认拉取前100条
            "fields": ["summary", "description", "customfield_10004", "priority", "status"] # 可根据实际 jira 字段调整
        }
        
        headers = {"Accept": "application/json"}
        auth = HTTPBasicAuth(username, api_token)

        resp = requests.post(search_url, json=payload, headers=headers, auth=auth, timeout=30)
        if resp.status_code != 200:
            return error_response(f"Jira API 请求失败: {resp.status_code} - {resp.text}", status_code=500)
            
        jira_data = resp.json()
        issues = jira_data.get("issues", [])
        
        if not issues:
            return error_response("根据提供的 JQL 未查找到任何 Bug", status_code=404)

        items = []
        for issue in issues:
            fields = issue.get("fields", {})
            
            # 提取信息，做容错处理
            bug_desc = fields.get("summary") or ""
            if not bug_desc:
                continue
                
            # 严重程度通常是优先级或者自定义字段
            priority_obj = fields.get("priority")
            sev = priority_obj.get("name", "P2") if priority_obj else "P2"
            
            items.append({
                "bug_desc": f"[{issue.get('key', '')}] {bug_desc}",
                "severity": _normalize_severity(sev),
                "frequency": "中", # Jira里可能没有频率字段，默认给个中
            })

        if is_preview:
            evaluated_items = _evaluate_bugs_for_preview(items)
            return success_response(data={"items": evaluated_items})
            
        # 正式落库
        _ensure_bug_assets()
        raw_data = _load_json_file(BUG_RAW_FILE, {"items": []})
        raw_items = raw_data.get("items") if isinstance(raw_data, dict) else []
        if not isinstance(raw_items, list):
            raw_items = []
            
        vec_data = _load_json_file(VECTOR_DATA_FILE, {"items": []})
        vec_items = vec_data.get("items") if isinstance(vec_data, dict) else []
        if not isinstance(vec_items, list):
            vec_items = []
            
        imported = []
        for it in items:
            row = {
                "id": "B_" + uuid.uuid4().hex[:8],
                "bug_desc": it["bug_desc"],
                "severity": it.get("severity"),
                "frequency": it.get("frequency"),
            }
            raw_items.append(row)
            imported.append(row)
            vec_items.append(
                {
                    "id": "V_" + uuid.uuid4().hex[:8],
                    "content": row["bug_desc"],
                    "type": "bug",
                    "board_type": "",
                    "ref_id": row["id"],
                }
            )
            
        _save_json_file(BUG_RAW_FILE, {"items": raw_items})
        _save_json_file(VECTOR_DATA_FILE, {"items": vec_items})

        return success_response(data={"imported_count": len(imported), "items": imported})
        
    except Exception as e:
        logger.exception("api_bug_import_jira failed")
        return error_response(str(e), status_code=500)
# ==========================================================

@prd_audit_bp.route("/api/bug/patterns", methods=["GET", "POST"])
def api_bug_patterns():
    try:
        _ensure_bug_assets()
        if request.method == "GET":
            return success_response(data={"items": _load_bug_patterns()})
        data = request.get_json() or {}
        items = data.get("items") if isinstance(data, dict) else None
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)
        cleaned = []
        for i, p in enumerate(items):
            if not isinstance(p, dict):
                continue
            pid = str(p.get("pattern_id") or "").strip() or "P{:03d}".format(i + 1)
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            kws = [str(x).strip() for x in kws if str(x).strip()]
            if not kws:
                continue
            cleaned.append(
                {
                    "pattern_id": pid,
                    "keywords": kws,
                    "category": str(p.get("category") or "通用缺陷").strip(),
                    "design_gap": p.get("design_gap") if isinstance(p.get("design_gap"), list) else [],
                    "rule": str(p.get("rule") or "").strip() or "关键流程必须定义异常兜底策略",
                    "weight": float(p.get("weight") or 0.5),
                    "enabled": bool(p.get("enabled", True)),
                }
            )
        _save_json_file(BUG_PATTERN_FILE, {"items": cleaned})
        return success_response(data={"items": cleaned, "count": len(cleaned)})
    except Exception as e:
        logger.exception("api_bug_patterns failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/patterns/auto_from_bugs", methods=["POST"])
def api_bug_patterns_auto_from_bugs():
    try:
        _ensure_bug_assets()
        data = request.get_json() or {}
        items = data.get("items") if isinstance(data, dict) else None
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)

        bug_descs = []
        for it in items:
            if not isinstance(it, dict):
                continue
            d = str(it.get("bug_desc") or "").strip()
            if d:
                bug_descs.append(d)
        if not bug_descs:
            return error_response("未提供可用于生成模式的 bug_desc", status_code=400)

        existing = _load_bug_patterns()
        used_keywords = set()
        max_pid = 0
        for p in existing:
            if not isinstance(p, dict):
                continue
            pid = str(p.get("pattern_id") or "")
            m = re.match(r"P(\d+)$", pid)
            if m:
                max_pid = max(max_pid, int(m.group(1)))
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            for k in kws:
                ks = str(k).strip()
                if ks:
                    used_keywords.add(ks)

        stopwords = {
            "问题", "这个", "那个", "无法", "不能", "没有", "出现", "导致", "以及",
            "时候", "之后", "过程", "进行", "相关", "功能", "页面", "系统", "当前",
            "仍然", "展示", "可以", "一下", "或者", "然后", "其中", "这里", "那里",
            "用户", "业务", "场景", "模块", "逻辑", "状态", "操作", "流程",
            "投屏", "广告", "语音", "自助", "商k", "k歌", "导唱", "打分", "录音", "量贩",
            "第三屏", "竖版",
        }
        risk_terms = [
            "异常", "失败", "报错", "崩溃", "闪退", "卡死", "无响应", "黑屏",
            "不一致", "不同步", "错乱", "冲突", "越权", "超时", "丢失", "重叠",
            "不消失", "不能", "无法", "无反馈", "无应答",
        ]
        keyword_counter = Counter()

        for desc in bug_descs:
            parts = re.split(r"[，,。；;、\s]+", desc)
            candidates = []
            for part in parts:
                s = str(part).strip()
                if not s:
                    continue
                if len(s) < 2 or len(s) > 12:
                    continue
                if s in stopwords:
                    continue
                s_lower = s.lower()
                if s_lower in {"pk", "x9", "vx", "bug", "p0", "p1", "p2"}:
                    continue
                if re.fullmatch(r"[a-z]{1,3}\d{1,4}", s_lower):
                    continue
                if re.fullmatch(r"[a-z]{1,4}-\d{1,6}", s_lower):
                    continue
                if re.fullmatch(r"p\d{1,2}", s_lower):
                    continue
                if re.fullmatch(r"[a-z]{1,2}", s_lower):
                    continue
                if not any(t in s for t in risk_terms):
                    continue
                candidates.append(s)
            for kw in candidates[:3]:
                keyword_counter[kw] += 1

        top_keywords = [kw for kw, _ in keyword_counter.most_common(20)]
        added = []
        for kw in top_keywords:
            if kw in used_keywords:
                continue
            related = [d for d in bug_descs if kw in d]
            if not related:
                continue
            text_all = "\n".join(related)
            if any(x in text_all for x in ["唤醒", "语音", "识别", "应答"]):
                category = "语音交互"
                design_gap = ["状态流转缺失", "反馈机制缺失"]
                rule = "语音交互必须定义唤醒-识别-反馈闭环与异常兜底"
                weight = 0.9
            elif any(x in text_all for x in ["卡死", "无响应", "崩溃", "闪退"]):
                category = "稳定性"
                design_gap = ["超时恢复缺失", "异常处理缺失"]
                rule = "关键交互必须定义超时恢复、重试与降级策略"
                weight = 0.88
            elif any(x in text_all for x in ["不一致", "不同步", "错乱"]):
                category = "状态一致性"
                design_gap = ["状态机缺失"]
                rule = "多模块交互必须定义状态一致性与冲突裁决规则"
                weight = 0.86
            else:
                category = "通用缺陷"
                design_gap = ["异常处理缺失"]
                rule = "关键流程必须定义异常处理与回退策略"
                weight = 0.72

            max_pid += 1
            row = {
                "pattern_id": "P{:03d}".format(max_pid),
                "keywords": [kw],
                "category": category,
                "design_gap": design_gap,
                "rule": rule,
                "weight": weight,
                "enabled": True,
            }
            existing.append(row)
            added.append(row)
            used_keywords.add(kw)
            if len(added) >= 12:
                break

        _save_json_file(BUG_PATTERN_FILE, {"items": existing})
        return success_response(data={"added_count": len(added), "items": added})
    except Exception as e:
        logger.exception("api_bug_patterns_auto_from_bugs failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/analyze", methods=["POST"])
def api_bug_analyze():
    try:
        _ensure_bug_assets()
        data = request.get_json() or {}
        bug_inputs = []
        if isinstance(data.get("bugs"), list):
            for x in data.get("bugs"):
                if isinstance(x, dict):
                    d = str(x.get("bug_desc") or "").strip()
                else:
                    d = str(x or "").strip()
                if d:
                    bug_inputs.append(d)
        one = str(data.get("bug_desc") or "").strip()
        if one:
            bug_inputs.append(one)
        if not bug_inputs:
            return error_response("bug_desc 不能为空", status_code=400)

        analyses = []
        rules_data = _load_json_file(BUG_RULE_FILE, {"items": []})
        rules = rules_data.get("items") if isinstance(rules_data, dict) else []
        if not isinstance(rules, list):
            rules = []
        for d in bug_inputs:
            parsed = _analyze_bug_desc(d)
            if not parsed:
                continue
            analyses.append({"bug_desc": d, **parsed})
            rid = _next_rule_id(rules)
            rules.append(
                {
                    "rule_id": rid,
                    "rule_desc": parsed.get("rule") or "",
                    "source_pattern": parsed.get("pattern_id") or "",
                    "severity": "P0" if float(parsed.get("weight") or 0) >= 0.9 else ("P1" if float(parsed.get("weight") or 0) >= 0.75 else "P2"),
                }
            )
        _save_json_file(BUG_RULE_FILE, {"items": rules})
        return success_response(data={"items": analyses, "generated_rules": rules[-len(analyses):] if analyses else []})
    except Exception as e:
        logger.exception("api_bug_analyze failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/rules", methods=["GET", "POST"])
def api_bug_rules():
    try:
        _ensure_bug_assets()
        if request.method == "POST":
            data = request.get_json() or {}
            items = data.get("items") if isinstance(data, dict) else None
            if not isinstance(items, list):
                return error_response("items 必须为数组", status_code=400)
            cleaned = []
            for i, it in enumerate(items):
                if not isinstance(it, dict):
                    continue
                rid = str(it.get("rule_id") or "").strip() or "RBUG_{:03d}".format(i + 1)
                cleaned.append(
                    {
                        "rule_id": rid,
                        "category": str(it.get("category") or "General").strip() or "General",
                        "rule_desc": str(it.get("rule_desc") or "").strip() or "关键流程需补充规则",
                        "source_pattern": str(it.get("source_pattern") or "").strip(),
                        "severity": _normalize_severity(it.get("severity")),
                        "detect": str(it.get("detect") or "").strip(),
                        "risk": str(it.get("risk") or "").strip(),
                    }
                )
            _save_json_file(BUG_RULE_FILE, {"items": cleaned})
            return success_response(data={"items": cleaned, "count": len(cleaned)}, message="规则库已保存")
        data = _load_json_file(BUG_RULE_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        return success_response(data={"items": items, "count": len(items)})
    except Exception as e:
        logger.exception("api_bug_rules failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/rules/export", methods=["GET"])
def api_bug_rules_export():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RULE_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_rules.json",
        )
    except Exception as e:
        logger.exception("api_bug_rules_export failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/patterns/export", methods=["GET"])
def api_bug_patterns_export():
    try:
        _ensure_bug_assets()
        items = _load_bug_patterns()
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_patterns.json",
        )
    except Exception as e:
        logger.exception("api_bug_patterns_export failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/raw/export", methods=["GET"])
def api_bug_raw_export():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RAW_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_raw.json",
        )
    except Exception as e:
        logger.exception("api_bug_raw_export failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/import_template_csv", methods=["GET"])
def api_bug_import_template_csv():
    try:
        sio = io.StringIO()
        writer = csv.writer(sio)
        writer.writerow(["bug_desc", "severity", "frequency"])
        writer.writerow(["投屏过程中黑屏", "P0", "高"])
        writer.writerow(["广告切换后卡死", "P1", "中"])
        writer.writerow(["弱网下播放无响应", "P1", "中"])
        writer.writerow(["切歌后偶发无声", "P2", "低"])
        content = sio.getvalue().encode("utf-8-sig")
        buf = io.BytesIO(content)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="text/csv",
            as_attachment=True,
            download_name="bug_import_template.csv",
        )
    except Exception as e:
        logger.exception("api_bug_import_template_csv failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/bug/raw", methods=["GET"])
def api_bug_raw():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RAW_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        return success_response(data={"items": items, "count": len(items)})
    except Exception as e:
        logger.exception("api_bug_raw failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/vector/search", methods=["POST"])
def api_vector_search():
    try:
        data = request.get_json() or {}
        query = str(data.get("query") or "").strip()
        if not query:
            return error_response("query 不能为空", status_code=400)
        top_k = int(data.get("top_k") or 5)
        top_k = 1 if top_k < 1 else (20 if top_k > 20 else top_k)
        matches = _vector_search_local(
            query=query,
            top_k=top_k,
            type_filter=(data.get("type") or None),
            board_type=(data.get("board_type") or None),
        )
        return success_response(data={"query": query, "items": matches})
    except Exception as e:
        logger.exception("api_vector_search failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/prd/audit", methods=["POST"])
def api_prd_audit():
    try:
        data = request.get_json() or {}
        prd_text = str(data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        config_path = _pick_llm_config_path()
        use_llm = bool(data.get("use_llm", False))
        if not use_llm:
            config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

        last_obj = {}
        for chunk in pipeline.run_prd_audit_stream(prd_text, llm_config_path=config_path, timeout=180):
            try:
                obj = json.loads(chunk)
            except Exception:
                continue
            if isinstance(obj, dict) and "L3" in obj:
                last_obj = obj
        if not last_obj:
            return error_response("审计失败，未生成结果", status_code=500)

        patterns = _load_bug_patterns()
        bug_hits = []
        for p in patterns:
            if not isinstance(p, dict) or p.get("enabled") is False:
                continue
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            matched = [k for k in kws if str(k) and str(k) in prd_text]
            if matched:
                bug_hits.append(
                    {
                        "pattern_id": p.get("pattern_id"),
                        "category": p.get("category"),
                        "matched_keywords": matched,
                        "rule": p.get("rule"),
                        "weight": p.get("weight"),
                    }
                )
        bug_hits.sort(key=lambda x: float(x.get("weight") or 0), reverse=True)
        
        # 本地规则匹配：知识库风险提示（不需要大模型）
        knowledge_matches = _match_knowledge_base(prd_text, top_k=5)
        
        vector_hits = _vector_search_local(query=prd_text, top_k=8) if use_llm else []
        return success_response(
            data={
                "audit_result": last_obj,
                "bug_hits": bug_hits[:12],
                "knowledge_matches": knowledge_matches,  # 新增：知识库匹配结果
                "vector_hits": vector_hits,
            }
        )
    except Exception as e:
        logger.exception("api_prd_audit failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/multi_prd/audit", methods=["POST"])
def api_multi_prd_audit():
    """多 PRD 关联审计：批量解析 + 跨文档一致性检查（术语/API 对齐）"""
    try:
        data = request.get_json() or {}
        docs = data.get("docs") or []
        if not isinstance(docs, list) or not docs:
            return error_response("docs 不能为空", status_code=400)

        use_llm = bool(data.get("use_llm", False))
        config_path = _pick_llm_config_path()
        if not use_llm:
            config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

        from .system_model import extract_prd_structure
        import difflib

        def _clean_list(items: Any) -> List[str]:
            if not isinstance(items, list):
                return []
            out = []
            for x in items:
                s = str(x or "").strip()
                if not s or s == "【PRD未说明】":
                    continue
                out.append(s)
            return out

        def _norm_term(s: str) -> str:
            t = (s or "").strip()
            t = re.sub(r"\s+", " ", t)
            t = t.strip(" \t\r\n:：-—_（）()[]【】<>《》\"'`")
            return t.lower()

        def _extract_endpoints(text: str) -> List[Dict[str, str]]:
            raw = (text or "").replace("\r\n", "\n").replace("\r", "\n")
            eps = []
            # 常见格式：GET /api/xxx、POST /v1/xxx
            for m in re.finditer(r"\b(GET|POST|PUT|DELETE|PATCH|HEAD|OPTIONS)\s+(/[A-Za-z0-9_\-./{}]+)", raw, flags=re.IGNORECASE):
                method = m.group(1).upper()
                path = m.group(2)
                eps.append({"method": method, "path": path})
            # 兜底：/api/xxx 形式（无 method）
            for m in re.finditer(r"\b(/api/[A-Za-z0-9_\-./{}]+)\b", raw):
                path = m.group(1)
                eps.append({"method": "", "path": path})
            # 去重
            seen = set()
            out = []
            for e in eps:
                key = (e.get("method") or "", e.get("path") or "")
                if key in seen:
                    continue
                seen.add(key)
                out.append(e)
            return out

        parsed_docs = []
        all_terms = {}  # norm_term -> {raw_samples:set, docs:set}
        all_eps = {}    # (method,path) -> {docs:set}

        for idx, doc in enumerate(docs[:20]):
            if not isinstance(doc, dict):
                continue
            name = str(doc.get("name") or f"Doc{idx+1}").strip()[:80]
            content = str(doc.get("content") or "").strip()
            if not content:
                continue
            stage1 = extract_prd_structure(content, llm_config_path=config_path, timeout=90)
            fields = {
                "product_name": str(stage1.get("product_name") or "").strip(),
                "modules": _clean_list(stage1.get("modules")),
                "features": _clean_list(stage1.get("features")),
                "user_roles": _clean_list(stage1.get("user_roles")),
                "flows": _clean_list(stage1.get("flows")),
                "states": _clean_list(stage1.get("states")),
                "business_rules": _clean_list(stage1.get("business_rules")),
                "data_structures": _clean_list(stage1.get("data_structures")),
                "permissions": _clean_list(stage1.get("permissions")),
                "exceptions": _clean_list(stage1.get("exceptions")),
                "edge_cases": _clean_list(stage1.get("edge_cases")),
                "dependencies": _clean_list(stage1.get("dependencies")),
                "non_functional_requirements": _clean_list(stage1.get("non_functional_requirements")),
                "success_metrics": _clean_list(stage1.get("success_metrics")),
            }

            terms_raw = []
            for k in ("modules", "features", "user_roles", "states", "data_structures", "business_rules", "dependencies"):
                terms_raw.extend(fields.get(k) or [])
            term_norms = []
            for t in terms_raw:
                n = _norm_term(t)
                if not n:
                    continue
                term_norms.append(n)
                bucket = all_terms.get(n)
                if not bucket:
                    bucket = {"raw_samples": set(), "docs": set()}
                    all_terms[n] = bucket
                bucket["raw_samples"].add(t)
                bucket["docs"].add(name)

            endpoints = _extract_endpoints(content)
            for e in endpoints:
                key = (e.get("method") or "", e.get("path") or "")
                b = all_eps.get(key)
                if not b:
                    b = {"docs": set()}
                    all_eps[key] = b
                b["docs"].add(name)

            parsed_docs.append({
                "name": name,
                "content_chars": len(content),
                "stage1": fields,
                "term_count": len(set(term_norms)),
                "endpoints": endpoints,
            })

        if not parsed_docs:
            return error_response("未解析到有效文档", status_code=400)

        # 术语近似重复（相似度阈值）
        term_keys = list(all_terms.keys())
        similar_pairs = []
        for i in range(len(term_keys)):
            a = term_keys[i]
            if len(a) < 3:
                continue
            for j in range(i + 1, min(len(term_keys), i + 200)):
                b = term_keys[j]
                if len(b) < 3:
                    continue
                if abs(len(a) - len(b)) > 8:
                    continue
                ratio = difflib.SequenceMatcher(None, a, b).ratio()
                if ratio >= 0.88 and a != b:
                    da = sorted(list(all_terms[a]["docs"]))
                    db = sorted(list(all_terms[b]["docs"]))
                    similar_pairs.append({
                        "a": a, "b": b, "ratio": round(ratio, 3),
                        "docs_a": da[:6], "docs_b": db[:6],
                    })
        similar_pairs.sort(key=lambda x: (-x["ratio"], x["a"], x["b"]))
        similar_pairs = similar_pairs[:30]

        # Endpoint 对齐：找出在部分文档缺失的接口
        eps_rows = []
        doc_names = [d["name"] for d in parsed_docs]
        for (method, path), b in all_eps.items():
            present = b.get("docs") or set()
            missing = [n for n in doc_names if n not in present]
            if missing and len(missing) != len(doc_names):
                eps_rows.append({
                    "method": method or "-",
                    "path": path,
                    "present": sorted(list(present)),
                    "missing": missing,
                })
        eps_rows.sort(key=lambda x: (x["path"], x["method"]))
        eps_rows = eps_rows[:80]

        # 输出 Markdown 报告
        lines = []
        lines.append("# 多 PRD 关联审计报告")
        lines.append("")
        lines.append(f"- 文档数：{len(parsed_docs)}")
        lines.append(f"- 术语候选总数（去重）：{len(all_terms)}")
        lines.append(f"- Endpoint 候选总数（去重）：{len(all_eps)}")
        lines.append("")
        lines.append("## 1. 文档概览")
        lines.append("")
        lines.append("| 文档 | 字符数 | 模块数 | 流程数 | 状态数 | 规则数 | 端点数 |")
        lines.append("|---|---:|---:|---:|---:|---:|---:|")
        for d in parsed_docs:
            s = d.get("stage1") or {}
            lines.append("| {} | {} | {} | {} | {} | {} | {} |".format(
                d["name"],
                d["content_chars"],
                len(s.get("modules") or []),
                len(s.get("flows") or []),
                len(s.get("states") or []),
                len(s.get("business_rules") or []),
                len(d.get("endpoints") or []),
            ))
        lines.append("")

        lines.append("## 2. 跨文档接口对齐（缺失检查）")
        lines.append("")
        if not eps_rows:
            lines.append("- 未检测到明显的接口在部分文档缺失（或文档中未出现可识别的接口片段）。")
        else:
            lines.append("| 方法 | 路径 | 出现于 | 缺失于 |")
            lines.append("|---|---|---|---|")
            for r in eps_rows:
                lines.append("| {} | {} | {} | {} |".format(
                    r["method"],
                    r["path"],
                    "、".join(r["present"][:5]) + ("…" if len(r["present"]) > 5 else ""),
                    "、".join(r["missing"][:5]) + ("…" if len(r["missing"]) > 5 else ""),
                ))
        lines.append("")

        lines.append("## 3. 术语一致性（疑似同义/拼写不一致）")
        lines.append("")
        if not similar_pairs:
            lines.append("- 未检测到明显的相似术语对（阈值 0.88）。")
        else:
            lines.append("| 术语A | 术语B | 相似度 | A出现于 | B出现于 |")
            lines.append("|---|---|---:|---|---|")
            for p in similar_pairs:
                lines.append("| {} | {} | {} | {} | {} |".format(
                    p["a"],
                    p["b"],
                    p["ratio"],
                    "、".join(p["docs_a"]) or "-",
                    "、".join(p["docs_b"]) or "-",
                ))
        lines.append("")

        lines.append("## 4. 建议下一步（落地动作）")
        lines.append("")
        lines.append("- 建立统一“术语表”（模块/角色/状态/数据结构），将本报告第 3 节的疑似同义项合并为单一标准写法。")
        lines.append("- 建立统一“接口目录”（method+path），并要求前端/后端/数据 PRD 同步更新；本报告第 2 节可作为缺失清单。")
        lines.append("- 若需要更高精度的对齐（字段级、状态机级、一致性证明），建议进一步把 Stage1 的 blocks/source_map 用于精确锚点回溯。")
        lines.append("")

        report_md = "\n".join(lines)
        return success_response(data={"report_md": report_md, "docs": parsed_docs})
    except Exception as e:
        logger.exception("api_multi_prd_audit failed")
        return error_response(str(e), status_code=500)

# ---------- API：解析 PDF / Word ----------

@prd_audit_bp.route("/api/parse_pdf", methods=["POST"])
def api_parse_pdf():
    """解析 PDF 文件"""
    try:
        if "file" not in request.files:
            return error_response("未上传文件", status_code=400)
        file = request.files["file"]
        if not file.filename:
            return error_response("文件名为空", status_code=400)
        
        # 安全性增强：限制文件大小（例如 10MB）
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 10 * 1024 * 1024:
            return error_response("文件过大（超过 10MB），请拆分后上传或直接粘贴文本", status_code=400)
            
        if not file.filename.lower().endswith(".pdf"):
            return error_response("请上传 PDF 文件", status_code=400)

        from io import BytesIO
        file_data = file.read()
        file_stream = BytesIO(file_data)
        text_chunks = []
        page_errors = 0
        try:
            import pypdf
            reader = pypdf.PdfReader(file_stream, strict=False)
            for page in reader.pages:
                try:
                    text_chunks.append(page.extract_text() or "")
                except Exception:
                    page_errors += 1
        except ImportError:
            try:
                import PyPDF2
                file_stream.seek(0)
                reader = PyPDF2.PdfReader(file_stream)
                for page in reader.pages:
                    try:
                        text_chunks.append(page.extract_text() or "")
                    except Exception:
                        page_errors += 1
            except ImportError:
                return error_response("未安装 PDF 解析库 (pypdf 或 PyPDF2)。请联系管理员: pip install pypdf", status_code=500)
        except Exception as e:
            logger.error("PDF解析错误: %s", e)
            return error_response("PDF 解析失败: " + str(e), status_code=500)

        text = "\n".join([c for c in text_chunks if isinstance(c, str) and c.strip()]).strip()
        if text:
            data = {"text": text}
            if page_errors > 0:
                data["warning"] = f"部分页面解析失败（{page_errors}页），已返回可提取内容。"
            return success_response(data=data)
        fallback_hint = (
            "【PDF解析提示】该PDF可能是扫描件/图片版，未提取到可复制文字。\n"
            "建议：1）导出可复制文本的PDF；2）上传Word(.docx)；3）直接粘贴PRD正文。"
        )
        return success_response(data={"text": fallback_hint, "warning": "未提取到可复制文本，已返回导入提示。"})
    except Exception as e:
        logger.exception("PDF 上传解析失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/parse_docx", methods=["POST"])
def api_parse_docx():
    """解析 Word (.docx) 文件"""
    try:
        if "file" not in request.files:
            return error_response("未上传文件", status_code=400)
        file = request.files["file"]
        if not file.filename:
            return error_response("文件名为空", status_code=400)
            
        # 安全性增强：限制文件大小（例如 10MB）
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
        if size > 10 * 1024 * 1024:
            return error_response("Word 文件过大（超过 10MB），请拆分后上传或直接粘贴文本", status_code=400)

        if not file.filename.lower().endswith(".docx"):
            return error_response("请上传 .docx 文件", status_code=400)
        try:
            from docx import Document
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        except ImportError:
            return error_response("未安装 python-docx。请联系管理员: pip install python-docx", status_code=500)
        except Exception as e:
            logger.error("DOCX 解析错误: %s", e)
            return error_response("Word 解析失败: " + str(e), status_code=500)
        if not text.strip():
            return error_response("未从 Word 中提取到文本", status_code=400)
        return success_response(data={"text": text})
    except Exception as e:
        logger.exception("Word 上传解析失败")
        return error_response(str(e), status_code=500)

# ---------- API：导出报告 Word ----------

@prd_audit_bp.route("/api/export_report_docx", methods=["POST"])
def api_export_report_docx():
    """将评审报告 Markdown 或 stage3_json 转为 Word 并返回文件"""
    try:
        data = request.get_json() or {}
        content = (data.get("content") or "").strip()
        stage3_json = data.get("stage3_json") if isinstance(data.get("stage3_json"), dict) else None
        if not content and not stage3_json:
            return error_response("报告内容为空", status_code=400)
        try:
            from docx import Document
            import re

            def _clean(s: str) -> str:
                """移除不兼容 XML 的控制字符，避免 python-docx 报错。"""
                if not isinstance(s, str):
                    s = str(s)
                # 保留常规可见字符与换行，去掉 0x00-0x08,0x0b-0x0c,0x0e-0x1f
                return re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", "", s)

            doc = Document()
            if stage3_json and isinstance(stage3_json, dict):
                title = _clean(stage3_json.get("report_title") or "【审计报告】PRD：工具扫描+人工复核版")
            else:
                title = "PRD 评审报告"
            doc.add_heading(title, 0)
            if stage3_json:
                summary = stage3_json.get("summary") or {}
                defects = stage3_json.get("defects") or []
                core_summary = stage3_json.get("core_risk_summary") or {}
                merged_issues = stage3_json.get("merged_issues") or []
                doc.add_heading("一、总体结论", level=1)
                doc.add_paragraph(_clean(f"质量评分：{summary.get('quality_score', '【PRD未说明】')}/10"))
                doc.add_paragraph(_clean(f"风险等级：{summary.get('risk_level', '【PRD未说明】')}"))
                doc.add_paragraph(_clean(f"主要问题：{summary.get('main_problem', '【PRD未说明】')}"))
                doc.add_heading("核心风险摘要", level=2)
                doc.add_paragraph(_clean(f"一句话总结：{core_summary.get('one_liner', '【PRD未说明】')}"))
                for item in pipeline._ensure_list(core_summary.get("top3")):
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("二、漏洞与风险清单", level=1)
                stats = stage3_json.get("scan_stats") or {}
                doc.add_paragraph(
                    _clean(
                        f"扫描来源统计：规则库 {stats.get('rule', 0)} 条，LLM {stats.get('llm', 0)} 条，混合 {stats.get('hybrid', 0)} 条"
                    )
                )
                if isinstance(merged_issues, list) and merged_issues:
                    doc.add_heading("合并后的核心问题", level=2)
                    for i, m in enumerate(merged_issues, start=1):
                        doc.add_heading(
                            _clean(f"核心问题{i}：{m.get('name', '【PRD未说明】')}（{m.get('risk_level', 'P2')}）"),
                            level=3,
                        )
                        doc.add_paragraph(
                            _clean(
                                f"涉及锚点：{'; '.join(pipeline._ensure_list(m.get('anchors'))) or '【PRD未说明】'}"
                            )
                        )
                        doc.add_paragraph(_clean(f"问题描述：{m.get('description', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"风险分析：{m.get('reason', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"建议：{m.get('suggestion', '【PRD未说明】')}"))
                if isinstance(defects, list) and defects:
                    doc.add_heading("详细漏洞清单", level=2)
                    for i, d in enumerate(defects, start=1):
                        doc.add_heading(_clean(f"漏洞{i}"), level=2)
                        doc.add_paragraph(_clean(f"模块：{d.get('module', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"类型：{d.get('type', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"风险等级：{d.get('risk_level', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"来源：{d.get('source', 'llm')}"))
                        doc.add_paragraph(
                            _clean(f"锚点：{d.get('anchor', d.get('module', '【PRD未说明】'))}")
                        )
                        doc.add_paragraph(_clean(f"描述：{d.get('description', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"原因：{d.get('reason', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"建议：{d.get('suggestion', '【PRD未说明】')}"))
                else:
                    doc.add_paragraph("未发现漏洞")
                doc.add_heading("三、测试重点", level=1)
                for item in pipeline._ensure_list(stage3_json.get("test_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("四、研发重点", level=1)
                for item in pipeline._ensure_list(stage3_json.get("dev_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("五、项目风险", level=1)
                for item in pipeline._ensure_list(stage3_json.get("risks")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("六、计划建议", level=1)
                for item in pipeline._ensure_list(stage3_json.get("plan")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
            else:
                lines = content.replace("\r\n", "\n").split("\n")
                i = 0
                while i < len(lines):
                    line = _clean(lines[i])
                    stripped = line.strip()
                    if stripped.startswith("## "):
                        title = stripped[3:].strip()
                        doc.add_heading(title, level=1)
                    elif stripped.startswith("### "):
                        title = stripped[4:].strip()
                        doc.add_heading(title, level=2)
                    elif stripped:
                        doc.add_paragraph(line.rstrip())
                    else:
                        doc.add_paragraph()
                    i += 1
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(
                buf,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name="PRD评审报告.docx",
            )
        except ImportError:
            return error_response("未安装 python-docx。请联系管理员: pip install python-docx", status_code=500)
        except Exception as e:
            logger.exception("导出 Word 失败: %s", e)
            return error_response("导出 Word 失败: " + str(e), status_code=500)
    except Exception as e:
        logger.exception("export_report_docx 失败")
        return error_response(str(e), status_code=500)

# ---------- API：默认提示词、LLM 配置 ----------

@prd_audit_bp.route("/api/default_prd_prompt", methods=["GET"])
def api_get_default_prd_prompt():
    """返回 PRD 默认审计提示词（从 prd_audit 目录的 default 文件读取）"""
    try:
        if os.path.exists(DEFAULT_PRD_AUDIT_PROMPT_FILE):
            with open(DEFAULT_PRD_AUDIT_PROMPT_FILE, "r", encoding="utf-8") as f:
                s = f.read().strip()
                if s and len(s) > 100:
                    return success_response(data={"prompt": s})
        return success_response(data={"prompt": FALLBACK_PRD_PROMPT})
    except Exception as e:
        logger.warning("Failed to read default PRD prompt: %s", e)
        return success_response(data={"prompt": FALLBACK_PRD_PROMPT})

@prd_audit_bp.route("/api/audit_stats", methods=["GET"])
def api_audit_stats():
    """获取所有历史审计的统计数据，用于 Dashboard 大盘展示"""
    try:
        from .audit_learning import load_all_snapshots

        snapshots = load_all_snapshots(limit=5000)
        stats = []
        category_counts = {}

        for snap in snapshots:
            if not isinstance(snap, dict):
                continue
            ts = int(snap.get("created_at") or 0)
            extras = snap.get("extras") if isinstance(snap.get("extras"), dict) else {}
            summary = extras.get("summary") if isinstance(extras.get("summary"), dict) else {}
            guardrail = extras.get("guardrail") if isinstance(extras.get("guardrail"), dict) else {}
            guardrail_score = 0
            try:
                if guardrail.get("score") is not None:
                    guardrail_score = int(float(guardrail.get("score")))
            except Exception:
                guardrail_score = 0
            score = summary.get("score", 0)
            try:
                score = float(score)
            except Exception:
                score = 0

            stage2 = snap.get("stage2_output") if isinstance(snap.get("stage2_output"), dict) else {}
            defects = stage2.get("defects") if isinstance(stage2.get("defects"), list) else []
            p0 = sum(1 for d in defects if isinstance(d, dict) and str(d.get("risk_level") or "").upper() == "P0")
            p1 = sum(1 for d in defects if isinstance(d, dict) and str(d.get("risk_level") or "").upper() == "P1")
            p2 = sum(1 for d in defects if isinstance(d, dict) and str(d.get("risk_level") or "").upper() == "P2")

            stats.append({
                "timestamp": ts,
                "score": score,
                "p0": p0,
                "p1": p1,
                "p2": p2,
                "guardrail_score": guardrail_score,
                "guardrail_score10": round(float(guardrail_score) / 10.0, 1),
            })

            for d in defects:
                if not isinstance(d, dict):
                    continue
                cat = d.get("category") or "其他"
                category_counts[cat] = category_counts.get(cat, 0) + 1

        stats.sort(key=lambda x: x["timestamp"])
        return success_response(data={"trends": stats, "distribution": category_counts, "total_audits": len(stats)})
    except Exception as e:
        logger.exception("获取审计统计失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/evaluate", methods=["POST"])
def api_guardrail_evaluate():
    try:
        data = request.get_json() or {}
        from .guardrail_engine import evaluate_guardrail

        prd_text = str(data.get("prd_text") or data.get("content") or "").strip()
        report_md = str(data.get("report_md") or "").strip()
        stage1_output = data.get("stage1_output") if isinstance(data.get("stage1_output"), dict) else {}
        stage2_output = data.get("stage2_output") if isinstance(data.get("stage2_output"), dict) else {}
        test_cases = data.get("test_cases") if isinstance(data.get("test_cases"), list) else []

        res = evaluate_guardrail(
            prd_text=prd_text,
            stage1_output=stage1_output,
            stage2_output=stage2_output,
            report_md=report_md,
            test_cases=test_cases,
        )
        return success_response(data=res)
    except Exception as e:
        logger.exception("guardrail evaluate failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/adversarial_cases", methods=["GET"])
def api_guardrail_adversarial_cases_get():
    try:
        from .audit_learning import load_adversarial_cases
        return success_response(data=load_adversarial_cases())
    except Exception as e:
        logger.exception("load adversarial cases failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/adversarial_cases", methods=["POST"])
def api_guardrail_adversarial_cases_save():
    try:
        from .audit_learning import save_adversarial_cases
        data = request.get_json() or {}
        items = data.get("items")
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)
        cleaned = []
        for it in items[:200]:
            if not isinstance(it, dict):
                continue
            cid = str(it.get("case_id") or "").strip()
            prd_text = str(it.get("prd_text") or "").strip()
            if not cid or not prd_text:
                continue
            cleaned.append({
                "case_id": cid[:64],
                "name": str(it.get("name") or cid).strip()[:120],
                "kind": str(it.get("kind") or "normal").strip()[:40],
                "use_llm": bool(it.get("use_llm", True)),
                "min_guardrail_score": int(float(it.get("min_guardrail_score") or 70)),
                "prd_text": prd_text[:200000],
            })
        save_adversarial_cases({"items": cleaned})
        return success_response(message="已保存对抗回归集", data={"count": len(cleaned)})
    except Exception as e:
        logger.exception("save adversarial cases failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/regression/run", methods=["POST"])
def api_guardrail_regression_run():
    try:
        data = request.get_json() or {}
        job_id = _job_queue.enqueue("guardrail_regression", {
            "timeout": int(data.get("timeout") or 120),
            "use_llm": bool(data.get("use_llm", True)),
            "case_ids": data.get("case_ids") if isinstance(data.get("case_ids"), list) else [],
        })
        return success_response(data={"job_id": job_id}, message="已加入回归队列")
    except Exception as e:
        logger.exception("guardrail regression run failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/regression/latest", methods=["GET"])
def api_guardrail_regression_latest():
    try:
        from .audit_learning import load_latest_regression_run
        return success_response(data=load_latest_regression_run())
    except Exception as e:
        logger.exception("guardrail regression latest failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/guardrail/incidents", methods=["GET"])
def api_guardrail_incidents():
    try:
        from .audit_learning import load_incident_samples
        limit = int(request.args.get("limit") or 50)
        return success_response(data={"items": load_incident_samples(limit=limit)})
    except Exception as e:
        logger.exception("guardrail incidents failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/llm_metrics", methods=["GET"])
def api_llm_metrics():
    """读取 LLM 调用指标（jsonl），并输出聚合后的统计数据"""
    try:
        since_days = int(request.args.get("since_days") or 30)
        max_lines = int(request.args.get("max_lines") or 20000)
        now_ts = int(time.time())
        since_ts = now_ts - max(0, since_days) * 86400

        metrics_path = os.path.join(os.path.dirname(__file__), "learning_repo", "llm_metrics.jsonl")
        if not os.path.exists(metrics_path):
            return success_response(data={
                "window": {"since_days": since_days, "since_ts": since_ts, "now_ts": now_ts},
                "total_calls": 0,
                "ok_calls": 0,
                "fail_calls": 0,
                "avg_duration_ms": 0,
                "token_total": 0,
                "by_day": [],
                "by_provider": {},
            })

        def _extract_total_tokens(ev: Dict[str, Any]) -> int:
            usage = ev.get("usage")
            if isinstance(usage, dict):
                v = usage.get("total_tokens")
                if isinstance(v, int):
                    return v
                v = usage.get("totalTokenCount")
                if isinstance(v, int):
                    return v
                v = usage.get("totalTokens")
                if isinstance(v, int):
                    return v
            return 0

        total_calls = 0
        ok_calls = 0
        fail_calls = 0
        sum_ms = 0
        token_total = 0
        by_day: Dict[str, Dict[str, Any]] = {}
        by_provider: Dict[str, Dict[str, Any]] = {}

        with open(metrics_path, "r", encoding="utf-8") as f:
            for i, line in enumerate(f):
                if i >= max(0, max_lines):
                    break
                raw = (line or "").strip()
                if not raw:
                    continue
                try:
                    ev = json.loads(raw)
                except Exception:
                    continue
                if not isinstance(ev, dict):
                    continue
                ts = int(ev.get("ts") or 0)
                if ts and ts < since_ts:
                    continue
                provider = str(ev.get("provider") or "unknown")
                ok = bool(ev.get("ok"))
                dur = int(ev.get("duration_ms") or 0)
                tok = _extract_total_tokens(ev)

                total_calls += 1
                ok_calls += 1 if ok else 0
                fail_calls += 0 if ok else 1
                sum_ms += dur
                token_total += tok

                day = time.strftime("%Y-%m-%d", time.localtime(ts or now_ts))
                d = by_day.get(day)
                if not d:
                    d = {"date": day, "calls": 0, "fails": 0, "duration_ms": 0, "token_total": 0}
                    by_day[day] = d
                d["calls"] += 1
                d["fails"] += 0 if ok else 1
                d["duration_ms"] += dur
                d["token_total"] += tok

                p = by_provider.get(provider)
                if not p:
                    p = {"calls": 0, "fails": 0, "duration_ms": 0, "token_total": 0}
                    by_provider[provider] = p
                p["calls"] += 1
                p["fails"] += 0 if ok else 1
                p["duration_ms"] += dur
                p["token_total"] += tok

        by_day_list = list(by_day.values())
        by_day_list.sort(key=lambda x: x.get("date") or "")
        avg_ms = int(sum_ms / total_calls) if total_calls else 0

        return success_response(data={
            "window": {"since_days": since_days, "since_ts": since_ts, "now_ts": now_ts},
            "total_calls": total_calls,
            "ok_calls": ok_calls,
            "fail_calls": fail_calls,
            "avg_duration_ms": avg_ms,
            "token_total": token_total,
            "by_day": by_day_list,
            "by_provider": by_provider,
        })
    except Exception as e:
        logger.exception("获取 LLM 指标失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/jobs", methods=["GET"])
def api_jobs_list():
    try:
        limit = int(request.args.get("limit") or 50)
        return success_response(data=_job_queue.list(limit=limit))
    except Exception as e:
        logger.exception("jobs list failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/jobs/<job_id>", methods=["GET"])
def api_jobs_get(job_id):
    try:
        job = _job_queue.get(str(job_id))
        if not job:
            return error_response("job 不存在", status_code=404)
        return success_response(data=job)
    except Exception as e:
        logger.exception("jobs get failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/jobs/<job_id>/cancel", methods=["POST"])
def api_jobs_cancel(job_id):
    try:
        ok = _job_queue.cancel(str(job_id))
        if not ok:
            return error_response("无法取消（可能已结束或不存在）", status_code=400)
        return success_response(message="已请求取消")
    except Exception as e:
        logger.exception("jobs cancel failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/feishu/watch", methods=["GET"])
def api_feishu_watch_list():
    try:
        return success_response(data=_load_watch_state())
    except Exception as e:
        logger.exception("feishu watch list failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/feishu/watch", methods=["POST"])
def api_feishu_watch_update():
    try:
        data = request.get_json() or {}
        action = str(data.get("action") or "upsert").strip().lower()
        doc_url = str(data.get("doc_url") or "").strip()
        name = str(data.get("name") or "").strip()
        enabled = bool(data.get("enabled", True))

        if action in ("upsert", "add", "update") and not doc_url:
            return error_response("doc_url 不能为空", status_code=400)

        state = _load_watch_state()
        items = state.get("items") if isinstance(state.get("items"), list) else []
        key = _match_watch_item(doc_url) if doc_url else str(data.get("key") or "")

        if action in ("remove", "delete"):
            state["items"] = [it for it in items if not (isinstance(it, dict) and str(it.get("key") or "") == key)]
            _save_watch_state(state)
            return success_response(message="已移除监听")

        target = None
        for it in items:
            if not isinstance(it, dict):
                continue
            if str(it.get("key") or "") == key:
                target = it
                break
        if not target:
            target = {"key": key}
            items.append(target)

        target["doc_url"] = doc_url
        target["name"] = name or target.get("name") or key
        target["enabled"] = enabled
        target["updated_at"] = int(time.time())
        state["items"] = items
        _save_watch_state(state)
        return success_response(data=target, message="已保存监听配置")
    except Exception as e:
        logger.exception("feishu watch update failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/feishu/watch/trigger", methods=["POST"])
def api_feishu_watch_trigger():
    try:
        data = request.get_json() or {}
        doc_url = str(data.get("doc_url") or "").strip()
        key = str(data.get("key") or "").strip()
        use_llm = bool(data.get("use_llm", True))
        if not doc_url and not key:
            return error_response("doc_url 或 key 至少提供一个", status_code=400)

        state = _load_watch_state()
        items = state.get("items") if isinstance(state.get("items"), list) else []
        target = None
        for it in items:
            if not isinstance(it, dict):
                continue
            if key and str(it.get("key") or "") == key:
                target = it
                break
            if doc_url and str(it.get("doc_url") or "") == doc_url:
                target = it
                break
        if target:
            doc_url = doc_url or str(target.get("doc_url") or "")
            name = str(target.get("name") or "")
            key = str(target.get("key") or "")
        else:
            name = ""
            if not doc_url:
                return error_response("未找到监听项", status_code=404)
            key = _match_watch_item(doc_url)
            target = {"key": key, "name": name or key, "doc_url": doc_url, "enabled": True}
            items.append(target)
            state["items"] = items

        if not bool(target.get("enabled", True)):
            return error_response("该监听项已禁用", status_code=400)

        job_id = _job_queue.enqueue("feishu_doc_audit", {"doc_url": doc_url, "name": name, "use_llm": use_llm})
        target["last_job_id"] = job_id
        target["last_trigger_at"] = int(time.time())
        _save_watch_state(state)
        return success_response(data={"job_id": job_id, "key": key})
    except Exception as e:
        logger.exception("feishu watch trigger failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/feishu/event", methods=["POST"])
def api_feishu_event():
    """
    飞书事件回调（轻量实现）：
    - 支持 url_verification challenge
    - 从事件载荷中提取 doc token / doc url，匹配 watch_list 后自动触发审计队列
    """
    try:
        body = request.get_json(silent=True) or {}
        if str(body.get("type") or "") == "url_verification":
            return jsonify({"challenge": body.get("challenge")})

        # 校验 token（如果配置了）
        cfg_path = _pick_llm_config_path()
        verify_token = ""
        listen_enabled = False
        if os.path.exists(cfg_path):
            try:
                with open(cfg_path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                verify_token = str(cfg.get("feishu_event_verify_token") or "").strip()
                listen_enabled = bool(cfg.get("feishu_listen_enabled"))
            except Exception:
                pass
        if not listen_enabled:
            return success_response(message="飞书监听未启用")
        if verify_token:
            token_in = ""
            if isinstance(body.get("token"), str):
                token_in = body.get("token")
            header = body.get("header") if isinstance(body.get("header"), dict) else {}
            if isinstance(header.get("token"), str):
                token_in = header.get("token")
            if token_in and token_in != verify_token:
                return error_response("verify token mismatch", status_code=403)

        def _collect_candidates(obj: Any, out: set) -> None:
            if isinstance(obj, dict):
                for k, v in obj.items():
                    lk = str(k).lower()
                    if lk in ("doc_token", "document_id", "docx_id", "obj_token", "object_token", "wiki_token"):
                        if isinstance(v, str) and v.strip():
                            out.add(v.strip())
                    if isinstance(v, str) and ("feishu.cn/docx/" in v or "feishu.cn/docs/" in v or "larksuite.cn/wiki/" in v or "feishu.cn/wiki/" in v):
                        out.add(v.strip())
                    _collect_candidates(v, out)
            elif isinstance(obj, list):
                for it in obj:
                    _collect_candidates(it, out)

        candidates = set()
        _collect_candidates(body, candidates)

        state = _load_watch_state()
        items = state.get("items") if isinstance(state.get("items"), list) else []
        triggered = []
        for it in items:
            if not isinstance(it, dict):
                continue
            if not bool(it.get("enabled", True)):
                continue
            doc_url = str(it.get("doc_url") or "").strip()
            key = str(it.get("key") or "")
            if not doc_url or not key:
                continue
            matched = False
            for c in candidates:
                if c in doc_url or c == key.split(":", 1)[-1]:
                    matched = True
                    break
            if not matched:
                continue
            last_ts = int(it.get("last_trigger_at") or 0)
            if last_ts and int(time.time()) - last_ts < 15:
                continue
            job_id = _job_queue.enqueue("feishu_doc_audit", {"doc_url": doc_url, "name": it.get("name") or "", "use_llm": True})
            it["last_job_id"] = job_id
            it["last_trigger_at"] = int(time.time())
            triggered.append({"key": key, "job_id": job_id})

        state["items"] = items
        _save_watch_state(state)
        return success_response(data={"triggered": triggered, "candidate_count": len(candidates)})
    except Exception as e:
        logger.exception("feishu event failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/push_to_feishu", methods=["POST"])
def api_push_to_feishu():
    """将审计结果推送到飞书 Webhook"""
    try:
        data = request.get_json() or {}
        title = data.get("title", "PRD 审计风险提醒")
        content = data.get("content", "")
        
        # 获取配置中的 Webhook
        config_path = _pick_llm_config_path()
        webhook_url = ""
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                config = json.load(f)
                webhook_url = config.get("feishu_webhook", "")
        
        if not webhook_url:
            return error_response("未配置飞书 Webhook 地址，请在 LLM 配置中设置", status_code=400)
            
        import requests
        # 构造飞书卡片消息
        payload = {
            "msg_type": "interactive",
            "card": {
                "config": {"wide_screen_mode": True},
                "header": {
                    "title": {"tag": "plain_text", "content": title},
                    "template": "red" if "P0" in content or "风险" in title else "orange"
                },
                "elements": [
                    {
                        "tag": "div",
                        "text": {"tag": "lark_md", "content": content}
                    },
                    {
                        "tag": "hr"
                    },
                    {
                        "tag": "note",
                        "elements": [{"tag": "plain_text", "content": f"推送时间: {time.strftime('%Y-%m-%d %H:%M:%S')}"}]
                    }
                ]
            }
        }
        
        resp = requests.post(webhook_url, json=payload, timeout=10)
        if resp.status_code == 200:
            return success_response(message="已成功推送到飞书")
        else:
            return error_response(f"飞书推送失败: {resp.text}", status_code=resp.status_code)
            
    except Exception as e:
        logger.exception("推送到飞书失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/rules", methods=["GET"])
def api_get_rules():
    """获取当前的审计规则库"""
    try:
        rules_path = os.path.join(os.path.dirname(__file__), "prd_scan_rules_v2.json")
        rules = _load_json_file(rules_path, {"rules": []})
        return success_response(data=rules)
    except Exception as e:
        logger.exception("获取审计规则失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/rules", methods=["POST"])
def api_update_rules():
    """更新审计规则库"""
    try:
        data = request.get_json() or {}
        if "rules" not in data:
            return error_response("缺少 rules 字段", status_code=400)
            
        rules_path = os.path.join(os.path.dirname(__file__), "prd_scan_rules_v2.json")
        _save_json_file(rules_path, data)
        return success_response(message="审计规则库已更新")
    except Exception as e:
        logger.exception("更新审计规则失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/rules/hit_stats", methods=["GET"])
def api_rule_hit_stats():
    """规则命中率统计：哪些规则高频命中，哪些从未命中"""
    try:
        since_days = int(request.args.get("since_days") or 30)
        limit = int(request.args.get("limit") or 2000)
        top_n = int(request.args.get("top_n") or 20)
        never_n = int(request.args.get("never_n") or 50)
        now_ts = int(time.time())
        since_ts = now_ts - max(0, since_days) * 86400

        rules_path = os.path.join(os.path.dirname(__file__), "prd_scan_rules_v2.json")
        rules_obj = _load_json_file(rules_path, {"rules": []})
        rules_list = rules_obj.get("rules") if isinstance(rules_obj, dict) else []
        rules_list = rules_list if isinstance(rules_list, list) else []

        rules_by_type = {}
        for r in rules_list:
            if not isinstance(r, dict):
                continue
            t = str(r.get("type") or "").strip()
            if not t:
                continue
            rules_by_type[t] = {
                "type": t,
                "level": str(r.get("level") or "").strip(),
                "description": str(r.get("description") or "").strip(),
                "suggestion": str(r.get("suggestion") or "").strip(),
            }

        from .audit_learning import load_all_snapshots

        snaps = load_all_snapshots(limit=max(1, min(limit, 5000)))
        snap_count = 0
        hit = {}
        unknown_hit = {}

        for snap in snaps:
            if not isinstance(snap, dict):
                continue
            created_at = int(snap.get("created_at") or 0)
            if created_at and created_at < since_ts:
                continue
            snap_count += 1
            stage2 = snap.get("stage2_output") if isinstance(snap.get("stage2_output"), dict) else {}
            defects = stage2.get("defects") if isinstance(stage2, dict) else []
            if not isinstance(defects, list):
                continue

            for d in defects:
                if not isinstance(d, dict):
                    continue
                t = str(d.get("type") or "").strip()
                if not t:
                    continue
                lv = str(d.get("risk_level") or d.get("level") or "P2").upper()

                bucket = hit if t in rules_by_type else unknown_hit
                item = bucket.get(t)
                if not item:
                    item = {
                        "type": t,
                        "count": 0,
                        "risk_levels": {"P0": 0, "P1": 0, "P2": 0},
                        "last_seen": 0,
                    }
                    bucket[t] = item

                item["count"] += 1
                if lv in item["risk_levels"]:
                    item["risk_levels"][lv] += 1
                else:
                    item["risk_levels"][lv] = int(item["risk_levels"].get(lv, 0)) + 1
                if created_at and created_at > int(item["last_seen"] or 0):
                    item["last_seen"] = created_at

        stats = []
        for t, meta in rules_by_type.items():
            h = hit.get(t) or {"count": 0, "risk_levels": {"P0": 0, "P1": 0, "P2": 0}, "last_seen": 0}
            stats.append({
                "type": t,
                "level": meta.get("level"),
                "description": meta.get("description"),
                "count": int(h.get("count") or 0),
                "risk_levels": h.get("risk_levels") or {"P0": 0, "P1": 0, "P2": 0},
                "last_seen": int(h.get("last_seen") or 0),
            })

        stats.sort(key=lambda x: (-int(x.get("count") or 0), str(x.get("type") or "")))
        never_hit = [x for x in stats if int(x.get("count") or 0) == 0]
        top_hit = [x for x in stats if int(x.get("count") or 0) > 0][: max(0, top_n)]

        unknown_sorted = list(unknown_hit.values())
        unknown_sorted.sort(key=lambda x: (-int(x.get("count") or 0), str(x.get("type") or "")))

        return success_response(data={
            "window": {"since_days": since_days, "since_ts": since_ts, "now_ts": now_ts},
            "snapshot_count": snap_count,
            "rule_total": len(rules_by_type),
            "hit_rule_count": sum(1 for x in stats if int(x.get("count") or 0) > 0),
            "never_hit_count": len(never_hit),
            "top_hit": top_hit,
            "never_hit": never_hit[: max(0, never_n)],
            "unknown_hit_top": unknown_sorted[: 20],
        })
    except Exception as e:
        logger.exception("规则命中率统计失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/optimize_prd", methods=["POST"])
def api_optimize_prd():
    """AI 一键优化 PRD 内容"""
    try:
        data = request.get_json() or {}
        prd_content = (data.get("prd_content") or "").strip()
        defects = data.get("defects") or []
        
        if not prd_content:
            return error_response("PRD 内容不能为空", status_code=400)
            
        from .prd_rule_engine import PRD_OPTIMIZATION_PROMPT
        from utils.llm_client import call_llm_with_retry
        
        # 简化缺陷列表，减少 token
        trimmed_defects = []
        for d in defects[:30]: # 最多处理前30个缺陷
            trimmed_defects.append({
                "type": d.get("type"),
                "module": d.get("module"),
                "description": d.get("description"),
                "suggestion": d.get("suggestion")
            })
            
        prompt = PRD_OPTIMIZATION_PROMPT.replace("{prd_content}", prd_content).replace("{defects_json}", json.dumps(trimmed_defects, ensure_ascii=False))
        
        try:
            optimized_prd = call_llm_with_retry(
                messages=[{"role": "system", "content": "你是一个资深产品专家，擅长将模糊的需求转化为可量化、逻辑闭环的专业 PRD。"},
                          {"role": "user", "content": prompt}],
                timeout=180, # 润色需要较长时间
                max_retries=1
            )
            return success_response({"optimized_prd": optimized_prd})
        except Exception as e:
            logger.exception("LLM optimize_prd failed")
            return error_response(f"AI 优化失败: {str(e)}", status_code=500)
            
    except Exception as e:
        logger.exception("api_optimize_prd failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/diff_snapshots", methods=["POST"])
def api_diff_snapshots():
    """对比两个版本的 PRD 审计快照，分析需求变更及测试影响"""
    try:
        data = request.get_json() or {}
        id1 = data.get("id1")
        id2 = data.get("id2")
        
        if not id1 or not id2:
            return error_response("必须提供两个快照 ID 进行对比", status_code=400)
            
        snapshots_dir = os.path.join(os.path.dirname(__file__), "learning_repo", "snapshots")
        s1_path = os.path.join(snapshots_dir, f"{id1}.json")
        s2_path = os.path.join(snapshots_dir, f"{id2}.json")
        
        if not os.path.exists(s1_path) or not os.path.exists(s2_path):
            return error_response("快照文件不存在", status_code=404)
            
        s1 = _load_json_file(s1_path, {})
        s2 = _load_json_file(s2_path, {})
        
        prd_text_1 = s1.get("prd_text", "") if isinstance(s1, dict) else ""
        prd_text_2 = s2.get("prd_text", "") if isinstance(s2, dict) else ""

        # 先用文本 diff 抽取“差异段”，再喂给 LLM，避免长 PRD 截断丢信息
        import difflib

        diff_lines = list(
            difflib.unified_diff(
                (prd_text_1 or "").splitlines(),
                (prd_text_2 or "").splitlines(),
                fromfile=f"{id1}",
                tofile=f"{id2}",
                lineterm="",
                n=3,
            )
        )

        # 过滤掉 header，只保留 hunks
        hunks = []
        for line in diff_lines:
            if line.startswith(("---", "+++", "@@")) or line.startswith(("+", "-", " ")):
                hunks.append(line)
        diff_text = "\n".join(hunks).strip()
        if not diff_text:
            diff_text = "（未检测到文本差异或差异为空）"

        # 控制输入长度，避免 token 过大
        max_chars = int(data.get("max_chars") or 12000)
        if len(diff_text) > max_chars:
            diff_text = diff_text[:max_chars] + "\n...（diff 已截断）"

        s1_meta = {
            "snapshot_id": s1.get("snapshot_id"),
            "created_at_str": s1.get("created_at_str"),
            "offline_mode": s1.get("offline_mode"),
        }
        s2_meta = {
            "snapshot_id": s2.get("snapshot_id"),
            "created_at_str": s2.get("created_at_str"),
            "offline_mode": s2.get("offline_mode"),
        }
        # extras 中包含大量产物，这里仅抽取少量用于上下文
        extras1 = s1.get("extras") if isinstance(s1.get("extras"), dict) else {}
        extras2 = s2.get("extras") if isinstance(s2.get("extras"), dict) else {}
        meta1 = extras1.get("architecture_scan") if isinstance(extras1.get("architecture_scan"), dict) else {}
        meta2 = extras2.get("architecture_scan") if isinstance(extras2.get("architecture_scan"), dict) else {}
        
        from utils.llm_client import call_llm_with_retry
        
        prompt = f"""你是一个资深的测试架构师。请基于“统一 diff 差异段”对比两个版本 PRD 的变化，并评估对回归测试的影响。

注意：
1) diff 已包含上下文（每段前后 3 行），请以 diff 为主进行分析；
2) 如果 diff 被截断，请在报告中标注“对比结果不完整”，并给出需要补充的信息清单；
3) 输出 Markdown，不要输出 JSON。

【版本 1 元信息】：{json.dumps(s1_meta, ensure_ascii=False)}
【版本 2 元信息】：{json.dumps(s2_meta, ensure_ascii=False)}

【版本 1 架构扫描（可选）】：{json.dumps(meta1, ensure_ascii=False)[:2000]}
【版本 2 架构扫描（可选）】：{json.dumps(meta2, ensure_ascii=False)[:2000]}

【PRD 差异段（unified diff）】：
{diff_text}

请输出《PRD 版本对比与测试影响分析报告》，包含：
1. 📌 主要需求变更点（新增/修改/删除）
2. 🔄 受影响的业务流程与接口（可推断即可，不确定要标注）
3. 🧪 测试范围调整建议（新增用例/回归用例/无需重测）
4. ⚠️ 风险预警（数据兼容、幂等、权限、并发、性能等）

输出报告："""

        try:
            diff_report = call_llm_with_retry(
                messages=[{"role": "user", "content": prompt}],
                timeout=120,
                max_retries=1
            )
            return success_response({"diff_report": diff_report})
        except Exception as e:
            logger.exception("LLM diff_snapshots failed")
            return error_response(f"AI 对比失败: {str(e)}", status_code=500)
            
    except Exception as e:
        logger.exception("api_diff_snapshots failed")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/export_test_cases", methods=["POST"])
def api_export_test_cases():
    """将生成的测试用例导出为 CSV 或 Excel"""
    try:
        data = request.get_json() or {}
        test_cases = data.get("test_cases") or []
        export_format = data.get("format", "csv").lower()
        
        if not test_cases:
            return error_response("没有可导出的用例数据", status_code=400)

        cols = ["case_id", "priority", "module", "feature", "precondition", "steps", "expected"]
        col_map = {
            "case_id": "用例ID",
            "priority": "优先级",
            "module": "所属模块",
            "feature": "功能点",
            "precondition": "前置条件",
            "steps": "测试步骤",
            "expected": "预期结果",
        }

        if export_format != "xlsx":
            buf = io.StringIO()
            writer = csv.writer(buf)
            writer.writerow([col_map[c] for c in cols])
            for item in test_cases:
                if not isinstance(item, dict):
                    continue
                writer.writerow([str(item.get(c, "") or "") for c in cols])
            csv_data = buf.getvalue()
            return Response(
                csv_data.encode("utf-8-sig"),
                mimetype="text/csv",
                headers={"Content-disposition": f"attachment; filename=PRD_Test_Cases_{int(time.time())}.csv"},
            )

        try:
            import pandas as pd
            from io import BytesIO
        except ImportError:
            return error_response("导出 Excel 需要安装 pandas 与 openpyxl（建议：pip install pandas openpyxl）", status_code=400)

        df = pd.DataFrame([x for x in test_cases if isinstance(x, dict)])
        df = df[[c for c in cols if c in df.columns]]
        df.rename(columns=col_map, inplace=True)

        output = BytesIO()
        with pd.ExcelWriter(output, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name="TestCases")
        output.seek(0)
        filename = f"PRD_Test_Cases_{int(time.time())}.xlsx"
        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name=filename,
        )
        
    except Exception as e:
        logger.exception("导出测试用例失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/test_cases/push", methods=["POST"])
def api_push_test_cases():
    """推送生成的测试用例到外部用例管理系统（通用 HTTP Webhook）"""
    try:
        data = request.get_json() or {}
        test_cases = data.get("test_cases") or []
        if not isinstance(test_cases, list) or not test_cases:
            return error_response("没有可推送的用例数据", status_code=400)

        title = (data.get("title") or "PRD 自动生成测试用例").strip()
        project = (data.get("project") or "").strip()
        push_url = (data.get("push_url") or "").strip()
        push_token = (data.get("push_token") or "").strip()

        config_path = _pick_llm_config_path()
        config = {}
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
            except Exception:
                config = {}

        if not push_url:
            push_url = str(config.get("case_system_push_url") or "").strip()
        if not project:
            project = str(config.get("case_system_project") or "").strip()
        if not push_token:
            push_token = str(config.get("case_system_push_token") or "").strip()

        if not push_url:
            return error_response("未配置用例系统推送地址，请在 LLM 配置中设置", status_code=400)

        payload = {
            "source": "prd_audit",
            "title": title,
            "project": project,
            "generated_at": int(time.time()),
            "test_cases": test_cases,
        }
        if isinstance(data.get("meta"), dict):
            payload["meta"] = data.get("meta")

        headers = {"Content-Type": "application/json"}
        if push_token and "****" not in push_token:
            headers["Authorization"] = f"Bearer {push_token}"

        import requests

        resp = requests.post(push_url, json=payload, headers=headers, timeout=20)
        content_type = (resp.headers.get("Content-Type") or "").lower()
        body_preview = resp.text[:2000] if isinstance(resp.text, str) else ""
        out = {
            "status_code": resp.status_code,
            "content_type": content_type,
            "body_preview": body_preview,
        }
        try:
            if "application/json" in content_type:
                out["body_json"] = resp.json()
        except Exception:
            pass

        if 200 <= resp.status_code < 300:
            return success_response(data=out, message="已推送到用例系统")
        return error_response(f"推送失败: HTTP {resp.status_code}", status_code=resp.status_code)
    except Exception as e:
        logger.exception("推送测试用例失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/check_dependencies", methods=["GET"])
def api_check_dependencies():
    """检查 PRD 审计模块所需的环境依赖库是否已安装"""
    results = {
        "pypdf": {"installed": False, "version": None, "required": True},
        "python-docx": {"installed": False, "version": None, "required": True},
        "requests": {"installed": False, "version": None, "required": True},
        "pandas": {"installed": False, "version": None, "required": False},
        "openpyxl": {"installed": False, "version": None, "required": False},
    }
    
    # 检查 pypdf
    try:
        import pypdf
        results["pypdf"]["installed"] = True
        results["pypdf"]["version"] = getattr(pypdf, "__version__", "unknown")
    except ImportError:
        try:
            import PyPDF2
            results["pypdf"]["installed"] = True
            results["pypdf"]["version"] = f"PyPDF2 {getattr(PyPDF2, '__version__', 'unknown')}"
        except ImportError:
            pass
            
    # 检查 python-docx
    try:
        import docx
        results["python-docx"]["installed"] = True
        results["python-docx"]["version"] = getattr(docx, "__version__", "unknown")
    except ImportError:
        pass
        
    # 检查 requests
    try:
        import requests
        results["requests"]["installed"] = True
        results["requests"]["version"] = getattr(requests, "__version__", "unknown")
    except ImportError:
        pass

    # 检查 pandas（XLSX 导出依赖）
    try:
        import pandas
        results["pandas"]["installed"] = True
        results["pandas"]["version"] = getattr(pandas, "__version__", "unknown")
    except ImportError:
        pass

    # 检查 openpyxl（XLSX 导出依赖）
    try:
        import openpyxl
        results["openpyxl"]["installed"] = True
        results["openpyxl"]["version"] = getattr(openpyxl, "__version__", "unknown")
    except ImportError:
        pass
        
    all_ok = all(v["installed"] for v in results.values() if v["required"])
    xlsx_ok = results["pandas"]["installed"] and results["openpyxl"]["installed"]
    return success_response(data={
        "dependencies": results,
        "all_ok": all_ok,
        "xlsx_ok": xlsx_ok,
        "message": "所有核心依赖已就绪" if all_ok else "部分核心依赖缺失，请检查后端环境"
    })

# ---------- API：保存为测试用例（写入 session，跳转用例管理） ----------

@prd_audit_bp.route("/api/prepare_save_to_cases", methods=["POST"])
def api_prepare_save_to_cases():
    """将当前报告写入 session，返回跳转 URL，用于「保存为测试用例」"""
    data = request.get_json() or {}
    report_md = (data.get("report_md") or data.get("content") or "").strip()
    if not report_md:
        return error_response("report_md 不能为空", status_code=400)
    session["prd_audit_report_md"] = report_md[:500000]
    session.modified = True
    return success_response(data={"redirect_url": url_for("test_case.index", _external=False) + "?from_prd=1#prd_review"})

@prd_audit_bp.route("/api/learning/status", methods=["GET"])
def api_learning_status():
    try:
        return success_response(data=get_learning_status())
    except Exception as e:
        logger.exception("获取学习状态失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/lane_stats", methods=["GET"])
def api_learning_lane_stats():
    try:
        limit = int(request.args.get("limit", 5000) or 5000)
        if limit < 1:
            limit = 1
        if limit > 20000:
            limit = 20000
        return success_response(data=get_learning_lane_stats(limit=limit))
    except Exception as e:
        logger.exception("获取分轨统计失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/quality_dashboard", methods=["GET"])
def api_learning_quality_dashboard():
    try:
        limit = int(request.args.get("limit", 5000) or 5000)
        if limit < 1:
            limit = 1
        if limit > 20000:
            limit = 20000
        return success_response(data=get_learning_quality_dashboard(limit=limit))
    except Exception as e:
        logger.exception("获取学习质量看板失败")
        return error_response(str(e), status_code=500)

# ---------- API：历史审计记录 (Snapshots) ----------

def _enrich_history_snapshot_index(idx: Dict[str, Any]) -> List[Dict[str, Any]]:
    from .audit_learning import SNAPSHOT_DIR, build_snapshot_index_entry

    snaps = idx.get("snapshots") if isinstance(idx, dict) else []
    snaps = snaps if isinstance(snaps, list) else []
    enriched: List[Dict[str, Any]] = []
    changed = False
    for item in snaps:
        if not isinstance(item, dict):
            continue
        merged = dict(item)
        need_refresh = not isinstance(item.get("preview"), dict) or not item.get("created_at_str")
        if need_refresh:
            file_path = os.path.join(SNAPSHOT_DIR, f"{item.get('snapshot_id')}.json")
            if os.path.exists(file_path):
                try:
                    with open(file_path, "r", encoding="utf-8") as f:
                        payload = json.load(f)
                    if isinstance(payload, dict):
                        merged = build_snapshot_index_entry(payload)
                        changed = True
                except Exception:
                    merged = dict(item)
        enriched.append(merged)
    if changed and isinstance(idx, dict):
        idx["snapshots"] = enriched[-5000:]
    return enriched

@prd_audit_bp.route("/api/history/snapshots", methods=["GET"])
def api_history_snapshots():
    """获取所有历史审计快照的简要列表"""
    try:
        from .audit_learning import INDEX_FILE
        if not os.path.exists(INDEX_FILE):
            return success_response(data={"snapshots": []})
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            idx = json.load(f)
        snaps = _enrich_history_snapshot_index(idx)
        if isinstance(idx, dict):
            with open(INDEX_FILE, "w", encoding="utf-8") as f:
                json.dump(idx, f, ensure_ascii=False, indent=2)
        snaps.sort(key=lambda x: x.get("created_at", 0), reverse=True)
        return success_response(data={"snapshots": snaps})
    except Exception as e:
        logger.exception("获取历史记录失败")
        return error_response(f"获取历史记录失败: {e}", status_code=500)

@prd_audit_bp.route("/api/history/snapshot/<snapshot_id>", methods=["GET"])
def api_history_snapshot_detail(snapshot_id):
    """获取单个历史快照的完整详情"""
    try:
        from .audit_learning import SNAPSHOT_DIR, build_snapshot_preview
        file_path = os.path.join(SNAPSHOT_DIR, f"{snapshot_id}.json")
        if not os.path.exists(file_path):
            return error_response("未找到该历史记录", status_code=404)
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, dict) and not isinstance(data.get("preview"), dict):
            data["preview"] = build_snapshot_preview(data)
        return success_response(data=data)
    except Exception as e:
        logger.exception("读取历史记录详情失败")
        return error_response(f"读取历史记录详情失败: {e}", status_code=500)

@prd_audit_bp.route("/api/review/state", methods=["GET"])
def api_review_state():
    try:
        snapshot_id = (request.args.get("snapshot_id") or "").strip()
        if not snapshot_id:
            return error_response("缺少 snapshot_id", status_code=400)
        state = _load_json_file(REVIEW_STATE_FILE, {"snapshots": {}})
        snaps = state.get("snapshots") if isinstance(state, dict) else {}
        snap_state = snaps.get(snapshot_id) if isinstance(snaps, dict) else None
        snap_state = snap_state if isinstance(snap_state, dict) else {"items": {}}
        items = snap_state.get("items") if isinstance(snap_state.get("items"), dict) else {}
        counts = {"NEW": 0, "OWNED": 0, "FIXED": 0, "WONTFIX": 0, "DISPUTED": 0}
        for v in items.values():
            if not isinstance(v, dict):
                continue
            st = str(v.get("status") or "NEW").upper()
            if st not in counts:
                st = "NEW"
            counts[st] += 1
        return success_response(data={"snapshot_id": snapshot_id, "items": items, "counts": counts})
    except Exception as e:
        logger.exception("读取评审状态失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/review/update", methods=["POST"])
def api_review_update():
    try:
        data = request.get_json() or {}
        snapshot_id = str(data.get("snapshot_id") or "").strip()
        defect = data.get("defect") if isinstance(data.get("defect"), dict) else {}
        if not snapshot_id:
            return error_response("缺少 snapshot_id", status_code=400)
        if not isinstance(defect, dict) or not defect:
            return error_response("缺少 defect", status_code=400)

        key = _defect_key(defect)
        status = str(data.get("status") or "").strip().upper() or None
        assignee = str(data.get("assignee") or "").strip()
        mentions = str(data.get("mentions") or "").strip()
        comment = str(data.get("comment") or "").strip()
        actor = str(data.get("actor") or "").strip()

        allowed = {"NEW", "OWNED", "FIXED", "WONTFIX", "DISPUTED"}
        if status and status not in allowed:
            return error_response("不支持的 status", status_code=400)

        os.makedirs(os.path.dirname(REVIEW_STATE_FILE), exist_ok=True)
        state = _load_json_file(REVIEW_STATE_FILE, {"snapshots": {}})
        if not isinstance(state, dict):
            state = {"snapshots": {}}
        snaps = state.get("snapshots")
        if not isinstance(snaps, dict):
            snaps = {}
            state["snapshots"] = snaps
        snap_state = snaps.get(snapshot_id)
        if not isinstance(snap_state, dict):
            snap_state = {"items": {}, "events": [], "updated_at": 0}
            snaps[snapshot_id] = snap_state
        items = snap_state.get("items")
        if not isinstance(items, dict):
            items = {}
            snap_state["items"] = items
        item = items.get(key)
        if not isinstance(item, dict):
            item = {"status": "NEW", "assignee": "", "mentions": "", "last_comment": "", "updated_at": 0}
            items[key] = item

        changed = {}
        if status is not None:
            item["status"] = status
            changed["status"] = status
        if assignee is not None:
            item["assignee"] = assignee
            changed["assignee"] = assignee
        if mentions is not None:
            item["mentions"] = mentions
            changed["mentions"] = mentions
        if comment:
            item["last_comment"] = comment
            changed["comment"] = comment

        now_ts = int(time.time())
        item["updated_at"] = now_ts
        snap_state["updated_at"] = now_ts

        events = snap_state.get("events")
        if not isinstance(events, list):
            events = []
            snap_state["events"] = events
        if changed:
            ev = {"ts": now_ts, "defect_key": key, "actor": actor, "changed": changed}
            events.append(ev)
            snap_state["events"] = events[-2000:]

        _save_json_file(REVIEW_STATE_FILE, state)
        return success_response(data={"snapshot_id": snapshot_id, "defect_key": key, "item": item})
    except Exception as e:
        logger.exception("更新评审状态失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/build_rule_draft", methods=["POST"])
def api_learning_build_rule_draft():
    try:
        data = request.get_json() or {}
        min_count = int(data.get("min_count", 2) or 2)
        max_new_rules = int(data.get("max_new_rules", 30) or 30)
        if min_count < 1:
            min_count = 1
        if max_new_rules < 1:
            max_new_rules = 1
        result = build_rule_draft_from_snapshots(min_count=min_count, max_new_rules=max_new_rules)
        return success_response(data=result, message="规则草案已生成")
    except Exception as e:
        logger.exception("生成规则草案失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/rule_candidates", methods=["GET"])
def api_learning_rule_candidates():
    try:
        return success_response(data=load_rule_candidates())
    except Exception as e:
        logger.exception("读取规则候选失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/apply_candidates", methods=["POST"])
def api_learning_apply_candidates():
    try:
        data = request.get_json() or {}
        selected = data.get("selected_rule_names")
        if not isinstance(selected, list):
            selected = []
        max_new_rules = int(data.get("max_new_rules", 100) or 100)
        if max_new_rules < 1:
            max_new_rules = 1
        if max_new_rules > 1000:
            max_new_rules = 1000
        result = apply_selected_candidates(selected_names=selected, max_new_rules=max_new_rules)
        return success_response(data=result, message="已生成可应用规则文件")
    except Exception as e:
        logger.exception("应用候选规则失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/publish_applied", methods=["POST"])
def api_learning_publish_applied():
    try:
        data = request.get_json() or {}
        create_backup = bool(data.get("create_backup", True))
        result = publish_applied_rules(create_backup=create_backup)
        return success_response(data=result, message="已发布为正式规则")
    except FileNotFoundError as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=400)
    except ValueError as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=400)
    except Exception as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/outline_owner_correction", methods=["POST"])
def api_learning_outline_owner_correction():
    try:
        data = request.get_json() or {}
        prd_text = str(data.get("prd_text") or "")
        flow_rows = data.get("flow_rows")
        role_rows = data.get("role_rows")
        meta = data.get("meta")
        if not isinstance(flow_rows, list) or not flow_rows:
            return error_response("flow_rows 不能为空", status_code=400)
        result = save_outline_owner_correction(
            prd_text=prd_text,
            flow_rows=flow_rows,
            role_rows=role_rows if isinstance(role_rows, list) else [],
            meta=meta if isinstance(meta, dict) else {},
        )
        return success_response(data=result, message="责任角色校正已保存")
    except Exception as e:
        logger.exception("保存责任角色校正失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/generate_test_code", methods=["POST"])
def api_generate_test_code():
    """根据测试矩阵生成 Pytest 测试代码（流式）"""
    try:
        data = request.get_json() or {}
        test_matrix = data.get("test_matrix")
        
        if not test_matrix or not isinstance(test_matrix, dict):
            return error_response("缺少有效的 test_matrix 数据", status_code=400)

        # 构造 Prompt
        prompt = (
            "你是一个高级测试工程师。请根据以下【功能测试矩阵】JSON数据，"
            "编写一份完整的 Python Pytest 自动化测试脚本。\n\n"
            "要求：\n"
            "1. 使用 pytest 框架，包含 fixture Setup/Teardown。\n"
            "2. 每个测试用例对应矩阵中的一个 case_id。\n"
            "3. 代码中包含详细注释（模块、预期结果、判断依据）。\n"
            "4. 对于 '缺失' 或 '待确认' 的项，使用 @pytest.mark.skip 或 @pytest.mark.xfail 标记，并注明原因。\n"
            "5. 生成的代码必须是可执行的 Python 代码块（不要 markdown ```python 标记，直接输出代码）。\n\n"
            f"测试矩阵数据：\n{json.dumps(test_matrix, ensure_ascii=False, indent=2)}"
        )

        from utils.llm_client import load_llm_config, stream_chat_content
        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
            # 强制检查 API Key
            if not (llm_config.get("api_key") or "").strip():
                raise ValueError("LLM API Key 未配置")
        except Exception as e:
            return error_response(f"生成代码需要配置大模型 API Key: {e}", status_code=400)

        def generate():
            try:
                # 调用流式生成
                for chunk in stream_chat_content(llm_config, prompt):
                    if chunk:
                        # 简单过滤 markdown 标记，防止前端渲染混乱
                        # chunk = chunk.replace("```python", "").replace("```", "")
                        yield chunk
            except Exception as e:
                logger.error(f"Test code generation failed: {e}")
                yield f"\n# 生成失败: {str(e)}"

        return Response(stream_with_context(generate()), mimetype="text/plain")

    except Exception as e:
        logger.exception("测试代码生成接口异常")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/export_feature_xmind", methods=["POST"])
def api_export_feature_xmind():
    """根据功能节点导出 .xmind 文件（本地生成，无需大模型）"""
    try:
        data = request.get_json() or {}
        nodes = data.get("nodes")
        if not isinstance(nodes, list) or not nodes:
            return error_response("nodes 不能为空", status_code=400)
        xmind_buf = _build_xmind_file(nodes)
        return send_file(
            xmind_buf,
            mimetype="application/vnd.xmind.workbook",
            as_attachment=True,
            download_name="PRD功能导图.xmind",
        )
    except Exception as e:
        logger.exception("导出 XMind 失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/knowledge_cards", methods=["GET", "POST"])
def api_knowledge_cards():
    """获取或保存功能知识卡片"""
    try:
        if request.method == "POST":
            data = request.get_json() or {}
            items = data.get("items")
            saved = _save_knowledge_cards(items)
            return success_response(data={"items": saved, "count": len(saved)}, message="能力库已保存")

        domain = (request.args.get("domain") or "").strip().lower()
        cards = _load_knowledge_cards()
        if domain and domain not in ["ktv", "ktv点歌系统", "all"]:
            cards = []
        return success_response(data={"items": cards, "count": len(cards)})
    except Exception as e:
        logger.exception("获取功能知识卡片失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/knowledge_cards/import", methods=["POST"])
def api_knowledge_cards_import():
    try:
        file = request.files.get("file")
        if file:
            filename = file.filename.lower()
            if filename.endswith(".csv") or filename.endswith(".txt"):
                content = file.read().decode("utf-8-sig")
                reader = csv.DictReader(io.StringIO(content))
                items = []
                for row in reader:
                    if not row.get("name"): continue
                    item = {
                        "capability_id": row.get("capability_id") or f"CAP_{uuid.uuid4().hex[:8]}",
                        "domain": row.get("domain", ""),
                        "name": row.get("name", ""),
                        "desc": row.get("desc", ""),
                        "actors": [x for x in (row.get("actors") or "").split("|") if x],
                        "pre_conditions": [x for x in (row.get("pre_conditions") or "").split("|") if x],
                        "post_conditions": [x for x in (row.get("post_conditions") or "").split("|") if x],
                        "rules": [x for x in (row.get("rules") or "").split("|") if x],
                        "required_evidence": [x for x in (row.get("required_evidence") or "").split("|") if x],
                        "bad_patterns": [x for x in (row.get("bad_patterns") or "").split("|") if x],
                        "suggestion_template": row.get("suggestion_template", "")
                    }
                    items.append(item)
                if not items:
                    return error_response("CSV 解析为空或格式错误", status_code=400)
                
                existing = _load_knowledge_cards()
                existing.extend(items)
                saved = _save_knowledge_cards(existing)
                return success_response(data={"items": saved, "count": len(saved)}, message=f"成功导入 {len(items)} 条")

        data = request.get_json() or {}
        items = data.get("items")
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)
            
        existing = _load_knowledge_cards()
        existing.extend(items)
        saved = _save_knowledge_cards(existing)
        return success_response(data={"items": saved, "count": len(saved)}, message=f"成功导入 {len(items)} 条")
    except Exception as e:
        logger.exception("导入能力库失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/knowledge_cards/export", methods=["GET"])
def api_knowledge_cards_export():
    try:
        cards = _load_knowledge_cards()
        format_type = request.args.get("format", "json").lower()
        
        if format_type == "csv":
            sio = io.StringIO()
            writer = csv.writer(sio)
            writer.writerow(["capability_id", "domain", "name", "desc", "actors", "pre_conditions", "post_conditions", "rules", "required_evidence", "bad_patterns", "suggestion_template"])
            for c in cards:
                writer.writerow([
                    c.get("capability_id", ""),
                    c.get("domain", ""),
                    c.get("name", ""),
                    c.get("desc", ""),
                    "|".join(c.get("actors", [])),
                    "|".join(c.get("pre_conditions", [])),
                    "|".join(c.get("post_conditions", [])),
                    "|".join(c.get("rules", [])),
                    "|".join(c.get("required_evidence", [])),
                    "|".join(c.get("bad_patterns", [])),
                    c.get("suggestion_template", "")
                ])
            buf = io.BytesIO(sio.getvalue().encode("utf-8-sig"))
            return send_file(
                buf,
                mimetype="text/csv",
                as_attachment=True,
                download_name="knowledge_cards.csv",
            )
            
        buf = io.BytesIO()
        buf.write(json.dumps({"items": cards}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="knowledge_cards.json",
        )
    except Exception as e:
        logger.exception("导出能力库失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/backups", methods=["GET"])
def api_learning_backups():
    try:
        limit = int(request.args.get("limit", 20) or 20)
        if limit < 1:
            limit = 1
        if limit > 200:
            limit = 200
        return success_response(data={"items": list_rule_backups(limit=limit)})
    except Exception as e:
        logger.exception("获取规则备份失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/learning/rollback_backup", methods=["POST"])
def api_learning_rollback_backup():
    try:
        data = request.get_json() or {}
        backup_file_name = (data.get("backup_file_name") or "").strip()
        create_backup = bool(data.get("create_backup", True))
        if not backup_file_name:
            return error_response("backup_file_name 不能为空", status_code=400)
        result = rollback_rules_from_backup(backup_file_name=backup_file_name, create_backup=create_backup)
        return success_response(data=result, message="规则库已回滚")
    except FileNotFoundError as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=404)
    except ValueError as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=400)
    except Exception as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/rule_plugins", methods=["GET", "POST"])
def api_rule_plugins():
    try:
        if request.method == "POST":
            data = request.get_json() or {}
            profiles = data.get("profiles")
            if not isinstance(profiles, list):
                return error_response("profiles 必须为数组", status_code=400)
            saved = save_rule_plugin_profiles(profiles)
            return success_response(data={"profiles": saved, "count": len(saved)}, message="规则插件已保存")
        profiles = load_rule_plugin_profiles()
        return success_response(data={"profiles": profiles, "count": len(profiles)})
    except Exception as e:
        logger.exception("规则插件管理失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/rule_plugins/stats", methods=["GET"])
def api_rule_plugins_stats():
    try:
        limit = int(request.args.get("limit", 500) or 500)
        stats = get_plugin_usage_stats(limit=limit)
        return success_response(data=stats)
    except Exception as e:
        logger.exception("规则插件统计失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/prompt_center", methods=["GET", "POST"])
def api_prompt_center():
    try:
        if request.method == "POST":
            data = request.get_json() or {}
            saved = save_prompt_center(data)
            return success_response(data=saved, message="Prompt Center 已保存")
        data = load_prompt_center()
        return success_response(data=data)
    except Exception as e:
        logger.exception("Prompt Center 管理失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/prompt_center/stats", methods=["GET"])
def api_prompt_center_stats():
    try:
        limit = int(request.args.get("limit", 500) or 500)
        stats = get_prompt_evaluation_stats(limit=limit)
        return success_response(data=stats)
    except Exception as e:
        logger.exception("Prompt Center 统计失败")
        return error_response(str(e), status_code=500)

@prd_audit_bp.route("/api/gate/evaluate", methods=["POST"])
def api_gate_evaluate():
    try:
        data = request.get_json() or {}
        prd_text = (data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)
        cfg = _load_json_file(GATE_CONFIG_FILE, {"min_score": 80, "block_on_p0": True, "max_failed_rules": 12})
        result = run_outline_engine(prd_text, {}, {})
        dr = result.get("deterministic_rules") if isinstance(result.get("deterministic_rules"), dict) else {}
        defects = dr.get("defects") if isinstance(dr.get("defects"), list) else []
        score = int(dr.get("score", 0))
        failed_rules = int((dr.get("stats") or {}).get("failed_rules", len(defects)))
        has_p0 = any(str((d or {}).get("severity") or "").upper() == "P0" for d in defects if isinstance(d, dict))
        blocked = False
        reasons = []
        min_score = int(cfg.get("min_score", 80))
        max_failed = int(cfg.get("max_failed_rules", 12))
        block_on_p0 = bool(cfg.get("block_on_p0", True))
        if score < min_score:
            blocked = True
            reasons.append(f"score<{min_score} (actual={score})")
        if block_on_p0 and has_p0:
            blocked = True
            reasons.append("has_P0_defect")
        if failed_rules > max_failed:
            blocked = True
            reasons.append(f"failed_rules>{max_failed} (actual={failed_rules})")
        plugin = result.get("rule_plugin") if isinstance(result.get("rule_plugin"), dict) else {}
        return success_response(data={
            "score": score,
            "failed_rules": failed_rules,
            "has_p0": has_p0,
            "blocked": blocked,
            "reasons": reasons,
            "plugin": str(plugin.get("plugin_id") or "unknown"),
            "gate_config": {"min_score": min_score, "block_on_p0": block_on_p0, "max_failed_rules": max_failed},
        })
    except Exception as e:
        logger.exception("Gate 评估失败")
        return error_response(str(e), status_code=500)
