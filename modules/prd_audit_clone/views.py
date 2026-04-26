"""
PRD 审计独立模块：自包含实现，不依赖 test_case。
提供独立页与全部 API（生成、分析、解析 PDF/DOCX、导出 Word、默认提示词、LLM 配置、保存为测试用例）。
"""

import os
import json
import io
import re
import csv
import zipfile
import uuid
from collections import Counter
from flask import render_template, request, Response, stream_with_context, send_file
from utils.response import success_response, error_response
from utils.logger import setup_logger
from utils.llm_client import call_llm
from . import prd_audit_clone_bp as prd_audit_bp
from . import pipeline
from .outline_engine import run_outline_engine
from .outline_llm import run_outline_llm, merge_llm_with_local
from .feishu_client import is_feishu_doc_url, fetch_feishu_doc_content
from .audit_learning import (
    get_learning_status,
    get_learning_lane_stats,
    get_learning_quality_dashboard,
    build_rule_draft_from_snapshots,
    load_rule_candidates,
    apply_selected_candidates,
    publish_applied_rules,
    list_rule_backups,
    rollback_rules_from_backup,
)

logger = setup_logger("prd_audit_clone_module")

STORAGE_DIR = os.path.dirname(os.path.abspath(__file__))
os.makedirs(STORAGE_DIR, exist_ok=True)
DEFAULT_PRD_AUDIT_PROMPT_FILE = os.path.join(STORAGE_DIR, "prd_audit_prompt_default.txt")
# 独立副本：仅使用本目录下 llm_config.json（拷贝到其他机器时一起带走即可）
LLM_CONFIG_FILE = os.path.join(STORAGE_DIR, "llm_config.json")

FALLBACK_PRD_PROMPT = (
    "你是 PRD 审计专家。请基于输入 PRD 输出结构化风险报告，至少包含："
    "总体结论、漏洞与风险清单、测试重点、研发重点、计划建议。"
    "仅输出报告正文，不要开场白。\n\nPRD 内容：\n{content}"
)

KTV_CAPABILITY_CARDS = [
    {
        "capability_id": "KTV_PLAY_CTRL_001",
        "name": "切歌控制",
        "domain": "KTV点歌系统",
        "category": "播放控制",
        "priority": "P0",
        "trigger": "当前歌曲播放中，管理员点击“切歌”",
        "preconditions": [
            "管理员权限=是",
            "当前队列长度>=1",
        ],
        "system_behaviors": [
            "当前歌曲在500ms内停止播放",
            "队列下一首在2s内开始播放",
            "所有终端同步更新“当前播放歌曲”",
        ],
        "exceptions": [
            {"condition": "下一首解码失败", "action": "自动跳过并播放再下一首，记录错误码"},
            {"condition": "队列为空", "action": "进入待机页并提示“暂无歌曲”"},
        ],
        "logging": [
            "操作者ID",
            "房间ID",
            "歌曲ID",
            "操作时间",
            "结果码",
        ],
        "acceptance_criteria": [
            "100次切歌中，95次在2s内成功切换",
        ],
        "required_evidence": [
            "PRD中明确管理员切歌权限",
            "PRD中明确切歌时延SLA（500ms停止、2s起播）",
            "PRD中明确多终端同步规则",
            "PRD中明确异常分支（解码失败、队列为空）",
            "PRD中明确日志字段与验收口径",
        ],
        "bad_patterns": [
            "只写切歌功能，未写权限约束",
            "只写正常流程，未写异常回退",
            "未定义时延指标或验收标准",
            "未定义终端状态同步",
        ],
        "suggestion_template": "补充切歌控制时序图：鉴权校验→停止当前→拉起下一首→多端同步；并补充异常回退与日志字段定义。",
    }
]
KNOWLEDGE_CARDS_FILE = os.path.join(STORAGE_DIR, "knowledge_cards.json")
BUG_RAW_FILE = os.path.join(STORAGE_DIR, "bug_raw.json")
BUG_PATTERN_FILE = os.path.join(STORAGE_DIR, "bug_pattern.json")
BUG_RULE_FILE = os.path.join(STORAGE_DIR, "bug_rule.json")
VECTOR_DATA_FILE = os.path.join(STORAGE_DIR, "vector_data.json")

DEFAULT_BUG_PATTERNS = [
    {"pattern_id": "P001", "keywords": ["黑屏", "无画面", "无视频"], "category": "媒体异常", "design_gap": ["异常处理缺失", "资源释放缺失"], "rule": "媒体流程必须定义失败兜底策略", "weight": 0.9, "enabled": True},
    {"pattern_id": "P002", "keywords": ["卡死", "无响应", "假死"], "category": "稳定性", "design_gap": ["超时机制缺失", "状态回收缺失"], "rule": "关键交互必须定义超时与恢复策略", "weight": 0.88, "enabled": True},
    {"pattern_id": "P003", "keywords": ["崩溃", "闪退", "重启"], "category": "稳定性", "design_gap": ["异常隔离缺失"], "rule": "核心模块必须定义异常隔离与重试边界", "weight": 0.95, "enabled": True},
    {"pattern_id": "P004", "keywords": ["不同步", "状态错乱", "状态不一致"], "category": "状态管理", "design_gap": ["状态机缺失"], "rule": "多状态系统必须定义状态机与转移条件", "weight": 0.92, "enabled": True},
    {"pattern_id": "P005", "keywords": ["越权", "权限", "未授权"], "category": "权限安全", "design_gap": ["权限规则缺失"], "rule": "高风险操作必须定义角色鉴权和越权拦截", "weight": 0.9, "enabled": True},
    {"pattern_id": "P006", "keywords": ["并发", "抢占", "竞态"], "category": "并发冲突", "design_gap": ["并发裁决缺失"], "rule": "并发场景必须定义裁决优先级与幂等策略", "weight": 0.86, "enabled": True},
    {"pattern_id": "P007", "keywords": ["恢复失败", "无法恢复", "回退失败"], "category": "恢复机制", "design_gap": ["回退策略缺失"], "rule": "中断流程必须定义恢复或回退策略", "weight": 0.87, "enabled": True},
    {"pattern_id": "P008", "keywords": ["主板", "芯片", "平台差异"], "category": "平台兼容", "design_gap": ["平台差异约束缺失"], "rule": "涉及硬件能力的功能必须定义平台差异策略", "weight": 0.82, "enabled": True},
]

def _pick_llm_config_path() -> str:
    """返回本模块目录下的 LLM 配置（modules/prd_audit_clone/llm_config.json）。"""
    return LLM_CONFIG_FILE


def _normalize_knowledge_items(items):
    if not isinstance(items, list):
        return []
    out = []
    for it in items:
        if not isinstance(it, dict):
            continue
        capability_id = str(it.get("capability_id") or "").strip()
        name = str(it.get("name") or "").strip()
        if not capability_id or not name:
            continue
        out.append(it)
    return out


def _load_json_file(path, default_value):
    if os.path.exists(path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return default_value
    return default_value


def _save_json_file(path, data):
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _ensure_bug_assets():
    if not os.path.exists(BUG_PATTERN_FILE):
        _save_json_file(BUG_PATTERN_FILE, {"items": DEFAULT_BUG_PATTERNS})
    if not os.path.exists(BUG_RAW_FILE):
        _save_json_file(BUG_RAW_FILE, {"items": []})
    if not os.path.exists(BUG_RULE_FILE):
        _save_json_file(BUG_RULE_FILE, {"items": []})
    if not os.path.exists(VECTOR_DATA_FILE):
        _save_json_file(VECTOR_DATA_FILE, {"items": []})


def _normalize_severity(s):
    v = str(s or "").strip().upper()
    if v in ["P0", "P1", "P2"]:
        return v
    if "P0" in v:
        return "P0"
    if "P1" in v:
        return "P1"
    if "P2" in v:
        return "P2"
    raw = str(s or "").strip()
    if any(x in raw for x in ["A-严重", "A严重", "严重", "高危"]):
        return "P0"
    if any(x in raw for x in ["B-适中", "B适中", "中等", "中危"]):
        return "P1"
    if any(x in raw for x in ["C-微小", "C微小", "轻微", "低危"]):
        return "P2"
    return "P2"


def _normalize_frequency(s):
    v = str(s or "").strip()
    return v or "中"


def _tokenize_text(text):
    t = str(text or "").lower()
    parts = re.split(r"[^\w\u4e00-\u9fff]+", t)
    return [p for p in parts if p]


def _text_score(query, content):
    q = set(_tokenize_text(query))
    c = set(_tokenize_text(content))
    if not q or not c:
        return 0.0
    inter = len(q.intersection(c))
    union = len(q.union(c))
    s = inter / max(1, union)
    if str(query or "").strip() and str(query).strip() in str(content or ""):
        s += 0.2
    return round(min(1.0, s), 4)


def _next_rule_id(existing):
    mx = 0
    for it in existing:
        rid = str((it or {}).get("rule_id") or "")
        m = re.match(r"RBUG_(\d+)$", rid)
        if m:
            mx = max(mx, int(m.group(1)))
    return "RBUG_{:03d}".format(mx + 1)


def _load_bug_patterns():
    _ensure_bug_assets()
    data = _load_json_file(BUG_PATTERN_FILE, {"items": []})
    items = data.get("items") if isinstance(data, dict) else []
    return items if isinstance(items, list) else []


def _analyze_bug_desc(desc):
    text = str(desc or "").strip()
    if not text:
        return None
    patterns = _load_bug_patterns()
    hits = []
    for p in patterns:
        if not isinstance(p, dict):
            continue
        if p.get("enabled") is False:
            continue
        kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
        matched = [k for k in kws if str(k) and str(k) in text]
        if matched:
            hits.append((float(p.get("weight") or 0.5), p, matched))
    hits = sorted(hits, key=lambda x: x[0], reverse=True)
    if hits:
        _, top, matched = hits[0]
        return {
            "category": top.get("category") or "通用缺陷",
            "pattern_id": top.get("pattern_id") or "",
            "pattern": "、".join([str(x) for x in matched[:3]]),
            "design_gap": top.get("design_gap") if isinstance(top.get("design_gap"), list) else [],
            "rule": top.get("rule") or "关键流程必须定义异常兜底策略",
            "weight": float(top.get("weight") or 0.5),
        }
    return {
        "category": "通用缺陷",
        "pattern_id": "",
        "pattern": "未命中",
        "design_gap": ["异常处理缺失"],
        "rule": "关键流程必须定义异常兜底策略",
        "weight": 0.4,
    }


def _vector_search_local(query, top_k=5, type_filter=None, board_type=None):
    _ensure_bug_assets()
    data = _load_json_file(VECTOR_DATA_FILE, {"items": []})
    items = data.get("items") if isinstance(data, dict) else []
    if not isinstance(items, list):
        items = []
    out = []
    for it in items:
        if not isinstance(it, dict):
            continue
        if type_filter and str(it.get("type") or "") != str(type_filter):
            continue
        if board_type and str(it.get("board_type") or "") != str(board_type):
            continue
        content = str(it.get("content") or "")
        score = _text_score(query, content)
        if score <= 0:
            continue
        row = dict(it)
        row["score"] = score
        out.append(row)
    out.sort(key=lambda x: float(x.get("score") or 0), reverse=True)
    return out[: max(1, int(top_k or 5))]


def _load_knowledge_cards():
    if os.path.exists(KNOWLEDGE_CARDS_FILE):
        try:
            with open(KNOWLEDGE_CARDS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                items = _normalize_knowledge_items(data.get("items"))
            else:
                items = _normalize_knowledge_items(data)
            if items:
                return items
        except Exception as e:
            logger.warning("读取能力库文件失败: %s", e)
    return _normalize_knowledge_items(KTV_CAPABILITY_CARDS)


def _save_knowledge_cards(items):
    payload = {"items": _normalize_knowledge_items(items)}
    with open(KNOWLEDGE_CARDS_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return payload["items"]


def _build_xmind_tree(nodes):
    root = {
        "id": str(uuid.uuid4()),
        "title": "PRD功能清单",
        "children": {"attached": []},
    }
    stack = [root]
    for raw in nodes:
        title = str((raw or {}).get("title") or "").strip()
        if not title:
            continue
        level = int((raw or {}).get("level") or 1)
        if level < 1:
            level = 1
        if level > 6:
            level = 6
        while len(stack) > level:
            stack.pop()
        parent = stack[-1]
        topic = {"id": str(uuid.uuid4()), "title": title}
        parent.setdefault("children", {}).setdefault("attached", []).append(topic)
        stack.append(topic)
    return root


def _build_xmind_file(nodes):
    tree_root = _build_xmind_tree(nodes)
    content = [
        {
            "id": str(uuid.uuid4()),
            "class": "sheet",
            "title": "PRD功能导图",
            "rootTopic": tree_root,
        }
    ]
    metadata = {"creator": {"name": "Trae PRD Audit"}, "activeSheetId": content[0]["id"]}
    manifest = {"file-entries": {"content.json": {}, "metadata.json": {}}}

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content.json", json.dumps(content, ensure_ascii=False))
        zf.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False))
        zf.writestr("manifest.json", json.dumps(manifest, ensure_ascii=False))
    buf.seek(0)
    return buf


@prd_audit_bp.route("/")
def index():
    """独立 PRD 审计页"""
    return render_template("prd_audit_index.html")


@prd_audit_bp.route("/knowledge")
def knowledge_page():
    return render_template("prd_audit_knowledge.html")

@prd_audit_bp.route("/learning_mvp")
def learning_mvp_page():
    return render_template("prd_audit_learning_mvp.html")

@prd_audit_bp.route("/rules")
def rules_page():
    return render_template("prd_audit_rules.html")


@prd_audit_bp.route("/bug_patterns")
def bug_patterns_page():
    return render_template("prd_audit_bug_patterns.html")


@prd_audit_bp.route("/matrix_view")
def matrix_view_page():
    return render_template("prd_audit_matrix_view.html")


# ---------- API：生成（流式） ----------

@prd_audit_bp.route("/api/generate", methods=["POST"])
def api_generate():
    """PRD 三段式或单次调用生成（流式 NDJSON）"""
    try:
        data = request.get_json() or {}
        gen_type = data.get("type")
        content = data.get("content")
        use_llm = bool(data.get("use_llm", True))
        custom_prompt = data.get("prompt")
        report_level = (data.get("report_level") or "L3").upper()

        if not gen_type or gen_type != "prd_review":
            return error_response("type 必须为 prd_review", status_code=400)
        if not content or (isinstance(content, str) and not content.strip()):
            return error_response("content 不能为空", status_code=400)

        content = content.strip() if isinstance(content, str) else content

        if isinstance(content, str) and is_feishu_doc_url(content):
            ok, result = fetch_feishu_doc_content(content)
            if not ok:
                return error_response(
                    result if result else "无法拉取飞书文档，请复制内容粘贴或导出后上传",
                    status_code=400,
                )
            content = result
            logger.info("Feishu doc content fetched, length=%s", len(content))

        # 是否启用大模型：前端可选。use_llm = False 时强制使用本地规则体检模式。
        from utils.llm_client import load_llm_config
        config_path = _pick_llm_config_path()
        if use_llm:
            try:
                llm_config = load_llm_config(config_path)
                api_key = (llm_config.get("api_key") or "").strip()
                if not api_key:
                    logger.warning("LLM API Key 未配置，将使用本地规则体检模式：%s", config_path)
                    config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")
            except FileNotFoundError:
                logger.warning("LLM 配置文件不存在，将使用本地规则体检模式：%s", config_path)
                config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")
            except Exception as e:
                logger.warning("加载 LLM 配置失败，将使用本地规则体检模式: %s", e)
                config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")
        else:
            # 显式禁用 LLM：传入一个不存在的配置路径，触发 rule-first/offline 模式
            logger.info("PRD 审计本次调用被配置为不使用大模型（use_llm=false），将仅使用本地规则体检模式。")
            config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

        def generate():
            try:
                for chunk in pipeline.run_prd_audit_stream(
                    content,
                    llm_config_path=config_path,
                    timeout=180,
                    custom_prompt=custom_prompt.strip() if custom_prompt and custom_prompt.strip() else None,
                    report_level=report_level,
                ):
                    yield chunk
            except Exception as e:
                logger.exception("PRD generate failed")
                yield json.dumps({"type": "error", "text": str(e)}, ensure_ascii=False) + "\n"

        return Response(
            stream_with_context(generate()),
            content_type="application/x-ndjson; charset=utf-8",
        )
    except Exception as e:
        logger.exception("api_generate failed")
        return error_response(str(e), status_code=500)


# ---------- API：大纲专用（大模型，通用 PRD） ----------

@prd_audit_bp.route("/api/outline_llm", methods=["POST"])
def api_outline_llm():
    """仅生成「四支柱」认知大纲 JSON，适配任意 PRD；可选合并本地 outline_engine。"""
    try:
        data = request.get_json() or {}
        prd_text = (data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        merge_local = bool(data.get("merge_local", False))
        run_stage1 = data.get("run_stage1", True)
        if isinstance(run_stage1, str):
            run_stage1 = run_stage1.strip().lower() not in ("0", "false", "no", "")

        stage1_override = data.get("stage1_output")
        timeout = int(data.get("timeout", 120) or 120)
        if timeout < 30:
            timeout = 30
        if timeout > 600:
            timeout = 600

        from utils.llm_client import load_llm_config

        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
        except FileNotFoundError:
            return error_response("API Key 未配置（请编辑 modules/prd_audit_clone/llm_config.json）", status_code=400)
        if not (llm_config.get("api_key") or "").strip():
            return error_response("API Key 未配置（请编辑 modules/prd_audit_clone/llm_config.json）", status_code=400)

        from .system_model import extract_prd_structure

        if run_stage1:
            stage1 = extract_prd_structure(
                prd_text,
                llm_config_path=config_path,
                timeout=min(90, timeout),
            )
        else:
            stage1 = stage1_override if isinstance(stage1_override, dict) else {}
            if not stage1:
                return error_response("run_stage1=false 时需传入非空 stage1_output", status_code=400)

        llm_result = run_outline_llm(
            prd_text,
            stage1,
            config_path,
            timeout=timeout,
        )

        payload = {
            "llm": llm_result,
            "stage1_output": stage1,
        }
        if merge_local:
            local_engine = run_outline_engine(prd_text, stage1, {})
            payload["local_outline_engine"] = local_engine
            payload["merged"] = merge_llm_with_local(
                llm_result.get("llm_outline") or {},
                local_engine,
            )

        return success_response(
            data=payload,
            message="大纲生成完成" if llm_result.get("ok") else (llm_result.get("error") or "大纲生成未完全成功"),
        )
    except Exception as e:
        logger.exception("api_outline_llm failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/analyze_impact", methods=["POST"])
def api_analyze_impact():
    """架构级变更影响分析"""
    try:
        data = request.get_json() or {}
        change_desc = (data.get("change_desc") or "").strip()
        scan_result = data.get("scan_result")

        if not change_desc:
            return error_response("需求变更描述不能为空", status_code=400)
        if not scan_result or not isinstance(scan_result, dict):
            return error_response("缺少架构透视扫描数据", status_code=400)

        from .architecture_scanner import analyze_impact
        impact_result = analyze_impact(change_desc, scan_result)
        
        affected_modules = impact_result.get("affected_modules", [])
        affected_states = impact_result.get("affected_states", [])
        affected_apis = impact_result.get("affected_apis", [])
        affected_entities = impact_result.get("affected_entities", [])
        risk_assessment = impact_result.get("risk_assessment", {})

        # --- 将本地正则分析结果交由 LLM 进行二次润色与智能推理 ---
        from utils.llm_client import call_llm
        import json
        
        # 提取本地分析的上下文，喂给大模型
        context_data = {
            "change_desc": change_desc,
            "local_analysis": impact_result
        }
        
        prompt = f"""你是一个资深的软件架构师和测试专家。
用户提出了一个需求变更（一句话描述）。我已经通过代码/文档解析引擎，提取了系统当前的模块、状态机、API和实体数据，并做了一次初步的影响分析。

【用户变更需求】：
{change_desc}

【本地初步分析结果】：
{json.dumps(context_data['local_analysis'], ensure_ascii=False, indent=2)}

请你基于以上信息，结合你对软件工程的理解，输出一份**更智能、更具可读性、且具有深度业务推理**的《变更影响分析报告》。
注意：
1. 本地分析可能会漏掉一些隐含的业务逻辑关联（比如“星耀屏广告”可能与“投屏”、“点歌”存在抢占关系），请发挥你的推理能力进行补充。
2. 即使本地分析显示“None”或“无明显影响”，如果你认为在实际业务中可能会有影响，请明确指出并解释原因。
3. 请直接输出 Markdown 格式的报告。包含以下部分：
   - 📌 变更核心逻辑解读
   - 🎯 综合风险评估（整体影响、风险等级）
   - 🧩 受影响的模块与接口（说明影响原因）
   - 🔄 受影响的状态机与数据流
   - 🧪 回归测试策略建议与高风险点

输出报告："""

        try:
            llm_report = call_llm(messages=[{"role": "user", "content": prompt}])
            return success_response({"impact_report": llm_report})
        except Exception as e:
            logger.exception("LLM impact analysis failed")
            # 降级：如果大模型调用失败，退回到本地拼接的简单报告
            md = [f"### 变更影响分析报告：{change_desc}\n"]
            
            md.append("#### 1. 综合风险评估")
            md.append(f"- **整体影响范围**: {risk_assessment.get('overall_impact', '未知')}")
            md.append(f"- **回归测试建议级别**: {risk_assessment.get('regression_level', '未知')}")
            md.append(f"- **高风险点**: {risk_assessment.get('high_risk_points', 0)} 个")
            md.append("\n**测试策略建议**:")
            for advice in risk_assessment.get("testing_advice", []):
                md.append(f"- {advice}")
                
            md.append("\n#### 2. 受影响的模块")
            if affected_modules:
                for m in affected_modules:
                    md.append(f"- **{m.get('name')}** (影响: {m.get('impact_scope')}, 风险: {m.get('risk')})")
                    if m.get('dependencies'):
                        md.append(f"  - 依赖方: {', '.join(m.get('dependencies'))}")
            else:
                md.append("- 无明显受影响的模块")

            md.append("\n#### 3. 受影响的状态/流程")
            if affected_states:
                for s in affected_states:
                    md.append(f"- **{s.get('state')}** (模块: {s.get('module')}, 影响: {s.get('impact_scope')})")
                    for t in s.get('related_transitions', []):
                        md.append(f"  - 关联跃迁: {t}")
            else:
                md.append("- 无明显受影响的状态")

            md.append("\n#### 4. 受影响的 API 接口")
            if affected_apis:
                for a in affected_apis:
                    md.append(f"- **{a.get('api')}** (影响: {a.get('impact_scope')})")
                    md.append(f"  - 调用关系: {a.get('caller')} -> {a.get('callee')}")
            else:
                md.append("- 无明显受影响的 API")

            md.append("\n#### 5. 受影响的数据实体")
            if affected_entities:
                for e in affected_entities:
                    md.append(f"- **{e.get('entity')}** (影响: {e.get('impact_scope')})")
                    md.append(f"  - 变更动作: {e.get('action')}")
                    if e.get('related_fields'):
                        md.append(f"  - 关联字段: {', '.join(e.get('related_fields'))}")
            else:
                md.append("- 无明显受影响的数据实体")

            return success_response({
                "impact_report": "\n".join(md),
                "raw_data": impact_result
            })
    except Exception as e:
        logger.exception("api_analyze_impact failed")
        return error_response(str(e), status_code=500)


# ---------- API：分析（同步，返回完整 JSON） ----------

@prd_audit_bp.route("/api/analyze_prd", methods=["POST"])
def api_analyze_prd():
    """非流式 PRD 分析，返回 raw_report_markdown、stage1/2/3、summary 等"""
    try:
        data = request.get_json() or {}
        prd_text = (data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        from utils.llm_client import load_llm_config
        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
        except FileNotFoundError:
            return error_response("API Key 未配置（请编辑 modules/prd_audit_clone/llm_config.json）", status_code=400)
        if not (llm_config.get("api_key") or "").strip():
            return error_response("API Key 未配置（请编辑 modules/prd_audit_clone/llm_config.json）", status_code=400)

        merged_report, stage1_output, stage2_output, stage3_output = pipeline.run_prd_audit_sync(
            prd_text, llm_config_path=config_path, timeout=180
        )
        try:
            from .test_matrix_generator import TestMatrixGenerator
            test_matrix = TestMatrixGenerator(stage1_output, stage2_output).generate()
        except Exception:
            logger.exception("api_analyze_prd: generate test_matrix failed")
            test_matrix = {}
        try:
            from .test_points_engine import run_test_points_engine
            test_points_obj = run_test_points_engine(
                prd_text=prd_text,
                stage1_output=stage1_output if isinstance(stage1_output, dict) else {},
                stage2_output=stage2_output if isinstance(stage2_output, dict) else {},
                outline_engine={},
                platform_impact={},
                dependency_analysis={},
                test_matrix=test_matrix if isinstance(test_matrix, dict) else {},
            )
        except Exception:
            logger.exception("api_analyze_prd: generate test_points failed")
            test_points_obj = {}

        score = stage3_output.get("summary", {}).get("quality_score")
        reader_guide = pipeline._build_reader_guide(
            stage1_output if isinstance(stage1_output, dict) else {},
            stage3_output if isinstance(stage3_output, dict) else {},
        )
        quality_summary = {
            "overall": score,
            "dimensions": {},
            "rule_engine_score": score,
        }
        states = pipeline._ensure_list(stage1_output.get("states"))
        state_diagram_mermaid = ""
        if states:
            cleaned = [s.replace(" ", "_") for s in states]
            lines = ["stateDiagram-v2", f"  [*] --> {cleaned[0]}"]
            for i in range(len(cleaned) - 1):
                lines.append(f"  {cleaned[i]} --> {cleaned[i + 1]}")
            state_diagram_mermaid = "\n".join(lines)
        defects = stage3_output.get("defects") or []
        prd_gaps = "\n".join([f"- {d.get('module', '【PRD未说明】')}：{d.get('description', '【PRD未说明】')}" for d in defects[:20]])
        risks = "\n".join([f"- {x}" for x in (stage3_output.get("risks") or [])])
        rd_focus = "\n".join([f"- {x}" for x in (stage3_output.get("dev_focus") or [])])
        test_focus = "\n".join([f"- {x}" for x in (stage3_output.get("test_focus") or [])])
        plan_list = stage3_output.get("plan") or []
        plan_level = "P1"
        if any((d.get("risk_level") or "").upper() == "P0" for d in defects):
            plan_level = "P0"
        elif defects and all((d.get("risk_level") or "").upper() == "P2" for d in defects):
            plan_level = "P2"
        plan = {
            "level": plan_level,
            "order_advice": plan_list[:2],
            "prep_advice": plan_list[2:4],
        }
        section_1 = "## 一、总体结论\n\n- 质量评分：{}/10\n- 风险等级：{}\n- 主要问题：{}".format(
            stage3_output.get("summary", {}).get("quality_score", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("risk_level", "【PRD未说明】"),
            stage3_output.get("summary", {}).get("main_problem", "【PRD未说明】"),
        )
        section_2 = "## 二、漏洞与风险清单\n\n" + (prd_gaps or "- 未发现漏洞")
        section_4 = "## 四、测试重点\n\n" + (test_focus or "- 【PRD未说明】")
        section_5 = "## 五、研发重点\n\n" + (rd_focus or "- 【PRD未说明】")
        section_6 = "## 六、计划建议\n\n" + ("\n".join([f"- {x}" for x in plan_list]) if plan_list else "- 【PRD未说明】")
        result = {
            "raw_report_markdown": merged_report,
            "status": ["Stage1完成", "Stage2完成", "Stage3完成", "Stage4完成"],
            "summary": section_1,
            "modules": stage1_output.get("modules") or [],
            "prd_gaps": prd_gaps,
            "risks": risks,
            "rd_focus": rd_focus,
            "test_focus": test_focus,
            "test_points": test_points_obj,
            "plan": plan,
            "quality": quality_summary,
            "stage1_output": stage1_output,
            "stage2_output": stage2_output,
            "stage3_output": stage3_output,
            "reader_guide": reader_guide if isinstance(reader_guide, dict) else {},
            "test_matrix": test_matrix if isinstance(test_matrix, dict) else {},
            "system_model": {
                "states": stage1_output.get("states") or [],
                "events": [],
                "rules": stage1_output.get("business_rules") or [],
                "data_rules": stage1_output.get("exceptions") or [],
                "data_structures": stage1_output.get("data_structures") or [],
                "permissions": stage1_output.get("permissions") or [],
                "edge_cases": stage1_output.get("edge_cases") or [],
                "dependencies": stage1_output.get("dependencies") or [],
                "transitions": [],
                "high_priority_events": [],
            },
            "rule_analysis": stage2_output,
            "test_point_matrix": (test_points_obj.get("test_point_matrix") if isinstance(test_points_obj, dict) else {}) or {"version": "v1", "items": [], "stats": {"total": 0, "p0": 0, "p1": 0, "p2": 0, "by_module": {}}},
            "sections": {
                "1": section_1,
                "2": section_2,
                "4": section_4,
                "5": section_5,
                "6": section_6,
            },
            "state_diagram_mermaid": state_diagram_mermaid,
        }
        return success_response(data=result)
    except Exception as e:
        logger.exception("api_analyze_prd failed")
        return error_response(str(e), status_code=500)


def _evaluate_bugs_for_preview(items: list) -> list:
    """对即将导入的 Bug 列表进行 AI 预审"""
    if not items:
        return items

    # 准备 LLM 评估的 Prompt
    bug_texts = []
    for i, it in enumerate(items):
        desc = str(it.get("bug_desc", "")).strip()
        sev = str(it.get("severity", "P2")).strip()
        bug_texts.append(f"[{i}] 描述: {desc} | 严重度: {sev}")

    prompt = f"""你是一个高级测试架构师，正在审查一批准备导入到“PRD 审计系统”的 Bug 缺陷记录。
我们的目标是通过这些历史 Bug 来训练系统的 PRD 审计规则，使其能发现 PRD 中“缺失的业务逻辑、异常处理、边界条件和状态流转等设计缺陷”。

请逐一评估以下 Bug 是否有导入价值，并严格按以下标准分类：
1. "推荐导入": 逻辑漏洞、状态流转错误、并发冲突、异常流程未闭环、越权等设计/场景缺失类 Bug。
2. "建议修正": 有价值，但描述过于简略，缺少上下文或具体表现，建议补充完整场景后导入。
3. "建议废弃": 纯 UI/样式细节偏差、拼写错误、环境配置问题、明显的代码手误等（对 PRD 约束意义不大的低级错误）。

输入 Bug 列表：
{chr(10).join(bug_texts)}

请返回 JSON 数组，每个元素包含：
- "index": 输入的序号（整数）
- "ai_status": "推荐导入" 或 "建议修正" 或 "建议废弃"
- "ai_reason": 简短的评估理由（15字以内）

只输出合法的 JSON 数组，不要包含任何 markdown 代码块标记，不要多余说明。"""

    # 回退到本地规则的辅助函数
    def _fallback_eval(bug_item):
        desc = bug_item.get("bug_desc", "").lower()
        if any(kw in desc for kw in ["ui", "颜色", "对其", "拼写", "环境", "配置", "错别字", "字体", "间距", "样式"]):
            return {"ai_status": "建议废弃", "ai_reason": "UI/环境类无关PRD"}
        elif len(desc) < 6:
            return {"ai_status": "建议修正", "ai_reason": "描述过于简略"}
        else:
            return {"ai_status": "推荐导入", "ai_reason": "本地规则通过"}

    try:
        config_path = _pick_llm_config_path()
        resp = call_llm([{"role": "user", "content": prompt}], config_path=config_path, timeout=30)
        resp_text = str(resp or "").strip()
        # 清理可能存在的 markdown 代码块
        if resp_text.startswith("```json"):
            resp_text = resp_text[7:]
        if resp_text.startswith("```"):
            resp_text = resp_text[3:]
        if resp_text.endswith("```"):
            resp_text = resp_text[:-3]
        resp_text = resp_text.strip()
        
        eval_list = json.loads(resp_text)
        if isinstance(eval_list, list):
            eval_map = {int(x.get("index", -1)): x for x in eval_list if isinstance(x, dict)}
            for i, it in enumerate(items):
                ev = eval_map.get(i)
                if ev:
                    it["ai_status"] = ev.get("ai_status", "推荐导入")
                    it["ai_reason"] = ev.get("ai_reason", "AI认为有价值")
                else:
                    fb = _fallback_eval(it)
                    it.update(fb)
        else:
            raise ValueError("LLM 返回的不是 JSON 数组")
    except Exception as e:
        logger.warning(f"AI 预审 Bug 失败，回退到本地启发式规则: {e}")
        for it in items:
            fb = _fallback_eval(it)
            it.update(fb)
            
    return items


@prd_audit_bp.route("/api/bug/import", methods=["POST"])
def api_bug_import():
    try:
        _ensure_bug_assets()
        is_preview = request.args.get("preview", "false").lower() == "true"
        items = []
        if "file" in request.files:
            f = request.files["file"]
            raw = f.read()
            txt = raw.decode("utf-8-sig", errors="ignore")
            sample = "\n".join(txt.splitlines()[:3])
            delimiter = "\t" if ("\t" in sample and sample.count("\t") >= sample.count(",")) else ","
            reader = csv.DictReader(io.StringIO(txt), delimiter=delimiter)
            for row in reader:
                if not isinstance(row, dict):
                    continue
                bug_desc = str(
                    row.get("bug_desc")
                    or row.get("desc")
                    or row.get("概要")
                    or row.get("问题描述")
                    or row.get("summary")
                    or ""
                ).strip()
                if not bug_desc:
                    continue
                sev = (
                    row.get("severity")
                    or row.get("自定义字段(严重程度)")
                    or row.get("严重程度")
                    or row.get("等级")
                    or row.get("自定义字段(优先级)")
                    or row.get("优先级")
                )
                items.append(
                    {
                        "bug_desc": bug_desc,
                        "severity": _normalize_severity(sev),
                        "frequency": _normalize_frequency(row.get("frequency")),
                    }
                )
        else:
            data = request.get_json() or {}
            if isinstance(data.get("items"), list):
                for row in data.get("items"):
                    if not isinstance(row, dict):
                        continue
                    bug_desc = str(row.get("bug_desc") or "").strip()
                    if not bug_desc:
                        continue
                    items.append(
                        {
                            "bug_desc": bug_desc,
                            "severity": _normalize_severity(row.get("severity")),
                            "frequency": _normalize_frequency(row.get("frequency")),
                        }
                    )
            text_blob = str(data.get("text") or "").strip()
            if text_blob:
                # 尝试检测是否为 Jira 样式的 TSV/CSV 文本
                # 特征：第一行包含 "问题关键字" 或 "概要" 且含有 tab 或 多个空格
                first_line = text_blob.splitlines()[0] if text_blob else ""
                is_jira_format = False
                delimiter = None
                
                if "问题关键字" in first_line or "概要" in first_line:
                    if "\t" in first_line:
                        delimiter = "\t"
                        is_jira_format = True
                    elif "  " in first_line: # 可能是多个空格分隔
                        # 尝试用正则替换多个空格为 tab，或者直接解析
                        is_jira_format = True
                        delimiter = "regex_spaces"

                if is_jira_format:
                    lines = text_blob.splitlines()
                    # 如果是正则空格分隔，先处理一下 header
                    if delimiter == "regex_spaces":
                        # 简单策略：每一行都用正则 split，然后取对应 index
                        # 这是一个比较脆弱的策略，但对于复制粘贴的文本可能有效
                        # 更好的方式是看 header 的列位置
                        header_parts = re.split(r'\s{2,}|\t', lines[0].strip())
                        # 寻找 "概要" 和 "自定义字段(严重程度)" / "自定义字段(优先级)" 的 index
                        summary_idx = -1
                        severity_idx = -1
                        priority_idx = -1
                        
                        for idx, col in enumerate(header_parts):
                            c = col.strip()
                            if c == "概要": summary_idx = idx
                            if "严重程度" in c: severity_idx = idx
                            if "优先级" in c: priority_idx = idx
                        
                        if summary_idx != -1:
                            for line in lines[1:]:
                                if not line.strip(): continue
                                parts = re.split(r'\s{2,}|\t', line.strip())
                                # Parts 长度可能不一致，因为有些列可能为空导致合并
                                # 这里只能尽力而为。如果 parts 长度不够，可能无法准确提取
                                # 但用户给的例子看起来对齐得还不错
                                if len(parts) > summary_idx:
                                    bug_desc = parts[summary_idx].strip()
                                    sev = "P2"
                                    if severity_idx != -1 and len(parts) > severity_idx:
                                        sev = parts[severity_idx].strip()
                                    if priority_idx != -1 and len(parts) > priority_idx:
                                        # 优先级可能覆盖严重程度
                                        p = parts[priority_idx].strip()
                                        if p: sev = p
                                    
                                    if bug_desc:
                                        items.append({
                                            "bug_desc": bug_desc, 
                                            "severity": _normalize_severity(sev), 
                                            "frequency": "中"
                                        })
                    else:
                        # 标准 TSV 解析
                        try:
                            reader = csv.DictReader(io.StringIO(text_blob), delimiter=delimiter)
                            for row in reader:
                                bug_desc = str(
                                    row.get("bug_desc")
                                    or row.get("desc")
                                    or row.get("概要")
                                    or row.get("问题描述")
                                    or row.get("summary")
                                    or ""
                                ).strip()
                                if not bug_desc: continue
                                sev = (
                                    row.get("severity")
                                    or row.get("自定义字段(严重程度)")
                                    or row.get("严重程度")
                                    or row.get("等级")
                                    or row.get("自定义字段(优先级)")
                                    or row.get("优先级")
                                )
                                items.append({
                                    "bug_desc": bug_desc,
                                    "severity": _normalize_severity(sev),
                                    "frequency": "中"
                                })
                        except Exception:
                            # 解析失败回退到普通按行处理
                            pass

                # 如果没有被识别为 Jira 格式，或者解析后 items 仍为空（解析失败），则回退到逐行处理
                if not items:
                    for line in text_blob.splitlines():
                        d = str(line or "").strip()
                        if d:
                            items.append({"bug_desc": d, "severity": "P2", "frequency": "中"})

        if not items:
            return error_response("未解析到可导入 Bug", status_code=400)

        if is_preview:
            evaluated_items = _evaluate_bugs_for_preview(items)
            return success_response(data={"items": evaluated_items})

        raw_data = _load_json_file(BUG_RAW_FILE, {"items": []})
        raw_items = raw_data.get("items") if isinstance(raw_data, dict) else []
        if not isinstance(raw_items, list):
            raw_items = []
        vec_data = _load_json_file(VECTOR_DATA_FILE, {"items": []})
        vec_items = vec_data.get("items") if isinstance(vec_data, dict) else []
        if not isinstance(vec_items, list):
            vec_items = []
        imported = []
        for it in items:
            row = {
                "id": "B_" + uuid.uuid4().hex[:8],
                "bug_desc": it["bug_desc"],
                "severity": _normalize_severity(it.get("severity")),
                "frequency": _normalize_frequency(it.get("frequency")),
            }
            raw_items.append(row)
            imported.append(row)
            vec_items.append(
                {
                    "id": "V_" + uuid.uuid4().hex[:8],
                    "content": row["bug_desc"],
                    "type": "bug",
                    "board_type": "",
                    "ref_id": row["id"],
                }
            )
        _save_json_file(BUG_RAW_FILE, {"items": raw_items})
        _save_json_file(VECTOR_DATA_FILE, {"items": vec_items})
        return success_response(data={"imported_count": len(imported), "items": imported})
    except Exception as e:
        logger.exception("api_bug_import failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/patterns", methods=["GET", "POST"])
def api_bug_patterns():
    try:
        _ensure_bug_assets()
        if request.method == "GET":
            return success_response(data={"items": _load_bug_patterns()})
        data = request.get_json() or {}
        items = data.get("items") if isinstance(data, dict) else None
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)
        cleaned = []
        for i, p in enumerate(items):
            if not isinstance(p, dict):
                continue
            pid = str(p.get("pattern_id") or "").strip() or "P{:03d}".format(i + 1)
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            kws = [str(x).strip() for x in kws if str(x).strip()]
            if not kws:
                continue
            cleaned.append(
                {
                    "pattern_id": pid,
                    "keywords": kws,
                    "category": str(p.get("category") or "通用缺陷").strip(),
                    "design_gap": p.get("design_gap") if isinstance(p.get("design_gap"), list) else [],
                    "rule": str(p.get("rule") or "").strip() or "关键流程必须定义异常兜底策略",
                    "weight": float(p.get("weight") or 0.5),
                    "enabled": bool(p.get("enabled", True)),
                }
            )
        _save_json_file(BUG_PATTERN_FILE, {"items": cleaned})
        return success_response(data={"items": cleaned, "count": len(cleaned)})
    except Exception as e:
        logger.exception("api_bug_patterns failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/patterns/auto_from_bugs", methods=["POST"])
def api_bug_patterns_auto_from_bugs():
    try:
        _ensure_bug_assets()
        data = request.get_json() or {}
        items = data.get("items") if isinstance(data, dict) else None
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)

        bug_descs = []
        for it in items:
            if not isinstance(it, dict):
                continue
            d = str(it.get("bug_desc") or "").strip()
            if d:
                bug_descs.append(d)
        if not bug_descs:
            return error_response("未提供可用于生成模式的 bug_desc", status_code=400)

        existing = _load_bug_patterns()
        used_keywords = set()
        max_pid = 0
        for p in existing:
            if not isinstance(p, dict):
                continue
            pid = str(p.get("pattern_id") or "")
            m = re.match(r"P(\d+)$", pid)
            if m:
                max_pid = max(max_pid, int(m.group(1)))
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            for k in kws:
                ks = str(k).strip()
                if ks:
                    used_keywords.add(ks)

        stopwords = {
            "问题", "这个", "那个", "无法", "不能", "没有", "出现", "导致", "以及",
            "时候", "之后", "过程", "进行", "相关", "功能", "页面", "系统", "当前",
            "仍然", "展示", "可以", "一下", "或者", "然后", "其中", "这里", "那里",
            "用户", "业务", "场景", "模块", "逻辑", "状态", "操作", "流程",
            "投屏", "广告", "语音", "自助", "商k", "k歌", "导唱", "打分", "录音", "量贩",
            "第三屏", "竖版",
        }
        risk_terms = [
            "异常", "失败", "报错", "崩溃", "闪退", "卡死", "无响应", "黑屏",
            "不一致", "不同步", "错乱", "冲突", "越权", "超时", "丢失", "重叠",
            "不消失", "不能", "无法", "无反馈", "无应答",
        ]
        keyword_counter = Counter()

        for desc in bug_descs:
            parts = re.split(r"[，,。；;、\s]+", desc)
            candidates = []
            for part in parts:
                s = str(part).strip()
                if not s:
                    continue
                if len(s) < 2 or len(s) > 12:
                    continue
                if s in stopwords:
                    continue
                s_lower = s.lower()
                if s_lower in {"pk", "x9", "vx", "bug", "p0", "p1", "p2"}:
                    continue
                if re.fullmatch(r"[a-z]{1,3}\d{1,4}", s_lower):
                    continue
                if re.fullmatch(r"[a-z]{1,4}-\d{1,6}", s_lower):
                    continue
                if re.fullmatch(r"p\d{1,2}", s_lower):
                    continue
                if re.fullmatch(r"[a-z]{1,2}", s_lower):
                    continue
                if not any(t in s for t in risk_terms):
                    continue
                candidates.append(s)
            for kw in candidates[:3]:
                keyword_counter[kw] += 1

        top_keywords = [kw for kw, _ in keyword_counter.most_common(20)]
        added = []
        for kw in top_keywords:
            if kw in used_keywords:
                continue
            related = [d for d in bug_descs if kw in d]
            if not related:
                continue
            text_all = "\n".join(related)
            if any(x in text_all for x in ["唤醒", "语音", "识别", "应答"]):
                category = "语音交互"
                design_gap = ["状态流转缺失", "反馈机制缺失"]
                rule = "语音交互必须定义唤醒-识别-反馈闭环与异常兜底"
                weight = 0.9
            elif any(x in text_all for x in ["卡死", "无响应", "崩溃", "闪退"]):
                category = "稳定性"
                design_gap = ["超时恢复缺失", "异常处理缺失"]
                rule = "关键交互必须定义超时恢复、重试与降级策略"
                weight = 0.88
            elif any(x in text_all for x in ["不一致", "不同步", "错乱"]):
                category = "状态一致性"
                design_gap = ["状态机缺失"]
                rule = "多模块交互必须定义状态一致性与冲突裁决规则"
                weight = 0.86
            else:
                category = "通用缺陷"
                design_gap = ["异常处理缺失"]
                rule = "关键流程必须定义异常处理与回退策略"
                weight = 0.72

            max_pid += 1
            row = {
                "pattern_id": "P{:03d}".format(max_pid),
                "keywords": [kw],
                "category": category,
                "design_gap": design_gap,
                "rule": rule,
                "weight": weight,
                "enabled": True,
            }
            existing.append(row)
            added.append(row)
            used_keywords.add(kw)
            if len(added) >= 12:
                break

        _save_json_file(BUG_PATTERN_FILE, {"items": existing})
        return success_response(data={"added_count": len(added), "items": added})
    except Exception as e:
        logger.exception("api_bug_patterns_auto_from_bugs failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/analyze", methods=["POST"])
def api_bug_analyze():
    try:
        _ensure_bug_assets()
        data = request.get_json() or {}
        bug_inputs = []
        if isinstance(data.get("bugs"), list):
            for x in data.get("bugs"):
                if isinstance(x, dict):
                    d = str(x.get("bug_desc") or "").strip()
                else:
                    d = str(x or "").strip()
                if d:
                    bug_inputs.append(d)
        one = str(data.get("bug_desc") or "").strip()
        if one:
            bug_inputs.append(one)
        if not bug_inputs:
            return error_response("bug_desc 不能为空", status_code=400)

        analyses = []
        rules_data = _load_json_file(BUG_RULE_FILE, {"items": []})
        rules = rules_data.get("items") if isinstance(rules_data, dict) else []
        if not isinstance(rules, list):
            rules = []
        for d in bug_inputs:
            parsed = _analyze_bug_desc(d)
            if not parsed:
                continue
            analyses.append({"bug_desc": d, **parsed})
            rid = _next_rule_id(rules)
            rules.append(
                {
                    "rule_id": rid,
                    "rule_desc": parsed.get("rule") or "",
                    "source_pattern": parsed.get("pattern_id") or "",
                    "severity": "P0" if float(parsed.get("weight") or 0) >= 0.9 else ("P1" if float(parsed.get("weight") or 0) >= 0.75 else "P2"),
                }
            )
        _save_json_file(BUG_RULE_FILE, {"items": rules})
        return success_response(data={"items": analyses, "generated_rules": rules[-len(analyses):] if analyses else []})
    except Exception as e:
        logger.exception("api_bug_analyze failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/rules", methods=["GET", "POST"])
def api_bug_rules():
    try:
        _ensure_bug_assets()
        if request.method == "POST":
            data = request.get_json() or {}
            items = data.get("items") if isinstance(data, dict) else None
            if not isinstance(items, list):
                return error_response("items 必须为数组", status_code=400)
            cleaned = []
            for i, it in enumerate(items):
                if not isinstance(it, dict):
                    continue
                rid = str(it.get("rule_id") or "").strip() or "RBUG_{:03d}".format(i + 1)
                cleaned.append(
                    {
                        "rule_id": rid,
                        "category": str(it.get("category") or "General").strip() or "General",
                        "rule_desc": str(it.get("rule_desc") or "").strip() or "关键流程需补充规则",
                        "source_pattern": str(it.get("source_pattern") or "").strip(),
                        "severity": _normalize_severity(it.get("severity")),
                        "detect": str(it.get("detect") or "").strip(),
                        "risk": str(it.get("risk") or "").strip(),
                    }
                )
            _save_json_file(BUG_RULE_FILE, {"items": cleaned})
            return success_response(data={"items": cleaned, "count": len(cleaned)}, message="规则库已保存")
        data = _load_json_file(BUG_RULE_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        return success_response(data={"items": items, "count": len(items)})
    except Exception as e:
        logger.exception("api_bug_rules failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/rules/export", methods=["GET"])
def api_bug_rules_export():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RULE_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_rules.json",
        )
    except Exception as e:
        logger.exception("api_bug_rules_export failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/patterns/export", methods=["GET"])
def api_bug_patterns_export():
    try:
        _ensure_bug_assets()
        items = _load_bug_patterns()
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_patterns.json",
        )
    except Exception as e:
        logger.exception("api_bug_patterns_export failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/raw/export", methods=["GET"])
def api_bug_raw_export():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RAW_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        buf = io.BytesIO()
        buf.write(json.dumps({"items": items}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="bug_raw.json",
        )
    except Exception as e:
        logger.exception("api_bug_raw_export failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/import_template_csv", methods=["GET"])
def api_bug_import_template_csv():
    try:
        sio = io.StringIO()
        writer = csv.writer(sio)
        writer.writerow(["bug_desc", "severity", "frequency"])
        writer.writerow(["投屏过程中黑屏", "P0", "高"])
        writer.writerow(["广告切换后卡死", "P1", "中"])
        writer.writerow(["弱网下播放无响应", "P1", "中"])
        writer.writerow(["切歌后偶发无声", "P2", "低"])
        content = sio.getvalue().encode("utf-8-sig")
        buf = io.BytesIO(content)
        buf.seek(0)
        return send_file(
            buf,
            mimetype="text/csv",
            as_attachment=True,
            download_name="bug_import_template.csv",
        )
    except Exception as e:
        logger.exception("api_bug_import_template_csv failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/bug/raw", methods=["GET"])
def api_bug_raw():
    try:
        _ensure_bug_assets()
        data = _load_json_file(BUG_RAW_FILE, {"items": []})
        items = data.get("items") if isinstance(data, dict) else []
        if not isinstance(items, list):
            items = []
        return success_response(data={"items": items, "count": len(items)})
    except Exception as e:
        logger.exception("api_bug_raw failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/vector/search", methods=["POST"])
def api_vector_search():
    try:
        data = request.get_json() or {}
        query = str(data.get("query") or "").strip()
        if not query:
            return error_response("query 不能为空", status_code=400)
        top_k = int(data.get("top_k") or 5)
        top_k = 1 if top_k < 1 else (20 if top_k > 20 else top_k)
        matches = _vector_search_local(
            query=query,
            top_k=top_k,
            type_filter=(data.get("type") or None),
            board_type=(data.get("board_type") or None),
        )
        return success_response(data={"query": query, "items": matches})
    except Exception as e:
        logger.exception("api_vector_search failed")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/prd/audit", methods=["POST"])
def api_prd_audit():
    try:
        data = request.get_json() or {}
        prd_text = str(data.get("prd_text") or data.get("content") or "").strip()
        if not prd_text:
            return error_response("prd_text 不能为空", status_code=400)

        config_path = _pick_llm_config_path()
        use_llm = bool(data.get("use_llm", False))
        if not use_llm:
            config_path = os.path.join(STORAGE_DIR, "__llm_disabled__.json")

        last_obj = {}
        for chunk in pipeline.run_prd_audit_stream(prd_text, llm_config_path=config_path, timeout=180):
            try:
                obj = json.loads(chunk)
            except Exception:
                continue
            if isinstance(obj, dict) and "L3" in obj:
                last_obj = obj
        if not last_obj:
            return error_response("审计失败，未生成结果", status_code=500)

        patterns = _load_bug_patterns()
        bug_hits = []
        for p in patterns:
            if not isinstance(p, dict) or p.get("enabled") is False:
                continue
            kws = p.get("keywords") if isinstance(p.get("keywords"), list) else []
            matched = [k for k in kws if str(k) and str(k) in prd_text]
            if matched:
                bug_hits.append(
                    {
                        "pattern_id": p.get("pattern_id"),
                        "category": p.get("category"),
                        "matched_keywords": matched,
                        "rule": p.get("rule"),
                        "weight": p.get("weight"),
                    }
                )
        bug_hits.sort(key=lambda x: float(x.get("weight") or 0), reverse=True)
        vector_hits = _vector_search_local(query=prd_text, top_k=8)
        return success_response(
            data={
                "audit_result": last_obj,
                "bug_hits": bug_hits[:12],
                "vector_hits": vector_hits,
            }
        )
    except Exception as e:
        logger.exception("api_prd_audit failed")
        return error_response(str(e), status_code=500)


# ---------- API：解析 PDF / Word ----------

@prd_audit_bp.route("/api/parse_pdf", methods=["POST"])
def api_parse_pdf():
    """解析 PDF 文件"""
    try:
        if "file" not in request.files:
            return error_response("未上传文件", status_code=400)
        file = request.files["file"]
        if not file.filename:
            return error_response("文件名为空", status_code=400)
        if not file.filename.lower().endswith(".pdf"):
            return error_response("请上传 PDF 文件", status_code=400)

        from io import BytesIO
        file_data = file.read()
        file_stream = BytesIO(file_data)
        text_chunks = []
        page_errors = 0
        try:
            import pypdf
            reader = pypdf.PdfReader(file_stream, strict=False)
            for page in reader.pages:
                try:
                    text_chunks.append(page.extract_text() or "")
                except Exception:
                    page_errors += 1
        except ImportError:
            try:
                import PyPDF2
                file_stream.seek(0)
                reader = PyPDF2.PdfReader(file_stream)
                for page in reader.pages:
                    try:
                        text_chunks.append(page.extract_text() or "")
                    except Exception:
                        page_errors += 1
            except ImportError:
                return error_response("未安装 PDF 解析库 (pypdf 或 PyPDF2)。请联系管理员: pip install pypdf", status_code=500)
        except Exception as e:
            logger.error("PDF解析错误: %s", e)
            return error_response("PDF 解析失败: " + str(e), status_code=500)

        text = "\n".join([c for c in text_chunks if isinstance(c, str) and c.strip()]).strip()
        if text:
            data = {"text": text}
            if page_errors > 0:
                data["warning"] = f"部分页面解析失败（{page_errors}页），已返回可提取内容。"
            return success_response(data=data)
        fallback_hint = (
            "【PDF解析提示】该PDF可能是扫描件/图片版，未提取到可复制文字。\n"
            "建议：1）导出可复制文本的PDF；2）上传Word(.docx)；3）直接粘贴PRD正文。"
        )
        return success_response(data={"text": fallback_hint, "warning": "未提取到可复制文本，已返回导入提示。"})
    except Exception as e:
        logger.exception("PDF 上传解析失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/parse_docx", methods=["POST"])
def api_parse_docx():
    """解析 Word (.docx) 文件"""
    try:
        if "file" not in request.files:
            return error_response("未上传文件", status_code=400)
        file = request.files["file"]
        if not file.filename:
            return error_response("文件名为空", status_code=400)
        if not file.filename.lower().endswith(".docx"):
            return error_response("请上传 .docx 文件", status_code=400)
        try:
            from docx import Document
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text.strip() for cell in row.cells)
        except ImportError:
            return error_response("未安装 python-docx。请联系管理员: pip install python-docx", status_code=500)
        except Exception as e:
            logger.error("DOCX 解析错误: %s", e)
            return error_response("Word 解析失败: " + str(e), status_code=500)
        if not text.strip():
            return error_response("未从 Word 中提取到文本", status_code=400)
        return success_response(data={"text": text})
    except Exception as e:
        logger.exception("Word 上传解析失败")
        return error_response(str(e), status_code=500)


# ---------- API：导出报告 Word ----------

@prd_audit_bp.route("/api/export_report_docx", methods=["POST"])
def api_export_report_docx():
    """将评审报告 Markdown 或 stage3_json 转为 Word 并返回文件"""
    try:
        data = request.get_json() or {}
        content = (data.get("content") or "").strip()
        stage3_json = data.get("stage3_json") if isinstance(data.get("stage3_json"), dict) else None
        if not content and not stage3_json:
            return error_response("报告内容为空", status_code=400)
        try:
            from docx import Document
            import re

            def _clean(s: str) -> str:
                """移除不兼容 XML 的控制字符，避免 python-docx 报错。"""
                if not isinstance(s, str):
                    s = str(s)
                # 保留常规可见字符与换行，去掉 0x00-0x08,0x0b-0x0c,0x0e-0x1f
                return re.sub(r"[\x00-\x08\x0b-\x0c\x0e-\x1f]", "", s)

            doc = Document()
            if stage3_json and isinstance(stage3_json, dict):
                title = _clean(stage3_json.get("report_title") or "【审计报告】PRD：工具扫描+人工复核版")
            else:
                title = "PRD 评审报告"
            doc.add_heading(title, 0)
            if stage3_json:
                summary = stage3_json.get("summary") or {}
                defects = stage3_json.get("defects") or []
                core_summary = stage3_json.get("core_risk_summary") or {}
                merged_issues = stage3_json.get("merged_issues") or []
                doc.add_heading("一、总体结论", level=1)
                doc.add_paragraph(_clean(f"质量评分：{summary.get('quality_score', '【PRD未说明】')}/10"))
                doc.add_paragraph(_clean(f"风险等级：{summary.get('risk_level', '【PRD未说明】')}"))
                doc.add_paragraph(_clean(f"主要问题：{summary.get('main_problem', '【PRD未说明】')}"))
                doc.add_heading("核心风险摘要", level=2)
                doc.add_paragraph(_clean(f"一句话总结：{core_summary.get('one_liner', '【PRD未说明】')}"))
                for item in pipeline._ensure_list(core_summary.get("top3")):
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("二、漏洞与风险清单", level=1)
                stats = stage3_json.get("scan_stats") or {}
                doc.add_paragraph(
                    _clean(
                        f"扫描来源统计：规则库 {stats.get('rule', 0)} 条，LLM {stats.get('llm', 0)} 条，混合 {stats.get('hybrid', 0)} 条"
                    )
                )
                if isinstance(merged_issues, list) and merged_issues:
                    doc.add_heading("合并后的核心问题", level=2)
                    for i, m in enumerate(merged_issues, start=1):
                        doc.add_heading(
                            _clean(f"核心问题{i}：{m.get('name', '【PRD未说明】')}（{m.get('risk_level', 'P2')}）"),
                            level=3,
                        )
                        doc.add_paragraph(
                            _clean(
                                f"涉及锚点：{'; '.join(pipeline._ensure_list(m.get('anchors'))) or '【PRD未说明】'}"
                            )
                        )
                        doc.add_paragraph(_clean(f"问题描述：{m.get('description', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"风险分析：{m.get('reason', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"建议：{m.get('suggestion', '【PRD未说明】')}"))
                if isinstance(defects, list) and defects:
                    doc.add_heading("详细漏洞清单", level=2)
                    for i, d in enumerate(defects, start=1):
                        doc.add_heading(_clean(f"漏洞{i}"), level=2)
                        doc.add_paragraph(_clean(f"模块：{d.get('module', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"类型：{d.get('type', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"风险等级：{d.get('risk_level', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"来源：{d.get('source', 'llm')}"))
                        doc.add_paragraph(
                            _clean(f"锚点：{d.get('anchor', d.get('module', '【PRD未说明】'))}")
                        )
                        doc.add_paragraph(_clean(f"描述：{d.get('description', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"原因：{d.get('reason', '【PRD未说明】')}"))
                        doc.add_paragraph(_clean(f"建议：{d.get('suggestion', '【PRD未说明】')}"))
                else:
                    doc.add_paragraph("未发现漏洞")
                doc.add_heading("三、测试重点", level=1)
                for item in pipeline._ensure_list(stage3_json.get("test_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("四、研发重点", level=1)
                for item in pipeline._ensure_list(stage3_json.get("dev_focus")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("五、项目风险", level=1)
                for item in pipeline._ensure_list(stage3_json.get("risks")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
                doc.add_heading("六、计划建议", level=1)
                for item in pipeline._ensure_list(stage3_json.get("plan")) or ["【PRD未说明】"]:
                    doc.add_paragraph(_clean(f"- {item}"))
            else:
                lines = content.replace("\r\n", "\n").split("\n")
                i = 0
                while i < len(lines):
                    line = _clean(lines[i])
                    stripped = line.strip()
                    if stripped.startswith("## "):
                        title = stripped[3:].strip()
                        doc.add_heading(title, level=1)
                    elif stripped.startswith("### "):
                        title = stripped[4:].strip()
                        doc.add_heading(title, level=2)
                    elif stripped:
                        doc.add_paragraph(line.rstrip())
                    else:
                        doc.add_paragraph()
                    i += 1
            import io
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(
                buf,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name="PRD评审报告.docx",
            )
        except ImportError:
            return error_response("未安装 python-docx。请联系管理员: pip install python-docx", status_code=500)
        except Exception as e:
            logger.exception("导出 Word 失败: %s", e)
            return error_response("导出 Word 失败: " + str(e), status_code=500)
    except Exception as e:
        logger.exception("export_report_docx 失败")
        return error_response(str(e), status_code=500)


# ---------- API：默认提示词、LLM 配置 ----------

@prd_audit_bp.route("/api/default_prd_prompt", methods=["GET"])
def api_get_default_prd_prompt():
    """返回 PRD 默认审计提示词（从 prd_audit 目录的 default 文件读取）"""
    try:
        if os.path.exists(DEFAULT_PRD_AUDIT_PROMPT_FILE):
            with open(DEFAULT_PRD_AUDIT_PROMPT_FILE, "r", encoding="utf-8") as f:
                s = f.read().strip()
                if s and len(s) > 100:
                    return success_response(data={"prompt": s})
        return success_response(data={"prompt": FALLBACK_PRD_PROMPT})
    except Exception as e:
        logger.warning("Failed to read default PRD prompt: %s", e)
        return success_response(data={"prompt": FALLBACK_PRD_PROMPT})


@prd_audit_bp.route("/api/llm_config", methods=["GET"])
def api_get_llm_config():
    """获取本模块 LLM 配置（modules/prd_audit_clone/llm_config.json）"""
    try:
        path = _pick_llm_config_path()
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                config = json.load(f)
                profiles = config.get("profiles") if isinstance(config.get("profiles"), dict) else {}
                default_profile = (config.get("default_profile") or config.get("llm_provider") or "deepseek").strip()
                active = profiles.get(default_profile) if isinstance(profiles.get(default_profile), dict) else None
                if active:
                    for k in ["llm_provider", "base_url", "api_key", "model"]:
                        if k in active and active.get(k) is not None:
                            config[k] = active.get(k)
                if config.get("api_key"):
                    config["api_key"] = config["api_key"][:3] + "****" + config["api_key"][-4:]
                if config.get("fallback_api_key"):
                    config["fallback_api_key"] = config["fallback_api_key"][:3] + "****" + config["fallback_api_key"][-4:]
                return success_response(data=config)
        return success_response(data={})
    except Exception as e:
        logger.exception("获取 LLM 配置失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/llm_config", methods=["POST"])
def api_save_llm_config():
    """保存本模块 LLM 配置（写入 modules/prd_audit_clone/llm_config.json）"""
    try:
        data = request.get_json() or {}
        save_path = _pick_llm_config_path()
        old_config = {}
        if os.path.exists(save_path):
            try:
                with open(save_path, "r", encoding="utf-8") as f:
                    old_config = json.load(f)
            except Exception:
                pass
        new_config = old_config.copy()
        profiles = new_config.get("profiles") if isinstance(new_config.get("profiles"), dict) else {}

        llm_provider = (data.get("llm_provider") or new_config.get("llm_provider") or "deepseek").strip()
        profile_key = llm_provider
        profile_old = profiles.get(profile_key) if isinstance(profiles.get(profile_key), dict) else {}
        profile_new = profile_old.copy()
        for k in ["llm_provider", "base_url", "model"]:
            if k in data:
                profile_new[k] = data.get(k)

        api_key = (data.get("api_key") or "").strip()
        if api_key and "****" not in api_key:
            profile_new["api_key"] = api_key
        elif "api_key" in profile_old:
            profile_new["api_key"] = profile_old.get("api_key", "")

        profiles[profile_key] = profile_new
        new_config["profiles"] = profiles
        new_config["default_profile"] = profile_key

        for k in ["llm_provider", "base_url", "model", "api_key"]:
            if k in profile_new and profile_new.get(k) is not None:
                new_config[k] = profile_new.get(k)

        fallback_key = (data.get("fallback_api_key") or "").strip()
        if fallback_key and "****" not in fallback_key:
            new_config["fallback_api_key"] = fallback_key
        elif "fallback_api_key" in old_config:
            new_config["fallback_api_key"] = old_config.get("fallback_api_key", "")

        for k in ["fallback_enabled", "fallback_provider", "fallback_base_url", "fallback_model"]:
            if k in data:
                new_config[k] = data.get(k)

        with open(save_path, "w", encoding="utf-8") as f:
            json.dump(new_config, f, ensure_ascii=False, indent=2)
        return success_response(message="配置已保存")
    except Exception as e:
        logger.exception("保存 LLM 配置失败")
        return error_response(str(e), status_code=500)


# ---------- API：保存为测试用例（写入 session，跳转用例管理） ----------

@prd_audit_bp.route("/api/prepare_save_to_cases", methods=["POST"])
def api_prepare_save_to_cases():
    """独立副本未集成平台用例管理：请手动复制报告。"""
    return error_response(
        "独立版（prd_audit_clone）未集成平台用例管理。请在本页复制报告 Markdown 后粘贴到用例系统。",
        status_code=501,
    )


@prd_audit_bp.route("/api/learning/status", methods=["GET"])
def api_learning_status():
    try:
        return success_response(data=get_learning_status())
    except Exception as e:
        logger.exception("获取学习状态失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/lane_stats", methods=["GET"])
def api_learning_lane_stats():
    try:
        limit = int(request.args.get("limit", 5000) or 5000)
        if limit < 1:
            limit = 1
        if limit > 20000:
            limit = 20000
        return success_response(data=get_learning_lane_stats(limit=limit))
    except Exception as e:
        logger.exception("获取分轨统计失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/quality_dashboard", methods=["GET"])
def api_learning_quality_dashboard():
    try:
        limit = int(request.args.get("limit", 5000) or 5000)
        if limit < 1:
            limit = 1
        if limit > 20000:
            limit = 20000
        return success_response(data=get_learning_quality_dashboard(limit=limit))
    except Exception as e:
        logger.exception("获取学习质量看板失败")
        return error_response(str(e), status_code=500)


# ---------- API：历史审计记录 (Snapshots) ----------

@prd_audit_bp.route("/api/history/snapshots", methods=["GET"])
def api_history_snapshots():
    """获取所有历史审计快照的简要列表"""
    try:
        from .audit_learning import INDEX_FILE
        if not os.path.exists(INDEX_FILE):
            return success_response(data={"snapshots": []})
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            idx = json.load(f)
        snaps = idx.get("snapshots", [])
        # 按时间倒序
        snaps.sort(key=lambda x: x.get("created_at", 0), reverse=True)
        return success_response(data={"snapshots": snaps})
    except Exception as e:
        logger.exception("获取历史记录失败")
        return error_response(f"获取历史记录失败: {e}", status_code=500)

@prd_audit_bp.route("/api/history/snapshot/<snapshot_id>", methods=["GET"])
def api_history_snapshot_detail(snapshot_id):
    """获取单个历史快照的完整详情"""
    try:
        from .audit_learning import SNAPSHOT_DIR
        file_path = os.path.join(SNAPSHOT_DIR, f"{snapshot_id}.json")
        if not os.path.exists(file_path):
            return error_response("未找到该历史记录", status_code=404)
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return success_response(data=data)
    except Exception as e:
        logger.exception("读取历史记录详情失败")
        return error_response(f"读取历史记录详情失败: {e}", status_code=500)


@prd_audit_bp.route("/api/learning/build_rule_draft", methods=["POST"])
def api_learning_build_rule_draft():
    try:
        data = request.get_json() or {}
        min_count = int(data.get("min_count", 2) or 2)
        max_new_rules = int(data.get("max_new_rules", 30) or 30)
        if min_count < 1:
            min_count = 1
        if max_new_rules < 1:
            max_new_rules = 1
        result = build_rule_draft_from_snapshots(min_count=min_count, max_new_rules=max_new_rules)
        return success_response(data=result, message="规则草案已生成")
    except Exception as e:
        logger.exception("生成规则草案失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/rule_candidates", methods=["GET"])
def api_learning_rule_candidates():
    try:
        return success_response(data=load_rule_candidates())
    except Exception as e:
        logger.exception("读取规则候选失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/apply_candidates", methods=["POST"])
def api_learning_apply_candidates():
    try:
        data = request.get_json() or {}
        selected = data.get("selected_rule_names")
        if not isinstance(selected, list):
            selected = []
        max_new_rules = int(data.get("max_new_rules", 100) or 100)
        if max_new_rules < 1:
            max_new_rules = 1
        if max_new_rules > 1000:
            max_new_rules = 1000
        result = apply_selected_candidates(selected_names=selected, max_new_rules=max_new_rules)
        return success_response(data=result, message="已生成可应用规则文件")
    except Exception as e:
        logger.exception("应用候选规则失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/publish_applied", methods=["POST"])
def api_learning_publish_applied():
    try:
        data = request.get_json() or {}
        create_backup = bool(data.get("create_backup", True))
        result = publish_applied_rules(create_backup=create_backup)
        return success_response(data=result, message="已发布为正式规则")
    except FileNotFoundError as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=400)
    except ValueError as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=400)
    except Exception as e:
        logger.exception("发布正式规则失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/generate_test_code", methods=["POST"])
def api_generate_test_code():
    """根据测试矩阵生成 Pytest 测试代码（流式）"""
    try:
        data = request.get_json() or {}
        test_matrix = data.get("test_matrix")
        
        if not test_matrix or not isinstance(test_matrix, dict):
            return error_response("缺少有效的 test_matrix 数据", status_code=400)

        # 构造 Prompt
        prompt = (
            "你是一个高级测试工程师。请根据以下【功能测试矩阵】JSON数据，"
            "编写一份完整的 Python Pytest 自动化测试脚本。\n\n"
            "要求：\n"
            "1. 使用 pytest 框架，包含 fixture Setup/Teardown。\n"
            "2. 每个测试用例对应矩阵中的一个 case_id。\n"
            "3. 代码中包含详细注释（模块、预期结果、判断依据）。\n"
            "4. 对于 '缺失' 或 '待确认' 的项，使用 @pytest.mark.skip 或 @pytest.mark.xfail 标记，并注明原因。\n"
            "5. 生成的代码必须是可执行的 Python 代码块（不要 markdown ```python 标记，直接输出代码）。\n\n"
            f"测试矩阵数据：\n{json.dumps(test_matrix, ensure_ascii=False, indent=2)}"
        )

        from utils.llm_client import load_llm_config, stream_chat_content
        config_path = _pick_llm_config_path()
        try:
            llm_config = load_llm_config(config_path)
            # 强制检查 API Key
            if not (llm_config.get("api_key") or "").strip():
                raise ValueError("LLM API Key 未配置")
        except Exception as e:
            return error_response(f"生成代码需要配置大模型 API Key: {e}", status_code=400)

        def generate():
            try:
                # 调用流式生成
                for chunk in stream_chat_content(llm_config, prompt):
                    if chunk:
                        # 简单过滤 markdown 标记，防止前端渲染混乱
                        # chunk = chunk.replace("```python", "").replace("```", "")
                        yield chunk
            except Exception as e:
                logger.error(f"Test code generation failed: {e}")
                yield f"\n# 生成失败: {str(e)}"

        return Response(stream_with_context(generate()), mimetype="text/plain")

    except Exception as e:
        logger.exception("测试代码生成接口异常")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/export_feature_xmind", methods=["POST"])
def api_export_feature_xmind():
    """根据功能节点导出 .xmind 文件（本地生成，无需大模型）"""
    try:
        data = request.get_json() or {}
        nodes = data.get("nodes")
        if not isinstance(nodes, list) or not nodes:
            return error_response("nodes 不能为空", status_code=400)
        xmind_buf = _build_xmind_file(nodes)
        return send_file(
            xmind_buf,
            mimetype="application/vnd.xmind.workbook",
            as_attachment=True,
            download_name="PRD功能导图.xmind",
        )
    except Exception as e:
        logger.exception("导出 XMind 失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/knowledge_cards", methods=["GET", "POST"])
def api_knowledge_cards():
    """获取或保存功能知识卡片"""
    try:
        if request.method == "POST":
            data = request.get_json() or {}
            items = data.get("items")
            saved = _save_knowledge_cards(items)
            return success_response(data={"items": saved, "count": len(saved)}, message="能力库已保存")

        domain = (request.args.get("domain") or "").strip().lower()
        cards = _load_knowledge_cards()
        if domain and domain not in ["ktv", "ktv点歌系统", "all"]:
            cards = []
        return success_response(data={"items": cards, "count": len(cards)})
    except Exception as e:
        logger.exception("获取功能知识卡片失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/knowledge_cards/import", methods=["POST"])
def api_knowledge_cards_import():
    try:
        file = request.files.get("file")
        if file:
            filename = file.filename.lower()
            if filename.endswith(".csv") or filename.endswith(".txt"):
                content = file.read().decode("utf-8-sig")
                reader = csv.DictReader(io.StringIO(content))
                items = []
                for row in reader:
                    if not row.get("name"): continue
                    item = {
                        "capability_id": row.get("capability_id") or f"CAP_{uuid.uuid4().hex[:8]}",
                        "domain": row.get("domain", ""),
                        "name": row.get("name", ""),
                        "desc": row.get("desc", ""),
                        "actors": [x for x in (row.get("actors") or "").split("|") if x],
                        "pre_conditions": [x for x in (row.get("pre_conditions") or "").split("|") if x],
                        "post_conditions": [x for x in (row.get("post_conditions") or "").split("|") if x],
                        "rules": [x for x in (row.get("rules") or "").split("|") if x],
                        "required_evidence": [x for x in (row.get("required_evidence") or "").split("|") if x],
                        "bad_patterns": [x for x in (row.get("bad_patterns") or "").split("|") if x],
                        "suggestion_template": row.get("suggestion_template", "")
                    }
                    items.append(item)
                if not items:
                    return error_response("CSV 解析为空或格式错误", status_code=400)
                
                existing = _load_knowledge_cards()
                existing.extend(items)
                saved = _save_knowledge_cards(existing)
                return success_response(data={"items": saved, "count": len(saved)}, message=f"成功导入 {len(items)} 条")

        data = request.get_json() or {}
        items = data.get("items")
        if not isinstance(items, list):
            return error_response("items 必须为数组", status_code=400)
            
        existing = _load_knowledge_cards()
        existing.extend(items)
        saved = _save_knowledge_cards(existing)
        return success_response(data={"items": saved, "count": len(saved)}, message=f"成功导入 {len(items)} 条")
    except Exception as e:
        logger.exception("导入能力库失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/knowledge_cards/export", methods=["GET"])
def api_knowledge_cards_export():
    try:
        cards = _load_knowledge_cards()
        format_type = request.args.get("format", "json").lower()
        
        if format_type == "csv":
            sio = io.StringIO()
            writer = csv.writer(sio)
            writer.writerow(["capability_id", "domain", "name", "desc", "actors", "pre_conditions", "post_conditions", "rules", "required_evidence", "bad_patterns", "suggestion_template"])
            for c in cards:
                writer.writerow([
                    c.get("capability_id", ""),
                    c.get("domain", ""),
                    c.get("name", ""),
                    c.get("desc", ""),
                    "|".join(c.get("actors", [])),
                    "|".join(c.get("pre_conditions", [])),
                    "|".join(c.get("post_conditions", [])),
                    "|".join(c.get("rules", [])),
                    "|".join(c.get("required_evidence", [])),
                    "|".join(c.get("bad_patterns", [])),
                    c.get("suggestion_template", "")
                ])
            buf = io.BytesIO(sio.getvalue().encode("utf-8-sig"))
            return send_file(
                buf,
                mimetype="text/csv",
                as_attachment=True,
                download_name="knowledge_cards.csv",
            )
            
        buf = io.BytesIO()
        buf.write(json.dumps({"items": cards}, ensure_ascii=False, indent=2).encode("utf-8"))
        buf.seek(0)
        return send_file(
            buf,
            mimetype="application/json",
            as_attachment=True,
            download_name="knowledge_cards.json",
        )
    except Exception as e:
        logger.exception("导出能力库失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/backups", methods=["GET"])
def api_learning_backups():
    try:
        limit = int(request.args.get("limit", 20) or 20)
        if limit < 1:
            limit = 1
        if limit > 200:
            limit = 200
        return success_response(data={"items": list_rule_backups(limit=limit)})
    except Exception as e:
        logger.exception("获取规则备份失败")
        return error_response(str(e), status_code=500)


@prd_audit_bp.route("/api/learning/rollback_backup", methods=["POST"])
def api_learning_rollback_backup():
    try:
        data = request.get_json() or {}
        backup_file_name = (data.get("backup_file_name") or "").strip()
        create_backup = bool(data.get("create_backup", True))
        if not backup_file_name:
            return error_response("backup_file_name 不能为空", status_code=400)
        result = rollback_rules_from_backup(backup_file_name=backup_file_name, create_backup=create_backup)
        return success_response(data=result, message="规则库已回滚")
    except FileNotFoundError as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=404)
    except ValueError as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=400)
    except Exception as e:
        logger.exception("回滚规则库失败")
        return error_response(str(e), status_code=500)
